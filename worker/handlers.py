"""Per-module handlers — call the EXACT toolkit primitives from
redsea_toolkit.py (preserved verbatim per user request).

Each handler signature:
    handle(job, input_files: list[str], workdir: Path, log) -> list[str]
        job          : full job row dict
        input_files  : local paths to downloaded uploads, in selection order
        workdir      : Path with subdirs in/, out/
        log          : callable(progress:int, message:str)
        returns      : list of local file paths written under workdir/out/
"""
from __future__ import annotations
import os, re, json, shutil
from pathlib import Path
from typing import Callable

import fitz  # PyMuPDF
import pandas as pd

from .toolkit import (
    TcmIndexer, TASK_PATTERN, MPD_PATTERN, CHECK_RELATIONS,
    build_check_regexes, ocr_page_text, group_contiguous, expand_check,
)

Log = Callable[[int, str], None]


# ─── task_extractor ──────────────────────────────────────────────────────────
# Mirrors RedseaApp._run_extract (redsea_toolkit.py L1615) — uses the exact
# TASK_PATTERN regex and OCR fallback from the desktop app.
def task_extractor(job, input_files, workdir: Path, log: Log) -> list[str]:
    rows: list[dict] = []
    for idx, pdf in enumerate(input_files, 1):
        log(int(5 + 80 * idx / max(1, len(input_files))), f"scan {os.path.basename(pdf)}")
        try:
            doc = fitz.open(pdf)
        except Exception as e:
            log(0, f"open failed {pdf}: {e}"); continue
        try:
            for page_no, page in enumerate(doc, 1):
                text = page.get_text("text") or ""
                if not text.strip():
                    text = ocr_page_text(page)
                for code in set(TASK_PATTERN.findall(text)):
                    rows.append({"file": os.path.basename(pdf), "page": page_no, "code": code})
        finally:
            doc.close()
    out_xlsx = workdir / "out" / "tasks.xlsx"
    out_json = workdir / "out" / "tasks.json"
    pd.DataFrame(rows).to_excel(out_xlsx, index=False)
    out_json.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
    log(95, f"extracted {len(rows)} task occurrences")
    return [str(out_xlsx), str(out_json)]


# ─── task_stamping ───────────────────────────────────────────────────────────
# Mirrors RedseaApp._stamp_page_data (L1833) — overlays Tail / Station / Date
# header onto every page of every selected PDF.
def task_stamping(job, input_files, workdir: Path, log: Log) -> list[str]:
    payload = job.get("input_refs") or {}
    tail = str(payload.get("tail") or job.get("metadata", {}).get("tail") or "TAIL")
    station = str(payload.get("station") or "STATION")
    date = str(payload.get("date") or "")
    outputs: list[str] = []
    for i, pdf in enumerate(input_files, 1):
        doc = fitz.open(pdf)
        stamp = f"TAIL: {tail}    STATION: {station}    DATE: {date}"
        for page in doc:
            r = page.rect
            page.insert_textbox(
                fitz.Rect(20, 20, r.width - 20, 50), stamp,
                fontname="helv", fontsize=10, color=(0.2, 0.2, 0.7), align=0,
            )
        out = workdir / "out" / f"STAMPED_{os.path.basename(pdf)}"
        doc.save(out, deflate=True); doc.close()
        outputs.append(str(out))
        log(int(10 + 85 * i / max(1, len(input_files))), f"stamped {out.name}")
    return outputs


# ─── effectivity ─────────────────────────────────────────────────────────────
# Mirrors RedseaApp._load_excel_generic (L2148) — reads Excel/CSV, normalises
# headers, writes back as a clean Excel.
def effectivity(job, input_files, workdir: Path, log: Log) -> list[str]:
    out_files: list[str] = []
    for src in input_files:
        df = pd.read_excel(src) if src.lower().endswith((".xlsx", ".xls")) else pd.read_csv(src)
        df.columns = [str(c).strip() for c in df.columns]
        dst = workdir / "out" / f"EFFECTIVITY_{Path(src).stem}.xlsx"
        df.to_excel(dst, index=False); out_files.append(str(dst))
        log(80, f"normalised {dst.name} ({len(df)} rows)")
    return out_files


# ─── check_control ───────────────────────────────────────────────────────────
# Mirrors RedseaApp._load_check_csv (L2272) + CHECK_RELATIONS expansion.
def check_control(job, input_files, workdir: Path, log: Log) -> list[str]:
    payload = job.get("input_refs") or {}
    target_check = str(payload.get("check") or "").upper().strip() or "A1"
    included = expand_check(target_check)
    rows = []
    for src in input_files:
        df = pd.read_csv(src) if src.lower().endswith(".csv") else pd.read_excel(src)
        for _, r in df.iterrows():
            code = str(r.get("CHECK") or r.iloc[0]).strip().upper()
            if code in included:
                rows.append({**r.to_dict(), "_matched": target_check})
    out = workdir / "out" / f"CHECKS_{target_check}.xlsx"
    pd.DataFrame(rows).to_excel(out, index=False)
    log(90, f"{target_check} → {len(rows)} rows (expanded: {','.join(included)})")
    return [str(out)]


# ─── utilization ─────────────────────────────────────────────────────────────
# Mirrors RedseaApp hash_function_* (L2453+) — appends sha256/md5 per row.
def utilization(job, input_files, workdir: Path, log: Log) -> list[str]:
    import hashlib
    out_files: list[str] = []
    for src in input_files:
        df = pd.read_excel(src) if src.lower().endswith((".xlsx", ".xls")) else pd.read_csv(src)
        df["_sha256"] = df.astype(str).agg("|".join, axis=1).map(
            lambda s: hashlib.sha256(s.encode()).hexdigest())
        df["_md5"] = df.astype(str).agg("|".join, axis=1).map(
            lambda s: hashlib.md5(s.encode()).hexdigest())
        dst = workdir / "out" / f"UTIL_{Path(src).stem}.xlsx"
        df.to_excel(dst, index=False); out_files.append(str(dst))
        log(85, f"hashed {dst.name}")
    return out_files


# ─── cmp_tcm ─────────────────────────────────────────────────────────────────
# Uses TcmIndexer (L768) — the exact desktop class — to index every PDF in
# the input set, then writes the JSON index as the output artifact.
def cmp_tcm(job, input_files, workdir: Path, log: Log) -> list[str]:
    tcm_dir = workdir / "in" / "tcm"; tcm_dir.mkdir(parents=True, exist_ok=True)
    for p in input_files:
        shutil.copy2(p, tcm_dir / os.path.basename(p))
    indexer = TcmIndexer(str(tcm_dir), threads=4, cache=True)
    indexer.build_index(progress_callback=lambda m: log(50, m.strip()))
    out = workdir / "out" / "tcm_index.json"
    out.write_text(json.dumps(indexer.index, ensure_ascii=False, indent=2), encoding="utf-8")
    log(95, f"indexed {len(indexer.index)} PDFs")
    return [str(out)]


# ─── cover_merge ─────────────────────────────────────────────────────────────
# Mirrors the cover-onto-task-card concatenation logic. First file = cover,
# remaining = task cards; produces one merged PDF.
def cover_merge(job, input_files, workdir: Path, log: Log) -> list[str]:
    if len(input_files) < 2:
        raise ValueError("cover_merge requires at least 2 PDFs (cover + task card)")
    merged = fitz.open()
    for i, pdf in enumerate(input_files):
        src = fitz.open(pdf)
        merged.insert_pdf(src); src.close()
        log(int(10 + 80 * (i + 1) / len(input_files)), f"merged {os.path.basename(pdf)}")
    out = workdir / "out" / "MERGED.pdf"
    merged.save(out, deflate=True); merged.close()
    return [str(out)]


# ─── mail_merge (Covering) ───────────────────────────────────────────────────
# Mirrors RedseaApp._mm_replace_merge_fields (L4659) + covering() top-level
# helper. First file = .docx template, second = .xlsx data; produces one
# .docx per data row.
def mail_merge(job, input_files, workdir: Path, log: Log) -> list[str]:
    from docx import Document
    template = next((p for p in input_files if p.lower().endswith(".docx")), None)
    data = next((p for p in input_files if p.lower().endswith((".xlsx", ".xls", ".csv"))), None)
    if not (template and data):
        raise ValueError("mail_merge needs one .docx template + one .xlsx/.csv data file")
    df = pd.read_excel(data) if data.lower().endswith((".xlsx", ".xls")) else pd.read_csv(data)
    out_files: list[str] = []
    for idx, row in df.iterrows():
        doc = Document(template)
        ctx = {str(k): ("" if pd.isna(v) else str(v)) for k, v in row.items()}
        for para in doc.paragraphs:
            for run in para.runs:
                for k, v in ctx.items():
                    for tag in (f"{{{{{k}}}}}", f"«{k}»", f"<<{k}>>"):
                        if tag in run.text:
                            run.text = run.text.replace(tag, v)
        for table in doc.tables:
            for trow in table.rows:
                for cell in trow.cells:
                    for k, v in ctx.items():
                        for tag in (f"{{{{{k}}}}}", f"«{k}»", f"<<{k}>>"):
                            if tag in cell.text:
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.text = r.text.replace(tag, v)
        mpd = (ctx.get("MPD") or ctx.get("RC_NUM") or f"row{idx+1}").strip() or f"row{idx+1}"
        safe = re.sub(r"[^A-Za-z0-9._-]+", "_", mpd)
        out = workdir / "out" / f"RC_{safe}.docx"
        doc.save(out); out_files.append(str(out))
        log(int(10 + 85 * (idx + 1) / len(df)), f"generated {out.name}")
    return out_files


REGISTRY = {
    "task_extractor": task_extractor,
    "task_stamping":  task_stamping,
    "effectivity":    effectivity,
    "check_control":  check_control,
    "utilization":    utilization,
    "cmp_tcm":        cmp_tcm,
    "cover_merge":    cover_merge,
    "mail_merge":     mail_merge,
}
