"""Per-module handlers — call the EXACT toolkit primitives from
redsea_toolkit.py (preserved verbatim per user request).

Each handler signature:
    handle(job, input_files: list[str], workdir: Path, log) -> list[str]
        job          : full job row dict
        input_files  : local paths to downloaded uploads, in selection order
        workdir      : Path with subdirs in/, out/
        log          : callable(progress:int, message:str)
        returns      : list of local file paths written under workdir/out/
"""
from __future__ import annotations
import os, re, json, shutil
from pathlib import Path
from typing import Callable

import fitz  # PyMuPDF
import pandas as pd

from .toolkit import (
    TcmIndexer, TASK_PATTERN, MPD_PATTERN, CHECK_RELATIONS,
    build_check_regexes, ocr_page_text, group_contiguous, expand_check,
    normalize_check_code,
)

Log = Callable[[int, str], None]


# ─── task_extractor ──────────────────────────────────────────────────────────
# Mirrors RedseaApp._run_extract (redsea_toolkit.py L1615) — full logic:
# 1. Find PDF by first-two-digits of the task code
# 2. Search for related subtasks in that PDF
# 3. Extract related tasks to separate PDFs
# 4. Fallback: extract pages containing the base task code
def task_extractor(job, input_files, workdir: Path, log: Log) -> list[str]:
    payload = job.get("input_refs") or {}
    code = str(payload.get("task_code") or "").strip()

    # ── Legacy mode: no task_code → scan all PDFs for all codes (original web behavior)
    if not code:
        rows: list[dict] = []
        for idx, pdf in enumerate(input_files, 1):
            log(int(5 + 80 * idx / max(1, len(input_files))), f"scan {os.path.basename(pdf)}")
            try:
                doc = fitz.open(pdf)
            except Exception as e:
                log(0, f"open failed {pdf}: {e}"); continue
            try:
                for page_no, page in enumerate(doc, 1):
                    text = page.get_text("text") or ""
                    if not text.strip():
                        text = ocr_page_text(page)
                    for found_code in set(TASK_PATTERN.findall(text)):
                        rows.append({"file": os.path.basename(pdf), "page": page_no, "code": found_code})
            finally:
                doc.close()
        out_xlsx = workdir / "out" / "tasks.xlsx"
        out_json = workdir / "out" / "tasks.json"
        pd.DataFrame(rows).to_excel(out_xlsx, index=False)
        out_json.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
        log(95, f"extracted {len(rows)} task occurrences")
        return [str(out_xlsx), str(out_json)]

    # ── Full mode: task_code provided → mirror _run_extract logic from app2.py
    first_two_digits = code.split('-')[0] if '-' in code else code[:2]
    log(10, f"Looking for PDF files starting with: {first_two_digits}")

    # Find matching PDF among uploaded files
    matching = [p for p in input_files if os.path.basename(p).startswith(first_two_digits)]
    if not matching:
        log(20, f"No PDF starting with '{first_two_digits}' among uploads, scanning all...")
        matching = input_files  # fallback: scan all uploaded PDFs

    pdf_path = matching[0]
    log(15, f"Using PDF: {os.path.basename(pdf_path)}")

    out_dir = workdir / "out"
    out_dir.mkdir(parents=True, exist_ok=True)
    outputs: list[str] = []

    # Step 1: Search for related subtasks
    log(20, f"Searching for related tasks to '{code}'...")
    m = re.match(r"^(\d{2}-\d{2,3})-\d{2}$", code.strip())
    related_tasks: list[str] = []
    if m:
        prefix = m.group(1) + "-"
        try:
            doc = fitz.open(pdf_path)
            found = set()
            for page in doc:
                text = page.get_text("text") or ""
                if not text.strip():
                    text = ocr_page_text(page)
                for m2 in TASK_PATTERN.finditer(text):
                    cand = m2.group(0)
                    if cand.startswith(prefix) and cand != code:
                        found.add(cand)
            related_tasks = sorted(found)
            doc.close()
        except Exception as e:
            log(25, f"Error scanning for related tasks: {e}")

    if related_tasks:
        log(30, f"Found {len(related_tasks)} related tasks: {', '.join(related_tasks)}")
        all_codes = [code] + related_tasks
        try:
            doc = fitz.open(pdf_path)
            for i, c in enumerate(all_codes):
                # Scan pages for this code (simple string match, skip index pages)
                page_matches = []
                for pn in range(doc.page_count):
                    text = doc[pn].get_text("text") or ""
                    upper = text.upper()
                    is_index = False
                    if any(mk in upper for mk in ("INDEX", "TABLE OF CONTENTS", "LIST OF EFFECTIVE PAGES")):
                        if len(TASK_PATTERN.findall(text)) > 10:
                            is_index = True
                    if c in text and not is_index:
                        page_matches.append(pn)
                if not page_matches:
                    continue
                out_doc = fitz.open()
                for run in group_contiguous(page_matches):
                    out_doc.insert_pdf(doc, from_page=run[0], to_page=run[-1])
                out_name = f"{c.replace('/', '_')}_related.pdf"
                out_path = out_dir / out_name
                out_doc.save(out_path, deflate=True)
                out_doc.close()
                outputs.append(str(out_path))
                log(30 + int(50 * (i + 1) / len(all_codes)), f"Extracted: {out_name}")
            doc.close()
        except Exception as e:
            log(50, f"Error extracting related tasks: {e}")

        if outputs:
            log(90, f"Successfully extracted {len(outputs)} related task PDF(s)")
            return outputs

    # Step 2: Fallback — extract base task pages only
    log(60, f"Searching for base task '{code}' in {os.path.basename(pdf_path)}...")
    try:
        doc = fitz.open(pdf_path)
        page_matches = []
        for pn in range(doc.page_count):
            text = doc[pn].get_text("text") or ""
            upper = text.upper()
            is_index = False
            if any(mk in upper for mk in ("INDEX", "TABLE OF CONTENTS", "LIST OF EFFECTIVE PAGES")):
                if len(TASK_PATTERN.findall(text)) > 10:
                    is_index = True
            if code in text and not is_index:
                page_matches.append(pn)

        if not page_matches:
            log(80, f"No matching pages found for '{code}'")
            # Still output an empty result file
            out_json = out_dir / "no_results.json"
            out_json.write_text(json.dumps({"task_code": code, "message": "No pages found"}, indent=2), encoding="utf-8")
            doc.close()
            return [str(out_json)]

        log(75, f"Found {len(page_matches)} pages with '{code}'")
        out_doc = fitz.open()
        for run in group_contiguous(page_matches):
            out_doc.insert_pdf(doc, from_page=run[0], to_page=run[-1])
        out_name = f"{code.replace('/', '_')}_extracted.pdf"
        out_path = out_dir / out_name
        out_doc.save(out_path, deflate=True)
        out_doc.close()
        doc.close()
        outputs.append(str(out_path))
        log(90, f"Extracted base task: {out_name}")
    except Exception as e:
        log(80, f"Error extracting base task: {e}")

    log(95, f"done — {len(outputs)} file(s)")
    return outputs


# ─── task_stamping ───────────────────────────────────────────────────────────
# Mirrors RedseaApp._stamp_page_data (redsea_toolkit.py L1833) — overlays Tail,
# Airline Card No (RC number derived from BOEING CARD NO), Station, and Date.
def task_stamping(job, input_files, workdir: Path, log: Log) -> list[str]:
    payload = job.get("input_refs") or {}
    tail = str(payload.get("tail") or job.get("metadata", {}).get("tail") or "TAIL")
    station = str(payload.get("station") or "STATION")
    date = str(payload.get("date") or "")
    outputs: list[str] = []
    
    out_dir = workdir / "out"
    out_dir.mkdir(parents=True, exist_ok=True)
    
    for i, pdf in enumerate(input_files, 1):
        try:
            doc = fitz.open(pdf)
        except Exception as e:
            log(0, f"Error opening {pdf}: {e}"); continue
            
        for page_num, page in enumerate(doc):
            try:
                airline_card_no = "RC"
                found_number = False
                
                # --- Method 1: Location-based text search ---
                search_terms = ["BOEING CARD NO.", "BOEING CARD NO"]
                search_instances = []
                for term in search_terms:
                    search_instances = page.search_for(term)
                    if search_instances:
                        break
                
                if search_instances:
                    label_rect = search_instances[0]
                    search_area = fitz.Rect(label_rect.x1 - 5, label_rect.y0 - 5, label_rect.x1 + 300, label_rect.y1 + 5)
                    extracted_text = page.get_text("text", clip=search_area).strip()
                    
                    if extracted_text:
                        boeing_match = re.search(r'([\d-]+)', extracted_text)
                        if boeing_match:
                            full_boeing_no = boeing_match.group(1).strip()
                            parts = full_boeing_no.split('-')
                            if len(parts) >= 3:
                                airline_card_no = f"RC{parts[0]}-{parts[1]}-{parts[2]}"
                                found_number = True
                            elif len(parts) >= 2:
                                airline_card_no = f"RC{parts[0]}-{parts[1]}"
                                found_number = True
                
                # --- Method 2: Fallback to full page text search ---
                if not found_number:
                    page_text = page.get_text()
                    boeing_match = re.search(r"BOEING\s+CARD\s+NO\.?\s*([\d-]+)", page_text, re.IGNORECASE)
                    if boeing_match:
                        full_boeing_no = boeing_match.group(1).strip()
                        parts = full_boeing_no.split('-')
                        if len(parts) >= 3:
                            airline_card_no = f"RC{parts[0]}-{parts[1]}-{parts[2]}"
                        elif len(parts) >= 2:
                            airline_card_no = f"RC{parts[0]}-{parts[1]}"
                
                # --- Stamping ---
                items_to_stamp = [
                    ("TAIL NUMBER", tail), 
                    ("AIRLINE CARD NO", airline_card_no), 
                    ("STATION", station),
                    ("DATE", date)
                ]
                for label, value in items_to_stamp:
                    if not value: continue
                    instances = page.search_for(label)
                    for inst in instances:
                        page.insert_text((inst.x0, inst.y1 + 10), value, fontsize=10, color=(0, 0, 0))
            except Exception as e:
                log(50, f"Error stamping page {page_num+1} of {os.path.basename(pdf)}: {e}")
                
        out = out_dir / f"STAMPED_{os.path.basename(pdf)}"
        doc.save(out, deflate=True)
        doc.close()
        outputs.append(str(out))
        log(int(10 + 85 * i / max(1, len(input_files))), f"stamped {out.name}")
        
    return outputs


# ─── effectivity ─────────────────────────────────────────────────────────────
# Mirrors RedseaApp._load_excel_generic (L2148) — reads Excel/CSV, normalises
# headers, writes back as a clean Excel.
def effectivity(job, input_files, workdir: Path, log: Log) -> list[str]:
    payload = job.get("input_refs") or {}
    if payload.get("data_source") == "db":
        log(50, "Database source is selected. (Pending DB Migrations implementation)")
        raise NotImplementedError("Database source for Effectivity is not yet implemented.")
        
    out_files: list[str] = []
    for src in input_files:
        df = pd.read_excel(src) if src.lower().endswith((".xlsx", ".xls")) else pd.read_csv(src)
        df.columns = [str(c).strip() for c in df.columns]
        dst = workdir / "out" / f"EFFECTIVITY_{Path(src).stem}.xlsx"
        df.to_excel(dst, index=False); out_files.append(str(dst))
        log(80, f"normalised {dst.name} ({len(df)} rows)")
    return out_files


# ─── check_control ───────────────────────────────────────────────────────────
# Mirrors RedseaApp._load_check_csv (L2272) + CHECK_RELATIONS expansion.
def check_control(job, input_files, workdir: Path, log: Log) -> list[str]:
    payload = job.get("input_refs") or {}
    if payload.get("data_source") == "db":
        log(50, "Database source is selected. (Pending DB Migrations implementation)")
        raise NotImplementedError("Database source for Check Control is not yet implemented.")
        
    target_check = str(payload.get("check") or "").upper().strip() or "A1"
    included = expand_check(target_check)
    rows = []
    for src in input_files:
        df = pd.read_csv(src) if src.lower().endswith(".csv") else pd.read_excel(src)
        for _, r in df.iterrows():
            code = str(r.get("CHECK") or r.iloc[0]).strip().upper()
            if code in included:
                rows.append({**r.to_dict(), "_matched": target_check})
    out = workdir / "out" / f"CHECKS_{target_check}.xlsx"
    pd.DataFrame(rows).to_excel(out, index=False)
    log(90, f"{target_check} → {len(rows)} rows (expanded: {','.join(included)})")
    return [str(out)]


# ─── utilization ─────────────────────────────────────────────────────────────
# Mirrors RedseaApp hash_function_* (L2453+) — appends sha256/md5 per row.
def utilization(job, input_files, workdir: Path, log: Log) -> list[str]:
    payload = job.get("input_refs") or {}
    if payload.get("data_source") == "db":
        log(50, "Database source is selected. (Pending DB Migrations implementation)")
        raise NotImplementedError("Database source for Utilization is not yet implemented.")
        
    import hashlib
    out_files: list[str] = []
    for src in input_files:
        df = pd.read_excel(src) if src.lower().endswith((".xlsx", ".xls")) else pd.read_csv(src)
        df["_sha256"] = df.astype(str).agg("|".join, axis=1).map(
            lambda s: hashlib.sha256(s.encode()).hexdigest())
        df["_md5"] = df.astype(str).agg("|".join, axis=1).map(
            lambda s: hashlib.md5(s.encode()).hexdigest())
        dst = workdir / "out" / f"UTIL_{Path(src).stem}.xlsx"
        df.to_excel(dst, index=False); out_files.append(str(dst))
        log(85, f"hashed {dst.name}")
    return out_files


def cmp_tcm(job, input_files, workdir: Path, log: Log) -> list[str]:
    payload = job.get("input_refs") or {}
    check_code = str(payload.get("check") or "").upper().strip()

    pdfs = [p for p in input_files if p.lower().endswith(".pdf")]
    excels = [p for p in input_files if p.lower().endswith((".xlsx", ".xls", ".csv"))]

    tcm_dir = workdir / "in" / "tcm"
    tcm_dir.mkdir(parents=True, exist_ok=True)
    for p in pdfs:
        shutil.copy2(p, tcm_dir / os.path.basename(p))

    log(10, f"Building TCM index from {len(pdfs)} PDFs...")
    indexer = TcmIndexer(str(tcm_dir), threads=4, cache=True)
    indexer.build_index(progress_callback=lambda m: log(50, m.strip()))
    
    out_files = []
    
    # If no excel/check provided, just return the index JSON
    if not excels or not check_code:
        out = workdir / "out" / "tcm_index.json"
        out.write_text(json.dumps(indexer.index, ensure_ascii=False, indent=2), encoding="utf-8")
        log(95, f"indexed {len(indexer.index)} PDFs (no Excel provided)")
        return [str(out)]

    log(60, f"Extracting tasks for check {check_code} from Excel...")
    excel_path = excels[0]
    try:
        df = pd.read_excel(excel_path) if excel_path.lower().endswith((".xlsx", ".xls")) else pd.read_csv(excel_path)
    except Exception as e:
        log(60, f"Error reading excel: {e}")
        return []

    # Column 24 holds the check code, column 0 the task code.
    #
    # App2 matches the check by EQUALITY on two tiers (app2.py:3247-3260, and
    # redsea_toolkit.py:3330-3341): first the whitespace-stripped, upper-cased
    # cell against the same form of the target, then _normalize_check_code of
    # both sides. It is never a containment test. A substring match turns every
    # check code into a prefix filter -- asking for "A1" silently absorbs "A10"
    # and "A11", emitting task cards the operator never requested.
    df_str = df.astype(str)
    tasks = set()
    normalized_target = re.sub(r"\s+", "", check_code).upper()
    normalized_target_variant = normalize_check_code(check_code)
    if df_str.shape[1] > 24:
        for idx in range(len(df_str)):
            cell_raw = str(df_str.iloc[idx, 24]).strip()
            matches = (
                re.sub(r"\s+", "", cell_raw).upper() == normalized_target
                or normalize_check_code(cell_raw) == normalized_target_variant
            )
            if not matches:
                continue
            task = str(df_str.iloc[idx, 0]).strip()
            if task and task != "nan":
                tasks.add(task)
    
    # Also add the new subtask expansion logic here to match the desktop
    expanded_tasks = set(tasks)
    for task in tasks:
        subtasks = indexer.find_related_subtasks(task)
        expanded_tasks.update(subtasks)

    sorted_tasks = sorted(list(expanded_tasks))
    log(70, f"Found {len(sorted_tasks)} tasks/subtasks for {check_code}")

    # For each task, extract from TCM
    for i, task in enumerate(sorted_tasks):
        log(70 + int(20 * i / max(1, len(sorted_tasks))), f"Extracting {task}")
        pdf_path, run = indexer.find_best_occurrence_for_task(task)
        if not pdf_path or not run:
            continue
            
        start, end = run[0], run[1] if isinstance(run, (list,tuple)) and len(run)>=2 else (run[0], run[-1])
        try:
            src = fitz.open(pdf_path)
            out_doc = fitz.open()
            out_doc.insert_pdf(src, from_page=start, to_page=end)
            
            out_pdf = workdir / "out" / f"{task}.pdf"
            out_doc.save(out_pdf, deflate=True)
            out_files.append(str(out_pdf))
            out_doc.close()
            src.close()
        except Exception as e:
            log(70 + int(20 * i / max(1, len(sorted_tasks))), f"Failed to extract {task}: {e}")

    log(95, f"Generated {len(out_files)} task PDFs")
    return out_files


# ─── cover_merge ─────────────────────────────────────────────────────────────
# Mirrors the cover-onto-task-card concatenation logic. First file = cover,
# remaining = task cards; produces one merged PDF.
def cover_merge(job, input_files, workdir: Path, log: Log) -> list[str]:
    if len(input_files) < 2:
        raise ValueError("cover_merge requires at least 2 PDFs (cover + task card)")
    merged = fitz.open()
    for i, pdf in enumerate(input_files):
        src = fitz.open(pdf)
        merged.insert_pdf(src); src.close()
        log(int(10 + 80 * (i + 1) / len(input_files)), f"merged {os.path.basename(pdf)}")
    out = workdir / "out" / "MERGED.pdf"
    merged.save(out, deflate=True); merged.close()
    return [str(out)]


# ─── mail_merge (Covering) ───────────────────────────────────────────────────
# Mirrors RedseaApp._mm_replace_merge_fields (L4659) + covering() top-level
# helper. First file = .docx template, second = .xlsx data; produces one
# .docx per data row.
def mail_merge(job, input_files, workdir: Path, log: Log) -> list[str]:
    from docx import Document
    template = next((p for p in input_files if p.lower().endswith(".docx")), None)
    data = next((p for p in input_files if p.lower().endswith((".xlsx", ".xls", ".csv"))), None)
    if not (template and data):
        raise ValueError("mail_merge needs one .docx template + one .xlsx/.csv data file")
    df = pd.read_excel(data) if data.lower().endswith((".xlsx", ".xls")) else pd.read_csv(data)
    out_files: list[str] = []
    for idx, row in df.iterrows():
        doc = Document(template)
        ctx = {str(k): ("" if pd.isna(v) else str(v)) for k, v in row.items()}
        for para in doc.paragraphs:
            for run in para.runs:
                for k, v in ctx.items():
                    for tag in (f"{{{{{k}}}}}", f"«{k}»", f"<<{k}>>"):
                        if tag in run.text:
                            run.text = run.text.replace(tag, v)
        for table in doc.tables:
            for trow in table.rows:
                for cell in trow.cells:
                    for k, v in ctx.items():
                        for tag in (f"{{{{{k}}}}}", f"«{k}»", f"<<{k}>>"):
                            if tag in cell.text:
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.text = r.text.replace(tag, v)
        mpd = (ctx.get("MPD") or ctx.get("RC_NUM") or f"row{idx+1}").strip() or f"row{idx+1}"
        safe = re.sub(r"[^A-Za-z0-9._-]+", "_", mpd)
        out = workdir / "out" / f"RC_Card_{safe}.docx"
        doc.save(out); out_files.append(str(out))
        log(int(10 + 85 * (idx + 1) / len(df)), f"generated {out.name}")
    return out_files


REGISTRY = {
    "task_extractor": task_extractor,
    "task_stamping":  task_stamping,
    "effectivity":    effectivity,
    "check_control":  check_control,
    "utilization":    utilization,
    "cmp_tcm":        cmp_tcm,
    "cover_merge":    cover_merge,
    "mail_merge":     mail_merge,
}
