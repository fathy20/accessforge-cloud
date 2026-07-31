import shutil
import tempfile
import unittest
from pathlib import Path

import pandas as pd
from docx import Document


class TestMailMergeParity(unittest.TestCase):
    mpd = "27-001-00"
    title = "Deterministic RC card"

    def setUp(self):
        self.tmpdir = Path(tempfile.mkdtemp(prefix="mail_merge_parity_"))
        self.template = self.tmpdir / "template.docx"
        self.data = self.tmpdir / "data.xlsx"
        self.workdir = self.tmpdir / "web-work"
        (self.workdir / "out").mkdir(parents=True)

        document = Document()
        document.add_paragraph("RC title: «TITLE»")
        document.add_paragraph("Unrelated template content")
        document.save(self.template)
        pd.DataFrame([{"MPD": self.mpd, "TITLE": self.title}]).to_excel(self.data, index=False)

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_one_row_guillemet_rendering_matches_desktop(self):
        from worker import toolkit
        from worker.handlers import mail_merge

        desktop = Document(self.template)
        toolkit.rt.RedseaApp._mm_manual_replace(object(), desktop, {"TITLE": self.title})
        desktop_text = [paragraph.text for paragraph in desktop.paragraphs]

        web_outputs = mail_merge({}, [str(self.template), str(self.data)], self.workdir, lambda *_args: None)
        self.assertEqual(len(web_outputs), 1)
        web_output = Path(web_outputs[0])
        self.assertEqual(web_output.name, f"RC_Card_{self.mpd}.docx")

        web = Document(web_output)
        self.assertEqual([paragraph.text for paragraph in web.paragraphs], desktop_text)
        self.assertEqual(web.paragraphs[0].text, f"RC title: {self.title}")
        self.assertEqual(web.paragraphs[1].text, "Unrelated template content")

        self.assertNotIn("«TITLE»", web.paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
