# -*- coding: utf-8 -*-
"""
REDSEA Toolkit – Aviation Maintenance Task Management System

A comprehensive tool for managing aviation maintenance tasks including:
- Task extraction from PDF documents
- Task stamping and processing with OCR support
- Effectivity and TCM (Task Card Management)
- Check control and utilization tracking 
- CMP/TCM task card generation with indexing
- Cover merge functionality
- Mail Merge (Covering) - Automated Word-Excel integration for RC Cards

Required Dependencies:
  - PyMuPDF (fitz) - PDF processing
  - pytesseract - OCR functionality
  - Pillow (PIL) - Image processing
  - pandas - Data manipulation
  - customtkinter - Modern UI components
  - python-docx - Word document manipulation
  - openpyxl - Excel file manipulation
  
Note: Tesseract OCR engine must be installed separately on the system.
Set TESSERACT_CMD below if tesseract is not in your system PATH.
"""

# Standard library imports
import os
import re
import sys
import glob
import time
import json
import pickle
import io
import zipfile
import xml.etree.ElementTree as ET
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed

# Third-party imports
import fitz  # PyMuPDF for PDF processing
import pytesseract  # OCR functionality
from PIL import Image  # Image processing
import pandas as pd  # Data manipulation

# Optional imports for Mail Merge functionality
try:
    from docx import Document  # Word document manipulation
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Mail Merge (Covering) module will be disabled.")
    print("To enable Mail Merge: pip install python-docx")

try:
    from mailmerge import MailMerge
    MAILMERGE_AVAILABLE = True
except ImportError:
    MAILMERGE_AVAILABLE = False
    print("Warning: docx-mailmerge not installed. Please install: pip install docx-mailmerge")

try:
    import openpyxl  # Excel file manipulation
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("Warning: openpyxl not installed. Some Excel features may be limited.")
    print("To enable full Excel support: pip install openpyxl")

# GUI imports
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import customtkinter as ctk

# ================== APPLICATION CONFIGURATION ==================

# Application settings
APP_TITLE = "REDSEA Toolkit – Aviation Maintenance Task Management"
APP_MIN_W, APP_MIN_H = 1400, 800

# Assets paths configuration
ASSETS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
ICON_PATH = os.path.join(ASSETS_PATH, "REDSEA Airlines Logo.ico")
LOGO_PATH = os.path.join(ASSETS_PATH, "REDSEA Airlines Logo.png")
BACKGROUND_PATH = os.path.join(ASSETS_PATH, "login_bg.jpg")

# Tesseract OCR configuration
# If tesseract binary is not in PATH, set the full path here (Windows example):
# TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
TESSERACT_CMD = None
if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# Regular expression patterns for task identification
# Match 3-part codes (e.g., 27-054-00) AND 4-part subtasks (e.g., 27-054-00-01)
TASK_PATTERN = re.compile(r"\b\d{2}-\d{2,3}-\d{2}(?:-\d{2})?\b")
MPD_PATTERN = re.compile(r"(\d{2}-\d{2,3}-\d{2}(?:-\d{2})?)")

# UI Color scheme - Modern Professional Theme
COLOR_PRIMARY = "#1a365d"      # Deep blue
COLOR_ACCENT = "#2d5a87"       # Medium blue  
COLOR_OK = "#38a169"           # Modern green
COLOR_WARN = "#ed8936"         # Modern orange
COLOR_DANGER = "#e53e3e"       # Modern red
COLOR_ACCENT_DARK = "#1a365d"  # Dark blue
COLOR_BG_LIGHT = "#f7fafc"     # Light background
COLOR_TEXT_DARK = "#2d3748"    # Dark text
COLOR_CARD_BG = "#ffffff"      # Card background
COLOR_BORDER = "#e2e8f0"        # Border color
COLOR_HOVER = "#4a5568"        # Hover color

# Application Control Variables
MINI_LAUNCHER_MODE = False  # Changed to False to prevent immediate exit
TARGET_MODULE = None        # Module to open directly
APP_INSTANCE = None         # Global app instance for better management
TASK_EXTRACTOR_THREADS = 3  # Parallel workers for Task Extractor page scans

# Aviation maintenance check relationships
# Defines which sub-checks are included when a main check is performed
CHECK_RELATIONS = {
    # A-Check relationships (Letter checks)
    "A1": ["A1"], "A2": ["A1", "A2"], "A3": ["A1", "A3"], 
    "A4": ["A1", "A2", "A4"], "A5": ["A1", "A5"], 
    "A6": ["A1", "A2", "A3", "A6"], "A7": ["A1"], 
    "A8": ["A1", "A2", "A4", "A8"], "A9": ["A1", "A3"], 
    "A10": ["A1", "A2", "A5", "A10"], "A11": ["A1"],
    
    # Calendar-based checks (time intervals)
    "120DY": ["120DY"],  # 120 day check
    "240DY": ["120DY", "240DY"],  # 240 day check (includes 120 day)
    "12MO": ["120DY", "12MO"],  # 12 month check
    "16MO": ["120DY", "240DY", "12MO", "16MO"],  # 16 month check (major)
    "2000FC": ["2000FC"],  # 2000 flight cycle check
    
    # C-Check relationships (heavy maintenance)
    "C1": ["C1"], "C2": ["C1", "C2"], "C3": ["C1", "C3"], 
    "C4": ["C1", "C2", "C4"], "C5": ["C1", "C5"], 
    "C6": ["C1", "C2", "C3", "C6"]
}

# ================== MINI LAUNCHER SYSTEM ==================

class MiniLauncherManager:
    """
    Professional Mini Launcher Manager for Quick Access
    Provides a clean, efficient interface for module navigation
    """
    
    def __init__(self, parent_app):
        self.parent_app = parent_app
        self.mini_window = None
        self.is_mini_mode = False
        self.current_hover = None
        self.animation_running = False
    
    def show_mini_launcher(self):
        """Display the professional mini launcher interface"""
        if self.mini_window:
            self.mini_window.destroy()
        
        # Create main launcher window
        self.mini_window = ctk.CTkToplevel(self.parent_app)
        self.mini_window.title("REDSEA Airlines - Quick Access Launcher")
        self.mini_window.geometry("380x650")
        self.mini_window.resizable(False, False)
        
        # Set application icon for mini launcher
        try:
            if os.path.exists(ICON_PATH):
                self.mini_window.iconbitmap(ICON_PATH)
        except Exception as e:
            print(f"Warning: Could not set mini launcher icon: {e}")
        
        # Enhanced window properties
        self.mini_window.attributes('-topmost', True)
        self.mini_window.after(500, lambda: self.mini_window.attributes('-topmost', False))
        
        # Center the window on screen
        self.center_mini_window()
        
        # Create Mini Launcher interface
        self.create_mini_interface()
        
        # Ensure staying in foreground
        self.mini_window.transient(self.parent_app)
        self.mini_window.grab_set()
        
        self.is_mini_mode = True
    
    def center_mini_window(self):
        """Center the Mini window on screen"""
        self.mini_window.update_idletasks()
        width = 380
        height = 650
        x = (self.mini_window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.mini_window.winfo_screenheight() // 2) - (height // 2)
        self.mini_window.geometry(f'{width}x{height}+{x}+{y}')
    
    def create_mini_interface(self):
        """Create professional mini launcher interface"""
        # Add background image behind mini UI if available
        try:
            if hasattr(self.parent_app, 'background_image') and self.parent_app.background_image:
                # keep reference to avoid garbage collection
                self.mini_bg_label = ctk.CTkLabel(
                    self.mini_window,
                    image=self.parent_app.background_image,
                    text=""
                )
                self.mini_bg_label.place(x=0, y=0, relwidth=1, relheight=1)
                self.mini_bg_label.lower()
        except Exception as e:
            print(f"Warning: Could not set mini launcher background: {e}")

        # Main container with gradient-like appearance
        main_frame = ctk.CTkFrame(
            self.mini_window, 
            fg_color="transparent",
            corner_radius=15
        )
        main_frame.pack(fill="both", expand=True, padx=15, pady=15)
        
        # Header section with logo area
        header_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=25, pady=(25, 15))
        
        # Add logo if available
        if hasattr(self.parent_app, 'logo_image') and self.parent_app.logo_image:
            logo_label = ctk.CTkLabel(
                header_frame,
                image=self.parent_app.logo_image,
                text=""
            )
            logo_label.pack(pady=(0, 10))
        
        # Application title with enhanced typography
        title_label = ctk.CTkLabel(
            header_frame,
            text="REDSEA Airlines",
            font=ctk.CTkFont(size=22, weight="bold"),
            text_color=COLOR_PRIMARY
        )
        title_label.pack()
        
        # Professional subtitle
        subtitle_label = ctk.CTkLabel(
            header_frame,
            text="Aviation Maintenance Toolkit",
            font=ctk.CTkFont(size=11, weight="normal"),
            text_color=COLOR_TEXT_DARK
        )
        subtitle_label.pack(pady=(2, 0))
        
        # Separator line
        separator = ctk.CTkFrame(main_frame, height=2, fg_color=COLOR_PRIMARY)
        separator.pack(fill="x", padx=40, pady=(15, 25))
        
        # Create scrollable frame for buttons
        scrollable_frame = ctk.CTkScrollableFrame(
            main_frame,
            fg_color="transparent",
            height=380
        )
        scrollable_frame.pack(fill="both", expand=True, padx=25)
        
        # Professional module buttons with icons
        modules = [
            ("📊", "Full System Access", "Complete system interface", lambda: self.switch_to_full_system(), COLOR_PRIMARY),
            ("📋", "Task Extractor", "Extract tasks from PDF documents", lambda: self.open_module("task_extract"), COLOR_ACCENT),
            ("🖨️", "Task Stamping", "Process and stamp task documents", lambda: self.open_module("task_stamping"), COLOR_OK),
            ("📈", "Effectivity Reports", "Generate effectivity and TCM reports", lambda: self.open_module("effectivity"), COLOR_WARN),
            ("✅", "Check Control", "Monitor maintenance checks", lambda: self.open_module("check_control"), COLOR_TEXT_DARK),
            ("📦", "CMP/TCM Tasks", "Manage task card generation", lambda: self.open_module("cmp_tcm"), COLOR_ACCENT_DARK),
            ("✉️", "Mail Merge (Covering)", "Automated Word-Excel Mail Merge for RC Cards", lambda: self.open_module("mail_merge"), "#2563eb")
        ]
        
        for icon, title, description, command, color in modules:
            self.create_professional_button(scrollable_frame, icon, title, description, command, color)
        
        # Separator line
        separator = ctk.CTkFrame(scrollable_frame, height=2, fg_color=COLOR_BG_LIGHT)
        separator.pack(fill="x", pady=15)
        
        # Exit button
        exit_btn = ctk.CTkButton(
            scrollable_frame,
            text="❌ Exit Application",
            command=self.close_mini_launcher,
            height=45,
            fg_color=COLOR_DANGER,
            hover_color="#e04040",
            font=ctk.CTkFont(size=12, weight="bold"),
            text_color="white"
        )
        exit_btn.pack(fill="x", pady=(10, 20))
        
        # Developer credit
        credit_label = ctk.CTkLabel(
            main_frame,
            text="BY: Eng Fathy Ahmed  and Khaled Haggag",
            font=ctk.CTkFont(size=11, weight="bold"),
            text_color=COLOR_PRIMARY
        )
        credit_label.pack(pady=(5, 10))
        
        # Version info
        version_label = ctk.CTkLabel(
            main_frame,
            text="Version 2.0 - Quick Launcher",
            font=ctk.CTkFont(size=9),
            text_color="gray"
        )
        version_label.pack(pady=(0, 10))
    
    def create_professional_button(self, parent, icon, title, description, command, color):
        """Create professional module access button"""
        # Button container for enhanced layout
        btn_container = ctk.CTkFrame(parent, fg_color="transparent")
        btn_container.pack(fill="x", pady=6)  # Increase spacing between buttons
        
        # Main button with professional styling
        btn = ctk.CTkButton(
            btn_container,
            text=f"{icon}  {title}",
            command=command,
            height=55,  # Increase button height slightly
            fg_color=color,
            hover_color=self.darken_color(color),
            font=ctk.CTkFont(size=12, weight="bold"),
            text_color="white",
            corner_radius=10,
            anchor="w"
        )
        btn.pack(fill="x", padx=5)
        
        # Add hover effects
        self.add_button_hover_effects(btn, color)
        
        return btn
    
    def add_button_hover_effects(self, button, original_color):
        """Add professional hover effects to buttons"""
        def on_enter(event):
            if not self.animation_running:
                button.configure(fg_color=self.lighten_color(original_color))
        
        def on_leave(event):
            if not self.animation_running:
                button.configure(fg_color=original_color)
        
        button.bind("<Enter>", on_enter, add="+")
        button.bind("<Leave>", on_leave, add="+")
    
    def lighten_color(self, color):
        """Create lighter version of color for hover effect"""
        color_map = {
            COLOR_PRIMARY: "#2a5a7a",
            COLOR_ACCENT: "#ff8660",
            COLOR_OK: "#2bc98a",
            COLOR_WARN: "#ffd15b",
            COLOR_TEXT_DARK: "#404040",
            COLOR_ACCENT_DARK: "#e6653a"
        }
        return color_map.get(color, "#555555")
    
    def darken_color(self, color):
        """جعل اللون أغمق للـ hover effect"""
        color_map = {
            COLOR_PRIMARY: "#164054",
            COLOR_ACCENT: COLOR_ACCENT_DARK,
            COLOR_OK: "#0f8c73",
            COLOR_WARN: "#d39e00",
            COLOR_TEXT_DARK: "#1a1a1a",
            COLOR_ACCENT_DARK: "#b8471c"
        }
        return color_map.get(color, "#555555")
    
    def open_module(self, module_name):
        """Open specific module in full system with smooth transition"""
        global TARGET_MODULE, MINI_LAUNCHER_MODE
        
        # Set target module and prepare for transition
        TARGET_MODULE = module_name
        MINI_LAUNCHER_MODE = False
        
        # Directly open the module without animation
        self.complete_transition()
    
    def animate_transition_to_full(self):
        """Smooth transition animation to full system"""
        self.animation_running = True
        
        # Fade out effect simulation
        if self.mini_window:
            self.mini_window.attributes('-alpha', 0.8)
            self.mini_window.after(100, lambda: self.mini_window.attributes('-alpha', 0.6))
            self.mini_window.after(200, lambda: self.mini_window.attributes('-alpha', 0.4))
            self.mini_window.after(300, lambda: self.mini_window.attributes('-alpha', 0.2))
            self.mini_window.after(400, self.complete_transition)
    
    def complete_transition(self):
        """Complete the transition to full system"""
        # Show full system with target module FIRST to prevent mainloop exit
        if hasattr(self.parent_app, 'show_full_system'):
            self.parent_app.show_full_system()
        else:
            self.parent_app.configure_full_mode()
            
        if self.mini_window:
            self.mini_window.destroy()
            self.mini_window = None
        
        self.is_mini_mode = False
        self.animation_running = False
            
        # Apply direct navigation to the target module
        global TARGET_MODULE
        if TARGET_MODULE:
            self.parent_app.open_target_module(TARGET_MODULE)
            
        # Ensure window is properly focused and maximized
        self.parent_app.focus_force()
        self.parent_app.lift()
        self.parent_app.state('zoomed')  # Ensure maximized state
    
    def switch_to_full_system(self):
        """Professional transition to full system interface"""
        global MINI_LAUNCHER_MODE, TARGET_MODULE
        
        MINI_LAUNCHER_MODE = False
        TARGET_MODULE = None  # No specific module, just open full system
        
        self.animate_transition_to_full()
    
    def close_mini_launcher(self):
        """Professional closure of mini launcher"""
        if self.mini_window:
            # Smooth close animation
            self.mini_window.attributes('-alpha', 0.5)
            self.mini_window.after(150, self.mini_window.destroy)
        
        # Graceful application exit
        if hasattr(self.parent_app, 'quit'):
            self.parent_app.after(200, self.parent_app.quit)
        else:
            self.parent_app.after(200, lambda: sys.exit(0))

def expand_check(check_name: str):
    """
    Expand a check name to include all related sub-checks.
    
    Args:
        check_name: The main check code (e.g., 'A8', 'C4', '240DY')
        
    Returns:
        List of check codes including the main check and all sub-checks
    """
    k = (check_name or "").strip().upper()
    return CHECK_RELATIONS.get(k, [k]) if k else []

# ================== IMAGE UTILITIES ==================

def load_background_image(path: str, size: tuple = None, opacity: float = 0.15):
    """Load and prepare background image as CTkImage with alpha opacity."""
    try:
        from PIL import Image
        
        if not os.path.exists(path):
            return None
            
        # Open and resize image
        img = Image.open(path)
        if size:
            img = img.resize(size, Image.Resampling.LANCZOS)
        
        # Ensure RGBA
        if img.mode != 'RGBA':
            img = img.convert('RGBA')
        
        # Apply alpha opacity
        try:
            r, g, b, a = img.split()
            a = a.point(lambda p: int(p * max(0.0, min(1.0, opacity))))
            img.putalpha(a)
        except Exception:
            pass
        
        # Return CTkImage for better compatibility with customtkinter
        return ctk.CTkImage(light_image=img, dark_image=img, size=img.size)
    except Exception as e:
        print(f"Warning: Could not load background image: {e}")
        return None

def load_logo_image(path: str, size: tuple = (60, 60)):
    """Load and prepare logo image"""
    try:
        from PIL import Image
        
        if not os.path.exists(path):
            return None
            
        img = Image.open(path)
        img = img.resize(size, Image.Resampling.LANCZOS)
        
        # Use CTkImage for better compatibility
        return ctk.CTkImage(light_image=img, dark_image=img, size=size)
    except Exception as e:
        print(f"Warning: Could not load logo image: {e}")
        return None

# ================== UTILITY FUNCTIONS ==================

def safe_make_dir(path: str):
    """Safely create directory structure, ignoring if it already exists"""
    os.makedirs(path, exist_ok=True)
    
def is_pdf(p: str) -> bool:
    """Check if file path points to a valid PDF file"""
    return bool(p) and p.lower().endswith(".pdf") and os.path.isfile(p)
    
def walk_pdfs_in_dir(folder: str):
    """Recursively find all PDF files in a directory"""
    if not folder or not os.path.isdir(folder):
        return
    for p in sorted(glob.glob(os.path.join(folder, "**", "*.pdf"), recursive=True)):
        yield p
        
def unique_path(path: str) -> str:
    """Generate a unique file path by adding number suffix if file exists"""
    base, ext = os.path.splitext(path)
    n = 1
    new = path
    while os.path.exists(new):
        new = f"{base}({n}){ext}"
        n += 1
    return new
    
def group_contiguous(pages):
    """Group contiguous page numbers into ranges"""
    if not pages:
        return []
    pages = sorted(set(pages))
    runs = []
    cur = [pages[0]]
    for i in range(1, len(pages)):
        if pages[i] == pages[i-1] + 1:
            cur.append(pages[i])
        else:
            runs.append(cur)
            cur = [pages[i]]
    runs.append(cur)
    return runs


def covering(template_path: str, data_path: str, output_path_prefix: str):
    """
    Interactive covering: read CSV by column indices and fill Word template using docxtpl.
    Uses positional column indices (A=0, B=1, etc.) to map data from Excel to Word placeholders.

    Args:
        template_path: Path to RC.docx template
        data_path: Path to RC INDEX.xlsx - RC.csv
        output_path_prefix: Prefix for output file (final name: {prefix}_{MPD}.docx)

    Returns:
        str: Path to saved file on success, None on failure
    """
    # Import here to keep module imports stable
    from docxtpl import DocxTemplate

    # Mapping: Excel column index (0-based) to Word placeholder name
    placeholder_index = {
        # Index : (Excel Col, Word Placeholder)
        0:  ('A/SEQ', 'SEQ'),           # Routine TASK Card & RELATED TASK
        1:  ('B/MPD', 'MPD'),           # RC. #
        2:  ('C/DATE', 'DATE'),         # RC. DATE:
        3:  ('D/TITLE', 'TITLE'),       # TITLE
        4:  ('E/MHR', 'MHR'),           # EST MHRS
        5:  ('F/WO', 'WO'),            # W/O
        6:  ('G/AC', 'AC_REG'),        # A/C REG.
        7:  ('H/ACSN', 'AC_MSN'),      # A/C MSN:
        8:  ('I/ZONE', 'ZONE'),        # AREA/ZONE:
        9:  ('J/ACCESS', 'ACCESS'),     # ACCESS PANELS:
        10: ('K/CYC', 'CYCLE'),        # NO cycle
        11: ('L/FHS', 'HOURS'),        # TOTAL HOURS:
        13: ('N/SOURC', 'SOURCE'),     # SOURCE
        14: ('O/CRIT', 'CRITICAL'),    # CRITICAL TASK
        15: ('P/RII', 'RII'),          # RII TASK
        16: ('Q/OTHER', 'OTHER'),      # OTHER
        17: ('R/CMP', 'CMP')           # CMP APPROVAL#
    }

    # Validate files exist
    if not os.path.isfile(template_path):
        print(f"Error: template not found: {template_path}")
        return None
    if not os.path.isfile(data_path):
        print(f"Error: data file not found: {data_path}")
        return None

    # Ask user for MPD
    mpd_input = input("Enter MPD (MPD ITEM NUMBER): ").strip()
    if not mpd_input:
        print("No MPD entered. Aborting.")
        return None

    # Read CSV (all as strings, no NA conversion)
    try:
        df = pd.read_csv(data_path, dtype=str, keep_default_na=False)
    except Exception as e:
        print(f"Failed to read CSV '{data_path}': {e}")
        return None

    # Ensure we have enough columns
    if df.shape[1] <= 1:  # Need at least 2 columns (need column B/MPD)
        print("Error: Data file must have at least 2 columns")
        return None

    # Search MPD in column B (index 1)
    col_series = df.iloc[:, 1].astype(str).str.strip()
    mask = col_series == mpd_input
    if not mask.any():
        # Try case-insensitive
        mask = col_series.str.upper() == mpd_input.upper()
        if not mask.any():
            print(f"MPD '{mpd_input}' not found in column B.")
            print("\nAvailable MPD values in column B:")
            available = col_series.unique()
            for i, v in enumerate(available, 1):
                if i > 10:  # Show first 10
                    print("...")
                    break
                print(f"  {i}. {v}")
            return None

    # Get matching row
    matched = df[mask]
    if len(matched) > 1:
        print(f"Warning: {len(matched)} rows matched MPD '{mpd_input}'. Using first match.")

    row = matched.iloc[0].tolist()

    # Build template context with all variations
    context = {}
    print("\n=== Data mapped from Excel columns ===")
    for idx, (col_name, ph_name) in placeholder_index.items():
        # Get value from Excel column safely
        value = ''
        try:
            if idx < len(row):
                v = row[idx]
                value = '' if pd.isna(v) else str(v).strip()
        except Exception:
            value = ''

        # Store value under simple name
        context[ph_name] = value
        
        # Also store under legacy names if needed
        if ph_name == 'MPD':
            context['RC_NUM'] = value
        elif ph_name == 'DATE':
            context['RC_DATE'] = value
        elif ph_name == 'MHR':
            context['EST_MHRS'] = value
        elif ph_name == 'AC_REG':
            context['AC_REG'] = value
            context['AC'] = value
        elif ph_name == 'AC_MSN':
            context['AC_MSN'] = value
            context['ACSN'] = value

        print(f"{col_name:<10} → {ph_name:<12}: {value}")

    # Load and render template
    try:
        doc = DocxTemplate(template_path)
        doc.render(context)
    except Exception as e:
        print(f"\nError rendering template: {e}")
        return None

    # Save with MPD in filename
    safe_mpd = "".join(ch if (ch.isalnum() or ch in '-_.') else '_' for ch in mpd_input)
    output_filename = f"{output_path_prefix}_{safe_mpd}.docx"
    try:
        doc.save(output_filename)
        print(f"\n✓ Successfully saved: {output_filename}")
        return output_filename
    except Exception as e:
        print(f"\nError saving file: {e}")
        return None

# OCR helper: convert page to PIL Image (png) and run pytesseract
def page_to_image(page, zoom=2.0):
    # zoom>1 increases resolution for better OCR
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img_bytes = pix.tobytes("png")
    return Image.open(io.BytesIO(img_bytes))

def ocr_page_text(page, lang="eng"):
    try:
        img = page_to_image(page, zoom=2.0)
        text = pytesseract.image_to_string(img, lang=lang)
        return text or ""
    except Exception:
        return ""

# Build regexes for a check code (handles A/C, C, and calendar checks robustly)
def build_check_regexes(check_code: str):
    code = (check_code or "").strip().upper()
    regexes = []
    if not code:
        return [re.compile(r"$^")]  # match nothing
    # A/C or single-letter checks like A1, C3
    if re.match(r"^[A-Z]\d+$", code):
        # match exact token (A1), and also forms like "A CHECK", "A-Check", "A1 CHECK"
        letter = code[0]
        regexes.append(re.compile(rf"\b{re.escape(code)}\b", re.IGNORECASE))
        regexes.append(re.compile(rf"\b{re.escape(letter)}\s*CHECK\b", re.IGNORECASE))
        regexes.append(re.compile(rf"\b{re.escape(code)}\s*CHECK\b", re.IGNORECASE))
    # Calendar checks like 120DY, 240DY, 12MO, 16MO
    elif re.match(r"^\d+(DY|MO)$", code):
        m = re.match(r"^(\d+)(DY|MO)$", code)
        if not m:
            regexes.append(re.compile(re.escape(code), re.IGNORECASE))
            return regexes
        num, unit = m.group(1), m.group(2)
        if unit == "DY":
            # common variants
            patterns = [
                fr"\b{num}\s*DY\b",
                fr"\b{num}\s*DAY\b",
                fr"\b{num}\s*DAYS\b",
                fr"\b{num}DY\b",
                fr"\b{num}DAY\b",
                fr"EVERY\s+{num}\s+DAY",
                fr"\b{num}\s*DAY\s*CHECK\b",
                fr"\b{num}\s*DAY(S)?\b"
            ]
        else:  # MO
            patterns = [
                fr"\b{num}\s*MO\b",
                fr"\b{num}\s*MONTH\b",
                fr"\b{num}\s*MONTHS\b",
                fr"\b{num}MO\b",
                fr"EVERY\s+{num}\s+MONTH",
                fr"\b{num}\s*MONTH\s*CHECK\b",
                fr"\b{num}\s*MO\b"
            ]
        regexes = [re.compile(p, re.IGNORECASE) for p in patterns]
        # also a fallback to literal code
        regexes.append(re.compile(re.escape(code), re.IGNORECASE))
    else:
        # generic fallback
        regexes.append(re.compile(re.escape(code), re.IGNORECASE))
        # also allow "X CHECK" where X is the letter part (for cases)
        if len(code) == 2 and code[0].isalpha() and code[1].isdigit():
            regexes.append(re.compile(rf"\b{re.escape(code[0])}\s*CHECK\b", re.IGNORECASE))
    return regexes

# ---------- Index structure ----------
INDEX_FILENAME_JSON = ".tcm_index.json"
INDEX_FILENAME_PKL = ".tcm_index.pkl"

class TcmIndexer:
    """
    Builds and loads index for a TCM folder. 
    Supports multithreaded scanning for improved performance.
    """
    
    def __init__(self, tcm_folder, threads=4, cache=True):
        if not tcm_folder or not os.path.isdir(tcm_folder):
            raise ValueError(f"Invalid TCM folder path: {tcm_folder}")
            
        self.tcm_folder = tcm_folder
        self.threads = max(1, threads)  # Ensure at least 1 thread
        self.cache = cache
        self.index = {}
        self._loaded_from_cache = False
        self.stop_requested = False

    def cache_path_json(self):
        return os.path.join(self.tcm_folder, INDEX_FILENAME_JSON)
    def cache_path_pkl(self):
        return os.path.join(self.tcm_folder, INDEX_FILENAME_PKL)

    def try_load_cache(self):
        if not self.cache: return False
        j=self.cache_path_json(); p=self.cache_path_pkl()
        try:
            if os.path.exists(j):
                with open(j,"r",encoding="utf-8") as f: self.index=json.load(f)
                self._loaded_from_cache=True; return True
            if os.path.exists(p):
                with open(p,"rb") as f: self.index=pickle.load(f)
                self._loaded_from_cache=True; return True
        except Exception:
            pass
        return False

    def save_cache(self):
        try:
            with open(self.cache_path_pkl(),"wb") as f:
                pickle.dump(self.index, f, protocol=pickle.HIGHEST_PROTOCOL)
        except Exception:
            try:
                with open(self.cache_path_json(),"w",encoding="utf-8") as f:
                    json.dump(self.index, f, ensure_ascii=False)
            except Exception:
                pass

    def scan_single_pdf(self, pdf_path):
        """Return mapping task->list of [start,end] runs for this pdf."""
        mapping = {}
        try:
            doc = fitz.open(pdf_path)
        except Exception:
            return mapping
        hits_by_task = {}
        try:
            for i, page in enumerate(doc):
                try:
                    text = page.get_text("text") or ""
                except Exception:
                    text = ""
                # if scanned page, try OCR for indexing (optional: can skip to speed)
                if not text.strip():
                    try:
                        text = ocr_page_text(page)
                    except Exception:
                        text = ""
                if not text: continue
                for m in TASK_PATTERN.findall(text):
                    hits_by_task.setdefault(m, []).append(i)
        finally:
            try: doc.close()
            except Exception: pass

        for t, pages in hits_by_task.items():
            runs = group_contiguous(pages)
            mapping[t] = [[r[0], r[-1]] for r in runs]
        return mapping

    def build_index(self, progress_callback=None):
        """Scan all PDFs in tcm_folder using ThreadPoolExecutor."""
        self.stop_requested = False
        self.index = {}
        pdfs = list(walk_pdfs_in_dir(self.tcm_folder))
        if progress_callback:
            progress_callback(f"Found {len(pdfs)} PDF(s) to index\n")
        if not pdfs:
            return
        with ThreadPoolExecutor(max_workers=self.threads) as ex:
            futures = {ex.submit(self.scan_single_pdf, p): p for p in pdfs}
            for i, fut in enumerate(as_completed(futures), 1):
                if self.stop_requested:
                    if progress_callback:
                        progress_callback(f"\n❌ Indexing stopped by user at {i}/{len(pdfs)} files\n")
                    # Cancel remaining futures if possible
                    for f in [f for f in futures if not f.done()]:
                        f.cancel()
                    break
                    
                p = futures[fut]
                try:
                    mapping = fut.result()
                    if mapping:
                        self.index[p] = mapping
                        if progress_callback:
                            progress_callback(f"Indexed: {os.path.basename(p)} -> {len(mapping)} task codes\n")
                    else:
                        if progress_callback:
                            progress_callback(f"Indexed: {os.path.basename(p)} -> 0 tasks\n")
                except Exception as e:
                    if progress_callback:
                        progress_callback(f"Indexing failed for {p}: {e}\n")
        
        if not self.stop_requested:
            self.save_cache()
            if progress_callback:
                progress_callback(f"\n✅ Indexing completed! Saved index with {len(self.index)} PDFs\n")
        if progress_callback:
            progress_callback("Indexing complete and cached.\n")

    def find_best_occurrence_for_task(self, task_code):
        """
        From index: choose the pdf+range that has the longest contiguous range for this task.
        Returns (pdf_path, [start,end]) or (None, []).
        """
        best_pdf = None; best_range = []
        for pdf_path, mapping in self.index.items():
            ranges = mapping.get(task_code)
            if not ranges: continue
            ranges_sorted = sorted(ranges, key=lambda r:(-(r[1]-r[0]+1), r[0]))
            r = ranges_sorted[0]
            length = r[1]-r[0]+1
            if length > (best_range[1]-best_range[0]+1 if best_range else 0):
                best_range = r; best_pdf = pdf_path
        return best_pdf, best_range

    def find_related_subtasks(self, task_code: str):
        """Return all task codes in index that look like subtasks of given task.

        A subtask is detected when it shares the same first two numeric parts
        (e.g., 27-054-00 is parent; 27-054-01, 27-054-02 ... are subtasks).
        """
        try:
            m = re.match(r"^(\d{2}-\d{2,3})-\d{2}$", task_code)
            if not m:
                return []
            prefix = m.group(1) + "-"
            related = set()
            for pdf_path, mapping in self.index.items():
                for code in mapping.keys():
                    if code.startswith(prefix):
                        related.add(code)
            return sorted(related)
        except Exception:
            return []

# ---------- GUI app ----------
class RedseaApp(ctk.CTk):
    """
    Main REDSEA application for aviation maintenance task management.
    Provides functionality for:
    - Task extraction from PDFs
    - Task stamping and processing 
    - Effectivity and TCM management
    - Check control and utilization tracking
    - CMP/TCM task card generation
    """
    
    def __init__(self):
        super().__init__()
        try:
            # Professional application setup
            self.title("REDSEA Airlines - Aviation Maintenance Toolkit")
            
            # Load visual assets
            self.load_visual_assets()
            
            # Initialize Mini Launcher Manager
            self.mini_launcher = MiniLauncherManager(self)
            
            # Check startup mode
            global MINI_LAUNCHER_MODE, TARGET_MODULE
            
            if MINI_LAUNCHER_MODE:
                # Start in professional mini launcher mode
                self.setup_mini_mode()
            else:
                # Start in full system mode
                self.setup_full_mode()
                
                # Open specific module if requested
                if TARGET_MODULE:
                    # فتح الوحدة فورًا بدون تأخير
                    self.open_target_module(TARGET_MODULE)
                    TARGET_MODULE = None  # Reset after use
            
            # Set professional appearance
            ctk.set_appearance_mode("light")
            ctk.set_default_color_theme("blue")

            
            # Initialize application state variables
            self.selected_pdf = None
            self.output_dir = None
            self.util_df = None
            self.check_df = None
            self.eff_df = None
            self.cmp_pdf = None
            self.tcm_dir = None
            self.out_cmp_tcm = None
            self.indexer = None
            self.covers_dir = None
            self.mpd_rsd_excel = None
            
            # Initialize window control variables
            self.zoom_level = 1.0
            self.is_fullscreen = False
            self.is_maximized = False
            self.original_geometry = f"{APP_MIN_W}x{APP_MIN_H}"
            
            # Flag to track if an operation is running and can be stopped
            self.operation_running = False
            
        except Exception as e:
            messagebox.showerror("Initialization Error", f"Failed to initialize application: {e}")
    
    def load_visual_assets(self):
        """Load visual assets (icon, logo, background)"""
        try:
            # Set application icon
            if os.path.exists(ICON_PATH):
                self.iconbitmap(ICON_PATH)
                print(f"Icon loaded successfully: {ICON_PATH}")
            else:
                print(f"Warning: Icon file not found: {ICON_PATH}")
                
            # Load logo image
            self.logo_image = load_logo_image(LOGO_PATH, size=(80, 80))
            if self.logo_image:
                print(f"Logo loaded successfully: {LOGO_PATH}")
            else:
                print(f"Warning: Logo file not found: {LOGO_PATH}")
                
            # Load background image
            self.background_image = load_background_image(BACKGROUND_PATH, opacity=0.15)
            if self.background_image:
                print(f"Background loaded successfully: {BACKGROUND_PATH}")
            else:
                print(f"Warning: Background file not found: {BACKGROUND_PATH}")
                
        except Exception as e:
            print(f"Warning: Could not load visual assets: {e}")
            self.logo_image = None
            self.background_image = None
    
    def setup_mini_mode(self):
        """Setup professional mini launcher mode"""
        # Configure window for mini mode
        self.geometry("320x480")
        self.resizable(False, False)
        self.center_window()
        
        # Hide window initially
        self.withdraw()
        
        # Show mini launcher
        self.mini_launcher.show_mini_launcher()
        
        # Set professional window properties
        self.after(100, lambda: self.mini_launcher.mini_window.focus_force())
    
    def setup_full_mode(self):
        """Setup professional full system mode"""
        # Configure window for full mode
        self.geometry(f"{APP_MIN_W}x{APP_MIN_H}")
        self.minsize(APP_MIN_W, APP_MIN_H)
        self.resizable(True, True)
        
        # Build professional interface
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self._build_sidebar()
        self._build_main_shell()
        
        # Show default tab
        self.show_tab("Task Extractor")
        
        # Show window with fade-in effect
        self.deiconify()
        self.center_window()
    
    def configure_full_mode(self):
        """Configure system for full mode with advanced window controls"""
        # Set initial window size
        self.geometry(f"{APP_MIN_W}x{APP_MIN_H}")
        self.minsize(APP_MIN_W, APP_MIN_H)
        
        # Enable full resizing capabilities
        self.resizable(True, True)
        
        # Show window first
        self.deiconify()
        
        # Start maximized immediately
        try:
            self.state('zoomed')  # Maximize window automatically
            self.is_maximized = True
        except Exception as e:
            print(f"Could not zoom window: {e}")
        self.is_maximized = True
        
        # Add window control handlers
        self.setup_window_controls()
        
        # Rebuild interface if needed
        if not hasattr(self, 'sidebar'):
            self.setup_full_mode()
    
    def setup_window_controls(self):
        """Setup advanced window control features"""
        # Bind keyboard shortcuts for window control
        self.bind('<F11>', self.toggle_fullscreen)
        self.bind('<Control-plus>', self.zoom_in)
        self.bind('<Control-minus>', self.zoom_out)
        self.bind('<Control-0>', self.reset_zoom)
        self.bind('<Alt-Return>', self.toggle_maximize)
        
        # Store original size for restore functionality
        self.original_geometry = f"{APP_MIN_W}x{APP_MIN_H}"
        self.is_fullscreen = False
        self.is_maximized = False
        self.zoom_level = 1.0
    
    def toggle_fullscreen(self, event=None):
        """Toggle fullscreen mode (F11)"""
        self.is_fullscreen = not self.is_fullscreen
        if self.is_fullscreen:
            self.state('zoomed')
            self.wm_attributes('-fullscreen', True)
        else:
            self.wm_attributes('-fullscreen', False)
            self.state('normal')
        return "break"
    
    def toggle_maximize(self, event=None):
        """Toggle maximize/restore window (Alt+Enter)"""
        if self.state() == 'zoomed':
            self.state('normal')
            self.geometry(self.original_geometry)
            self.is_maximized = False
        else:
            self.state('zoomed')
            self.is_maximized = True
        return "break"
    
    def zoom_in(self, event=None):
        """Zoom in the interface (Ctrl+Plus)"""
        if self.zoom_level < 2.0:
            self.zoom_level += 0.1
            self.apply_zoom()
        return "break"
    
    def zoom_out(self, event=None):
        """Zoom out the interface (Ctrl+Minus)"""
        if self.zoom_level > 0.5:
            self.zoom_level -= 0.1
            self.apply_zoom()
        return "break"
    
    def reset_zoom(self, event=None):
        """Reset zoom to 100% (Ctrl+0)"""
        self.zoom_level = 1.0
        self.apply_zoom()
        return "break"
    
    def apply_zoom(self):
        """Apply current zoom level to the interface"""
        try:
            # Calculate new dimensions based on zoom level
            base_width = APP_MIN_W
            base_height = APP_MIN_H
            new_width = int(base_width * self.zoom_level)
            new_height = int(base_height * self.zoom_level)
            
            # Update window size if not maximized
            if self.state() != 'zoomed':
                self.geometry(f"{new_width}x{new_height}")
            
            # Update font sizes for zoom effect
            self.update_fonts_for_zoom()
            
        except Exception as e:
            print(f"Error applying zoom: {e}")
    
    def update_fonts_for_zoom(self):
        """Update font sizes based on zoom level"""
        try:
            import tkinter.font as tkFont
            
            # Get default font size
            base_font_size = 10
            new_font_size = int(base_font_size * self.zoom_level)
            
            # Update default font
            default_font = tkFont.nametofont("TkDefaultFont")
            default_font.configure(size=new_font_size)
            
            # Force update of all widgets
            self.update()
            
        except Exception as e:
            print(f"Error updating fonts: {e}")
    
    def center_window(self):
        """Center window on screen"""
        self.update_idletasks()
        width = max(self.winfo_width(), APP_MIN_W)
        height = max(self.winfo_height(), APP_MIN_H)
        
        screen_w = self.winfo_screenwidth()
        screen_h = self.winfo_screenheight()
        
        x = max(0, (screen_w // 2) - (width // 2))
        y = max(0, (screen_h // 2) - (height // 2))
        self.geometry(f'{width}x{height}+{x}+{y}')
    
    def show_full_system(self):
        """Display full system maximized automatically"""
        self.configure_full_mode()
        # No need to show window controls info since we removed the controls
    
    def open_target_module(self, module_name):
        """فتح وحدة محددة"""
        module_map = {
            'task_extract': 'Task Extractor',
            'task_stamping': 'Task Stamping',
            'effectivity': 'EFFECTIVITY / TCM',
            'check_control': 'Check Control',
            'utilization': 'Utilization',
            'cmp_tcm': 'CMP / TCM Tasks',
            'mail_merge': 'Mail Merge (Covering)'
        }
        
        target_tab = module_map.get(module_name, 'Task Extractor')
        self.show_tab(target_tab)
    
    def switch_to_mini_mode(self):
        """Professional transition to mini launcher mode"""
        global MINI_LAUNCHER_MODE
        MINI_LAUNCHER_MODE = True
        
        # Smooth transition animation
        self.attributes('-alpha', 0.8)
        self.after(100, lambda: self.attributes('-alpha', 0.6))
        self.after(200, lambda: self.attributes('-alpha', 0.4))
        self.after(300, lambda: self.attributes('-alpha', 0.2))
        self.after(400, self.complete_mini_transition)
    
    def complete_mini_transition(self):
        """Complete transition to mini mode"""
        # Hide current window
        self.withdraw()
        
        # Reset alpha for next time
        self.attributes('-alpha', 1.0)
        
        # Show mini launcher
        self.mini_launcher.show_mini_launcher()

    def _safe_log(self, widget, text):
        try: self.after(0, lambda: widget.insert(tk.END, text))
        except Exception: pass
    def _safe_show_error(self, title, msg):
        try: self.after(0, lambda: messagebox.showerror(title, msg))
        except Exception: pass
    def _safe_show_info(self, title, msg):
        try: self.after(0, lambda: messagebox.showinfo(title, msg))
        except Exception: pass

    def _build_sidebar(self):
        """Build modern professional sidebar with enhanced navigation"""
        self.sidebar = ctk.CTkFrame(self, width=260, corner_radius=0, fg_color=COLOR_CARD_BG)
        self.sidebar.grid(row=0, column=0, sticky="nswe")
        self.sidebar.grid_rowconfigure(10, weight=1)
        
        # Modern header with branding
        header_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=(30, 15))
        
        # Main title with modern styling
        title_label = ctk.CTkLabel(
            header_frame, 
            text="REDSEA\nToolkit", 
            justify="left", 
            font=ctk.CTkFont(size=24, weight="bold"), 
            text_color=COLOR_PRIMARY
        )
        title_label.pack(anchor="w")
        
        # Subtitle with modern styling
        subtitle_label = ctk.CTkLabel(
            header_frame,
            text="Aviation Maintenance",
            font=ctk.CTkFont(size=12, weight="normal"),
            text_color=COLOR_TEXT_DARK
        )
        subtitle_label.pack(anchor="w", pady=(5, 0))
        
        # Modern navigation buttons
        nav_modules = [
            ("Task Extractor", "📋", "Extract and process tasks"),
            ("Task Stamping", "🖨️", "Stamp and enhance documents"),
            ("EFFECTIVITY / TCM", "📈", "Generate reports and analytics"),
            ("Check Control", "✅", "Monitor maintenance checks"),
            ("Utilization", "📉", "Track resource utilization"),
            ("CMP / TCM Tasks", "📦", "Manage task generation"),
            ("Mail Merge (Covering)", "📄", "Automated Mail Merge for RC Cards")
        ]
        
        for name, icon, tooltip in nav_modules:
            self.create_nav_button(name, icon, tooltip)
        
        # Modern separator
        separator = ctk.CTkFrame(self.sidebar, height=1, fg_color=COLOR_BORDER)
        separator.pack(fill="x", padx=20, pady=20)
        
        # Control buttons
        self.create_control_buttons()
        
        # Developer credit (bottom section) with modern styling
        credit_label = ctk.CTkLabel(
            self.sidebar,
            text="by: Eng Fathy Ahmed & Khaled Haggag",
            font=ctk.CTkFont(size=11, weight="bold"),
            text_color=COLOR_PRIMARY
        )
        credit_label.pack(side="bottom", padx=20, pady=(10, 15))
        
        # Modern exit button at bottom
        exit_btn = ctk.CTkButton(
            self.sidebar, 
            text="🚪  Exit Application",
            fg_color=COLOR_DANGER, 
            hover_color="#c53030",
            height=50,
            font=ctk.CTkFont(size=13, weight="bold"),
            text_color="white",
            corner_radius=12,
            command=self.safe_exit
        )
        exit_btn.pack(side="bottom", fill="x", padx=20, pady=(0, 20))
    
    def create_nav_button(self, name, icon, tooltip):
        """Create modern navigation button"""
        btn = ctk.CTkButton(
            self.sidebar,
            text=f"{icon}  {name}",
            height=50,
            fg_color=COLOR_ACCENT,
            hover_color=COLOR_HOVER,
            font=ctk.CTkFont(size=12, weight="bold"),
            anchor="w",
            corner_radius=10,
            command=lambda n=name: self.show_tab(n)
        )
        btn.pack(fill="x", padx=20, pady=6)
        return btn
    
    def create_control_buttons(self):
        """Create modern control buttons section"""
        # Switch to Mini Mode button
        mini_btn = ctk.CTkButton(
            self.sidebar,
            text="🔄  Switch to Quick Launcher",
            height=45,
            fg_color=COLOR_OK,
            hover_color="#2f855a",
            font=ctk.CTkFont(size=11, weight="bold"),
            corner_radius=10,
            command=self.switch_to_mini_mode
        )
        mini_btn.pack(fill="x", padx=20, pady=6)
        
        # Settings/About button
        about_btn = ctk.CTkButton(
            self.sidebar,
            text="ℹ️  About System",
            height=40,
            fg_color="transparent",
            border_width=2,
            border_color=COLOR_PRIMARY,
            text_color=COLOR_PRIMARY,
            hover_color=COLOR_BG_LIGHT,
            font=ctk.CTkFont(size=10, weight="bold"),
            corner_radius=10,
            command=self.show_about
        )
        about_btn.pack(fill="x", padx=20, pady=6)
    
    def add_tooltip(self, widget, text):
        """Add professional tooltip to widget"""
        def on_enter(event):
            tooltip = ctk.CTkToplevel()
            tooltip.wm_overrideredirect(True)
            tooltip.configure(fg_color="black")
            
            label = ctk.CTkLabel(
                tooltip,
                text=text,
                text_color="white",
                fg_color="black",
                font=ctk.CTkFont(size=9)
            )
            label.pack(padx=8, pady=4)
            
            # Position tooltip
            x = widget.winfo_rootx() + widget.winfo_width() + 10
            y = widget.winfo_rooty() + widget.winfo_height() // 2
            tooltip.geometry(f"+{x}+{y}")
            
            # Auto-hide
            tooltip.after(3000, tooltip.destroy)
        
        widget.bind("<Enter>", on_enter, add="+")
    
    def show_about(self):
        """Show modern about dialog"""
        about_text = """REDSEA Airlines Aviation Maintenance Toolkit
        
Version: 3.0 Modern Professional
Build: 2025.01.15

✨ Modern Features:
• Enhanced Task Extraction & Processing
• Advanced Document Stamping & OCR
• Smart Effectivity & TCM Management
• Intelligent Maintenance Check Control
• Real-time Resource Utilization Tracking
• Automated CMP/TCM Task Generation
• Related Tasks Auto-Discovery

🎨 Modern UI Design:
• Clean & Professional Interface
• Enhanced User Experience
• Responsive Layout Design
• Modern Color Scheme

© 2025 REDSEA Airlines
Professional Aviation Maintenance Solutions"""
        
        messagebox.showinfo("About REDSEA Toolkit", about_text)
    
    def safe_exit(self):
        """Safe application exit with confirmation"""
        result = messagebox.askyesno(
            "Exit Application",
            "Are you sure you want to exit REDSEA Toolkit?\n\nAny unsaved work will be lost.",
            icon="question"
        )
        if result:
            self.quit()
            self.destroy()

    def _build_main_shell(self):
        self.main = ctk.CTkFrame(self, corner_radius=0); self.main.grid(row=0, column=1, sticky="nsew")
        self.main.grid_rowconfigure(1, weight=1); self.main.grid_columnconfigure(0, weight=1)

        # Add airplane background using a Canvas so it stays visible behind CTk widgets
        try:
            from PIL import Image, ImageTk
            if os.path.exists(BACKGROUND_PATH):
                # Load original PIL image once
                self._bg_pil = Image.open(BACKGROUND_PATH).convert('RGBA')
                # Create canvas
                self.bg_canvas = tk.Canvas(self.main, highlightthickness=0, bd=0)
                self.bg_canvas.place(x=0, y=0, relwidth=1, relheight=1)

                def _redraw_bg(event=None):
                    try:
                        w = max(1, self.main.winfo_width()); h = max(1, self.main.winfo_height())
                        # Resize while preserving aspect ratio to cover
                        img = self._bg_pil.copy()
                        img_ratio = img.width / img.height
                        frame_ratio = w / h if h else img_ratio
                        if frame_ratio > img_ratio:
                            # fit width
                            new_w = w; new_h = int(w / img_ratio)
                        else:
                            new_h = h; new_w = int(h * img_ratio)
                        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                        # Center crop to frame size
                        x0 = (new_w - w) // 2; y0 = (new_h - h) // 2
                        img = img.crop((x0, y0, x0 + w, y0 + h))
                        # Convert to PhotoImage and draw
                        self._bg_photo = ImageTk.PhotoImage(img)
                        self.bg_canvas.delete('all')
                        self.bg_canvas.create_image(0, 0, anchor='nw', image=self._bg_photo)
                        self.bg_canvas.lower()
                    except Exception:
                        pass

                # Bind resize to redraw background
                self.main.bind('<Configure>', _redraw_bg)
                # Draw initial background after layout
                self.after(50, _redraw_bg)
        except Exception as e:
            print(f"Warning: Canvas background setup failed: {e}")
        
        self.header = ctk.CTkFrame(self.main, height=80, fg_color="transparent", corner_radius=0); self.header.grid(row=0,column=0,sticky="ew",padx=0,pady=(0,12))
        self.header.grid_columnconfigure(0, weight=1)
        
        # Modern header with logo and title
        header_content = ctk.CTkFrame(self.header, fg_color="transparent")
        header_content.grid(row=0, column=0, sticky="ew", padx=30, pady=20)
        header_content.grid_columnconfigure(1, weight=1)
        
        # Add logo to header if available
        if hasattr(self, 'logo_image') and self.logo_image:
            logo_label = ctk.CTkLabel(header_content, image=self.logo_image, text="")
            logo_label.grid(row=0, column=0, padx=(0, 20), sticky="w")
        else:
            # Add REDSEA Airlines icon if logo not available
            try:
                icon_path = os.path.join(ASSETS_PATH, "REDSEA Airlines Logo.ico")
                if os.path.exists(icon_path):
                    # Load icon and convert to PhotoImage
                    icon_image = tk.PhotoImage(file=icon_path)
                    icon_label = ctk.CTkLabel(header_content, image=icon_image, text="")
                    icon_label.grid(row=0, column=0, padx=(0, 20), sticky="w")
                    # Keep reference to prevent garbage collection
                    self.icon_image = icon_image
            except Exception as e:
                print(f"Could not load REDSEA Airlines icon: {e}")
        
        title_label = ctk.CTkLabel(header_content, text="Planning & Engineering", font=("Segoe UI",28,"bold"), text_color=COLOR_PRIMARY)
        title_label.grid(row=0, column=1, sticky="w")
        
        self.breadcrumb = ctk.CTkLabel(self.header, text="", font=("Segoe UI",14,"normal"), text_color=COLOR_TEXT_DARK); self.breadcrumb.grid(row=1,column=0,padx=30,pady=(0,15),sticky="w")
        self.content = ctk.CTkFrame(self.main, fg_color="transparent", corner_radius=0); self.content.grid(row=1,column=0,sticky="nsew",padx=20,pady=(0,20))
        self.content.grid_columnconfigure(0, weight=1); self.content.grid_rowconfigure(1, weight=1)
        
        # No extra content background overlay; Canvas already draws full background

    def clear_content(self):
        for w in self.content.winfo_children(): w.destroy()

    def show_tab(self, tab_name):
        self.clear_content(); self.breadcrumb.configure(text=tab_name)
        if tab_name=="Task Extractor": self._tab_task_extractor()
        elif tab_name=="Task Stamping": self._tab_task_stamping()
        elif tab_name=="EFFECTIVITY / TCM": self._tab_effectivity()
        elif tab_name=="Check Control": self._tab_check_control()
        elif tab_name=="Utilization": self._tab_utilization()
        elif tab_name=="CMP / TCM Tasks": self._tab_cmp_tcm_tasks()
        elif tab_name=="Mail Merge (Covering)": self._tab_mail_merge()

    # ---------- Task Extractor ----------
    def _tab_task_extractor(self):
        card = ctk.CTkFrame(self.content, corner_radius=16, fg_color=COLOR_CARD_BG); card.grid(row=0,column=0,sticky="ew",padx=12,pady=12); card.grid_columnconfigure((0,1,2,3),weight=1)
        ctk.CTkLabel(card, text="Task Extractor", font=("Segoe UI",22,"bold"), text_color=COLOR_PRIMARY).grid(row=0,column=0,padx=20,pady=(20,12),sticky="w")
        
        self.lbl_pdf_path = ctk.CTkLabel(card, text="No folder selected", font=("Segoe UI",11), text_color=COLOR_TEXT_DARK)
        ctk.CTkButton(card, text="📁 Source Folder…", fg_color=COLOR_OK, hover_color="#2f855a", height=40, font=("Segoe UI",12,"bold"), corner_radius=10, command=self._pick_pdf_folder).grid(row=1,column=0,padx=20,pady=10,sticky="w")
        self.lbl_pdf_path.grid(row=1,column=1,columnspan=3,padx=12,pady=10,sticky="w")
        
        self.output_dir=None; self.lbl_out_dir=ctk.CTkLabel(card,text="Output: not set", font=("Segoe UI",11), text_color=COLOR_TEXT_DARK)
        ctk.CTkButton(card,text="📂 Output Folder…",fg_color=COLOR_ACCENT,hover_color=COLOR_HOVER,height=40,font=("Segoe UI",12,"bold"),corner_radius=10,command=self._pick_output_folder).grid(row=2,column=0,padx=20,pady=10,sticky="w")
        self.lbl_out_dir.grid(row=2,column=1,columnspan=3,padx=12,pady=10,sticky="w")
        
        opt=ctk.CTkFrame(card, fg_color=COLOR_BG_LIGHT, corner_radius=12); opt.grid(row=3,column=0,columnspan=4,padx=20,pady=(12,20),sticky="ew"); opt.grid_columnconfigure((0,1,2,3),weight=1)
        ctk.CTkLabel(opt,text="Task:", font=("Segoe UI",12,"bold"), text_color=COLOR_PRIMARY).grid(row=0,column=0,padx=12,pady=12,sticky="w")
        self.task_entry = ctk.CTkEntry(opt, placeholder_text="e.g., 27-054-00", height=35, font=("Segoe UI",11), corner_radius=8); self.task_entry.grid(row=0,column=1,padx=12,pady=12,sticky="ew")
        self.skip_first_var = tk.BooleanVar(value=False); self.combine_var = tk.BooleanVar(value=True)
       # ctk.CTkCheckBox(opt, text="Skip first found page", variable=self.skip_first_var, font=("Segoe UI",10), text_color=COLOR_TEXT_DARK).grid(row=0,column=2,padx=12,pady=12,sticky="w")
       # ctk.CTkCheckBox(opt, text="Combine all pages in one PDF", variable=self.combine_var, font=("Segoe UI",10), text_color=COLOR_TEXT_DARK).grid(row=0,column=3,padx=12,pady=12,sticky="w")
        ctk.CTkButton(card,text="🚀 Run Extract",height=50,fg_color=COLOR_WARN,hover_color="#dd6b20",font=("Segoe UI",14,"bold"),corner_radius=12,command=self._run_extract).grid(row=4,column=0,padx=20,pady=(0,20),sticky="w")
        
        table = ctk.CTkFrame(self.content, corner_radius=16, fg_color=COLOR_CARD_BG); table.grid(row=1,column=0,sticky="nsew",padx=12,pady=12); table.grid_rowconfigure(0,weight=1); table.grid_columnconfigure(0,weight=1)
        
        # Copy button frame
        btn_frame = ctk.CTkFrame(table, fg_color="transparent")
        btn_frame.grid(row=0, column=0, sticky="ew", padx=15, pady=8)
        btn_frame.grid_columnconfigure(0, weight=1)
        
        ctk.CTkButton(btn_frame, text="📋 Copy Log", command=self._copy_task_log, height=35, font=("Segoe UI",11,"bold"), corner_radius=8).grid(row=0, column=0, padx=(0, 12), sticky="w")
        
        self.tx_task_log = tk.Text(table, height=12, font=("Consolas",10), bg=COLOR_BG_LIGHT, fg=COLOR_TEXT_DARK, relief="flat", bd=0); self.tx_task_log.grid(row=1,column=0,sticky="nsew", padx=15, pady=15); ttk.Scrollbar(table,command=self.tx_task_log.yview).grid(row=1,column=1,sticky="ns")
        self.tx_task_log.configure(yscrollcommand=lambda *args: None)

    def _copy_task_log(self):
        """Copy task log content to clipboard"""
        try:
            log_content = self.tx_task_log.get("1.0", tk.END)
            self.clipboard_clear()
            self.clipboard_append(log_content)
            self._safe_log(self.tx_task_log, "📋 Log copied to clipboard!\n")
        except Exception as e:
            self._safe_log(self.tx_task_log, f"❌ Error copying log: {e}\n")

    def _pick_pdf_folder(self):
        """Select a folder containing PDF files for task extraction"""
        try:
            folder = filedialog.askdirectory(
                title="Select Folder Containing PDF Files"
            )
            if folder and os.path.exists(folder):
                self.selected_pdf_folder = folder
                self.lbl_pdf_path.configure(text=folder)
                self._safe_log(self.tx_task_log, f"Selected folder: {os.path.basename(folder)}\n")
                
                # Count PDF files in the folder
                pdf_count = len([f for f in os.listdir(folder) if f.lower().endswith('.pdf')])
                self._safe_log(self.tx_task_log, f"Found {pdf_count} PDF files in folder\n")
            elif folder:
                self._safe_show_error("Folder Error", "Selected folder does not exist or is not accessible.")
        except Exception as e:
            self._safe_show_error("Selection Error", f"Failed to select folder: {e}")

    def _pick_pdf(self):
        """Select a PDF file for task extraction (legacy function - kept for compatibility)"""
        try:
            f = filedialog.askopenfilename(
                title="Select PDF File",
                filetypes=[("PDF Files", "*.pdf")]
            )
            if f and os.path.exists(f):
                self.selected_pdf = f
                self.lbl_pdf_path.configure(text=f)
                self._safe_log(self.tx_task_log, f"Selected PDF: {os.path.basename(f)}\n")
            elif f:
                self._safe_show_error("File Error", "Selected file does not exist or is not accessible.")
        except Exception as e:
            self._safe_show_error("Selection Error", f"Failed to select PDF file: {e}")
            
    def _pick_output_folder(self):
        """Select output folder for extracted tasks"""
        try:
            d = filedialog.askdirectory(title="Select Output Folder")
            if d and os.path.exists(d):
                self.output_dir = d
                self.lbl_out_dir.configure(text=d)
                self._safe_log(self.tx_task_log, f"Output folder: {d}\n")
            elif d:
                self._safe_show_error("Folder Error", "Selected folder does not exist or is not accessible.")
        except Exception as e:
            self._safe_show_error("Selection Error", f"Failed to select output folder: {e}")

    def _run_extract(self):
        def job():
            # Check if folder is selected
            if not hasattr(self, 'selected_pdf_folder') or not self.selected_pdf_folder or not os.path.exists(self.selected_pdf_folder):
                self._safe_show_error("Error","Please choose a valid source folder first."); return
            
            code = self.task_entry.get().strip()
            if not code:
                self._safe_show_error("Error","Please enter a task number."); return
            
            try:
                # Extract first two digits for PDF filename
                first_two_digits = code.split('-')[0] if '-' in code else code[:2]
                
                self._safe_log(self.tx_task_log, f"🔍 Looking for PDF files starting with: {first_two_digits}\n")
                
                # Search for PDF files that start with the first two digits
                pdf_files = [f for f in os.listdir(self.selected_pdf_folder) if f.lower().endswith('.pdf')]
                matching_files = [f for f in pdf_files if f.startswith(first_two_digits)]
                
                if not matching_files:
                    self._safe_log(self.tx_task_log, f"❌ No PDF files found starting with '{first_two_digits}'\n")
                    self._safe_log(self.tx_task_log, f"📁 Available PDF files in folder:\n")
                    for pdf_file in sorted(pdf_files):
                        self._safe_log(self.tx_task_log, f"  - {pdf_file}\n")
                    return
                
                # Use the first matching file
                pdf_filename = matching_files[0]
                pdf_path = os.path.join(self.selected_pdf_folder, pdf_filename)
                
                self._safe_log(self.tx_task_log, f"✅ Found matching PDF file: {pdf_filename}\n")
                
                # Search for related tasks in the specific PDF file
                self._safe_log(self.tx_task_log, f"🔍 Searching for related tasks to '{code}' in {pdf_filename}...\n")
                related_tasks = self._find_related_tasks(code, pdf_path=pdf_path)
                
                if related_tasks:
                    self._safe_log(self.tx_task_log, f"✅ Found {len(related_tasks)} related tasks: {', '.join(related_tasks)}\n")
                    
                    # Extract related tasks to PDF files
                    out_dir = self.output_dir or self.selected_pdf_folder
                    saved_count = self._extract_related_tasks_to_pdf(code, related_tasks, pdf_path, out_dir)
                    
                    if saved_count > 0:
                        self._safe_log(self.tx_task_log, f"✅ Successfully extracted {saved_count} related tasks to PDF files\n")
                        self._safe_log(self.tx_task_log, f"📁 Output directory: {out_dir}\n")
                        self._safe_show_info("Related Tasks Found", f"Found and extracted {saved_count} related tasks!\nLocation: {out_dir}")
                        return
                    else:
                        self._safe_log(self.tx_task_log, f"⚠️ No related tasks could be extracted\n")
                else:
                    self._safe_log(self.tx_task_log, f"ℹ️ No related tasks found for '{code}' in {pdf_filename}\n")
                
                # If no related tasks found, search for the base task in the PDF file
                self._safe_log(self.tx_task_log, f"🔄 Searching for base task '{code}' in {pdf_filename}...\n")
                
                doc = fitz.open(pdf_path)
                matches = self._scan_pages_for_code_parallel(doc, code, max_workers=TASK_EXTRACTOR_THREADS)
                
                if self.skip_first_var.get() and matches:
                    matches = matches[1:]
                
                if not matches:
                    self._safe_log(self.tx_task_log, f"❌ No matching pages found for '{code}' in {pdf_filename}\n")
                    doc.close()
                    return
                
                self._safe_log(self.tx_task_log, f"✅ Found {len(matches)} pages with '{code}' in {pdf_filename}\n")
                
                # Extract pages
                out_dir = self.output_dir or self.selected_pdf_folder
                safe_make_dir(out_dir)
                
                if self.combine_var.get():
                    out = fitz.open()
                    runs = group_contiguous(matches)
                    for r in runs:
                        out.insert_pdf(doc, from_page=r[0], to_page=r[-1])
                    out_name = f"{code.replace('/','_')}_extracted.pdf"
                    out_path = os.path.join(out_dir, out_name)
                    out.save(out_path)
                    self._safe_log(self.tx_task_log, f"✅ Saved combined: {out_name}\n")
                else:
                    for idx, pn in enumerate(matches, 1):
                        out = fitz.open()
                        out.insert_pdf(doc, from_page=pn, to_page=pn)
                        out_name = f"{code.replace('/','_')}_p{pn+1}_{idx}.pdf"
                        out_path = os.path.join(out_dir, out_name)
                        out.save(out_path)
                        self._safe_log(self.tx_task_log, f"✅ Saved page: {out_name}\n")
                
                doc.close()
                    
            except Exception as e:
                self._safe_show_error("Extract failed", str(e))
        threading.Thread(target=job, daemon=True).start()

    def _scan_pages_for_code_parallel(self, doc, code, max_workers=3):
        """Return list of page indices containing the code, excluding INDEX pages, using threads."""
        try:
            total = doc.page_count
            indices = list(range(total))
            results = []

            def check_page(pn: int):
                try:
                    page = doc[pn]
                    text = page.get_text()
                    if code in text and not self._is_index_page(text):
                       
                       
                       
                        return pn
                except Exception:
                    return None
                return None

            with ThreadPoolExecutor(max_workers=max_workers) as ex:
                futures = [ex.submit(check_page, pn) for pn in indices]
                for fut in as_completed(futures):
                    pn = fut.result()
                    if isinstance(pn, int):
                        results.append(pn)
            results.sort()
            return results
        except Exception:
            # Fallback to sequential scan on error
            matches = []
            for i, page in enumerate(doc):
                try:
                    text = page.get_text()
                    if code in text and not self._is_index_page(text):
                        matches.append(i)
                except Exception:
                    continue
            return matches
        
    # ─── Missing methods needed by _run_extract (L1615) ─────────────────────
    # These were called but never defined in the original app2.py.
    # Built using the same subtask regex from TcmIndexer.find_related_subtasks (L904).

    def _is_index_page(self, text):
        """Detect INDEX / table-of-contents pages that should be skipped.
        
        Heuristic: a page is an index page if it contains an 'INDEX' header
        AND has a high density of task codes (>10 codes on a single page
        usually means it's a listing, not the actual task content).
        """
        if not text:
            return False
        upper = text.upper()
        # Common index page markers
        if any(marker in upper for marker in ("INDEX", "TABLE OF CONTENTS", "LIST OF EFFECTIVE PAGES")):
            # Only flag as index if there's a high code density
            codes_on_page = TASK_PATTERN.findall(text)
            if len(codes_on_page) > 10:
                return True
        return False

    def _find_related_tasks(self, code, pdf_path=None):
        """Find related subtask codes within a PDF that share the same
        first two numeric parts as `code`.
        
        Uses the same prefix-matching logic as TcmIndexer.find_related_subtasks:
        e.g. 27-054-00 is the parent; 27-054-01, 27-054-02, ... are subtasks.
        
        Returns sorted list of related codes (excluding the original code).
        """
        m = re.match(r"^(\d{2}-\d{2,3})-\d{2}$", code.strip())
        if not m or not pdf_path:
            return []
        prefix = m.group(1) + "-"
        try:
            doc = fitz.open(pdf_path)
        except Exception:
            return []
        try:
            found = set()
            for page in doc:
                text = page.get_text("text") or ""
                if not text.strip():
                    text = self._stamp_extract_text_with_ocr(page)
                for m2 in TASK_PATTERN.finditer(text):
                    cand = m2.group(0)
                    if cand.startswith(prefix) and cand != code:
                        found.add(cand)
            return sorted(found)
        finally:
            doc.close()

    def _extract_related_tasks_to_pdf(self, code, related_tasks, pdf_path, out_dir):
        """Extract pages for the original task + all related subtasks into
        separate PDF files under out_dir.
        
        Returns the number of files successfully saved.
        """
        safe_make_dir(out_dir)
        all_codes = [code] + list(related_tasks)
        try:
            doc = fitz.open(pdf_path)
        except Exception:
            return 0
        saved = 0
        try:
            for c in all_codes:
                matches = self._scan_pages_for_code_parallel(doc, c, max_workers=TASK_EXTRACTOR_THREADS)
                if not matches:
                    continue
                out = fitz.open()
                for run in group_contiguous(matches):
                    out.insert_pdf(doc, from_page=run[0], to_page=run[-1])
                out_name = f"{c.replace('/', '_')}_related.pdf"
                out.save(os.path.join(out_dir, out_name))
                out.close()
                saved += 1
            return saved
        finally:
            doc.close()

    # ***** START OF MODIFICATION: Task Stamping Tab and Logic *****

    # --- Data and Helpers for Stamping ---
    TAIL_MAP = {
        "BTR": "SU-RSA",
        "ILF": "SU-RSB",
        "GUN": "SU-RSC",
        "GOT": "SU-RSD"
    }
    
    # Package Relations
    PACKAGE_RELATIONS = {
        "A10": ["A2", "A4", "A8", "A10"]
    }
    
    def _create_package_structure(self, package):
        """Create folder structure for package relations
        
        Args:
            package: The main package to create structure for (e.g., 'A10')
        """
        try:
            if package not in self.PACKAGE_RELATIONS:
                self._safe_show_error("Error", f"Package {package} not found in relations")
                return
            
            # Get the folder path from user
            base_folder = filedialog.askdirectory(title=f"Select folder to create {package} structure")
            if not base_folder:
                return
                
            # Create main package folder
            main_folder = os.path.join(base_folder, package)
            os.makedirs(main_folder, exist_ok=True)
            
            # Create subfolders for related packages
            created = []
            for related_pkg in self.PACKAGE_RELATIONS[package]:
                pkg_folder = os.path.join(main_folder, related_pkg)
                os.makedirs(pkg_folder, exist_ok=True)
                created.append(related_pkg)
            
            self._safe_show_info("Success", f"Created folders for packages: {', '.join(created)}")
            
        except Exception as e:
            self._safe_show_error("Error", f"Failed to create package structure: {str(e)}")

    def _process_package_relations(self, package):
        """Process package relations and create folder structure
        
        Args:
            package: The main package to process (e.g., 'A10')
        """
        if package not in self.PACKAGE_RELATIONS:
            self._safe_show_error("Error", f"Package {package} not found in relations")
            return
            
        # Create main package folder
        main_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), package)
        os.makedirs(main_folder, exist_ok=True)
        
        # Create subfolders for related packages
        for related_pkg in self.PACKAGE_RELATIONS[package]:
            pkg_folder = os.path.join(main_folder, related_pkg)
            os.makedirs(pkg_folder, exist_ok=True)
            self._safe_log(self.tx_stamp_log, f"Created folder for package: {related_pkg}\n")
        
        self._safe_show_info("Complete", f"Created folder structure for {package} package relations")
    
    def _stamp_extract_text_with_ocr(self, page):
        """Use OCR to extract text from page when no direct text is available"""
        try:
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img)
            return text
        except Exception as e:
            self._safe_log(self.tx_stamp_log, f"OCR Error: {e}\n")
            return ""

    def _stamp_page_data(self, doc, tail_number, station_text, date_text):
        """
        : Stamp data on the page
        - TAIL NUMBER
        - AIRLINE CARD NO (RC + first 3 parts of BOEING CARD NO)
        - STATION (from user input)
        - DATE (from user input)
        """
        for page_num, page in enumerate(doc):
            try:
                airline_card_no = "RC"
                found_number = False

                self._safe_log(self.tx_stamp_log, f"   - Page {page_num + 1}: Processing...\n")
                
                # --- Method 1: Location-based text search ---
                search_terms = ["BOEING CARD NO.", "BOEING CARD NO"]
                search_instances = []
                for term in search_terms:
                    search_instances = page.search_for(term)
                    if search_instances:
                        self._safe_log(self.tx_stamp_log, f"     > Method 1: Found label '{term}'.\n")
                        break
                
                if search_instances:
                    label_rect = search_instances[0]
                    search_area = fitz.Rect(label_rect.x1 - 5, label_rect.y0 - 5, label_rect.x1 + 300, label_rect.y1 + 5)
                    extracted_text = page.get_text("text", clip=search_area).strip()
                    
                    if extracted_text:
                        self._safe_log(self.tx_stamp_log, f"     > Text found in area: '{extracted_text.replace(chr(10), ' ')}'\n")
                        boeing_match = re.search(r'([\d-]+)', extracted_text)
                        if boeing_match:
                            full_boeing_no = boeing_match.group(1).strip()
                            self._safe_log(self.tx_stamp_log, f"     > Extracted number: {full_boeing_no}\n")
                            parts = full_boeing_no.split('-')
                            if len(parts) >= 3:
                                airline_card_no = f"RC{parts[0]}-{parts[1]}-{parts[2]}"
                                found_number = True
                            elif len(parts) >= 2:
                                airline_card_no = f"RC{parts[0]}-{parts[1]}"
                                found_number = True
                        else:
                            self._safe_log(self.tx_stamp_log, f"     > WARNING: No number pattern found in the area.\n")
                    else:
                        self._safe_log(self.tx_stamp_log, f"     > WARNING: Found label, but the area next to it is empty.\n")

                # --- Method 2: Fallback to full page text search ---
                if not found_number:
                    self._safe_log(self.tx_stamp_log, f"     > Method 1 failed. Trying Method 2 (Full page search).\n")
                    page_text = page.get_text()
                    boeing_match = re.search(r"BOEING\s+CARD\s+NO\.?\s*([\d-]+)", page_text, re.IGNORECASE)
                    if boeing_match:
                        full_boeing_no = boeing_match.group(1).strip()
                        self._safe_log(self.tx_stamp_log, f"     > Method 2 found number: {full_boeing_no}\n")
                        parts = full_boeing_no.split('-')
                        if len(parts) >= 3:
                            airline_card_no = f"RC{parts[0]}-{parts[1]}-{parts[2]}"
                        elif len(parts) >= 2:
                            airline_card_no = f"RC{parts[0]}-{parts[1]}"
                    else:
                        self._safe_log(self.tx_stamp_log, f"     > Method 2 failed. Number will be 'RC'.\n")
                
                # --- Stamping ---
                items_to_stamp = [
                    ("TAIL NUMBER", tail_number), 
                    ("AIRLINE CARD NO", airline_card_no), 
                    ("STATION", station_text),
                    ("DATE", date_text) # Added DATE here
                ]
                for label, value in items_to_stamp:
                    if not value: continue
                    instances = page.search_for(label)
                    for inst in instances:
                        page.insert_text((inst.x0, inst.y1 + 10), value, fontsize=10, color=(0, 0, 0))

            except Exception as e:
                self._safe_log(self.tx_stamp_log, f"   - Stamping Error on page {page_num + 1}: {e}\n")

    def _stamp_process_single_pdf(self, input_file, output_folder, station_value, date_value):
        try:
            self._safe_log(self.tx_stamp_log, f"\nProcessing file: {os.path.basename(input_file)}\n")
            temp_file = os.path.join(output_folder, f"temp_{os.path.basename(input_file)}")
            
            # Open and process the document
            doc = fitz.open(input_file)
            tasks = {}
            plane_code_from_cover = None

            # Detect plane code from the first few pages
            for i, page in enumerate(doc):
                if i > 4: break
                text = page.get_text() or self._stamp_extract_text_with_ocr(page)
                for code in self.TAIL_MAP.keys():
                    if code in text:
                        plane_code_from_cover = code
                        self._safe_log(self.tx_stamp_log, f" > Detected Plane Code from Cover: {plane_code_from_cover} -> {self.TAIL_MAP[plane_code_from_cover]}\n")
                        break
                if plane_code_from_cover: break

            # Find all tasks and their pages in the document
            for i, page in enumerate(doc):
                text = page.get_text() or self._stamp_extract_text_with_ocr(page)
                match = re.search(r"(\d{2,3}-\d{3}-\d{2}-\d{2})", text)
                if match:
                    task_no = match.group(1)
                    if task_no not in tasks: tasks[task_no] = {"pages": []}
                    tasks[task_no]["pages"].append(i)

            if not tasks:
                self._safe_log(self.tx_stamp_log, " > No tasks found in this document.\n")
                doc.close()
                return

            self._safe_log(self.tx_stamp_log, f" > Found {len(tasks)} unique tasks. Saving and stamping...\n")
            
            # Create a separate, stamped PDF for each task
            for task, data in tasks.items():
                new_doc = fitz.open()
                for pno in data["pages"]:
                    new_doc.insert_pdf(doc, from_page=pno, to_page=pno)

                if plane_code_from_cover:
                    tail_number = self.TAIL_MAP.get(plane_code_from_cover, "")
                    self._stamp_page_data(new_doc, tail_number, station_value, date_value)
                else:
                    self._safe_log(self.tx_stamp_log, f" > Warning: Could not determine plane for task {task}. Stamping skipped.\n")

                # Save to temporary file first
                new_doc.save(temp_file)
                new_doc.close()
                
                # Verify the temporary file is valid
                with fitz.open(temp_file) as _:
                    pass
                    
                # Replace the original file
                os.replace(temp_file, input_file)
                self._safe_log(self.tx_stamp_log, f"   - Updated: {os.path.basename(input_file)}\n")

            doc.close()
        except Exception as e:
            self._safe_log(self.tx_stamp_log, f"An error occurred while processing {os.path.basename(input_file)}: {e}\n")

    def _stamp_process_folder(self, input_folder, output_folder, station_value, date_value):
        try:
            # Create a temporary folder for processing
            temp_folder = os.path.join(output_folder, "_temp_processing")
            os.makedirs(temp_folder, exist_ok=True)
            
            # Ensure we have write permissions in the input folder
            test_file = os.path.join(input_folder, "test_write_permission")
            try:
                with open(test_file, 'w') as f:
                    f.write('')
                os.remove(test_file)
            except:
                self._safe_show_error("Permission Error", "Cannot write to input folder. Please check folder permissions.")
                return
            
            # Get all PDF files in the input folder
            pdf_files = list(walk_pdfs_in_dir(input_folder))
            if not pdf_files:
                self._safe_log(self.tx_stamp_log, "No PDF files found in the selected folder.\n")
                self._safe_show_info("Done", "No PDF files were found to process.")
                return

            self._safe_log(self.tx_stamp_log, f"Found {len(pdf_files)} PDF(s) to process.\n")
            processed = 0
            
            for pdf_file in pdf_files:
                try:
                    # Generate temp file path
                    temp_path = os.path.join(temp_folder, f"temp_{os.path.basename(pdf_file)}")
                    
                    # Process the file to temporary location
                    self._stamp_process_single_pdf(pdf_file, temp_folder, station_value, date_value)
                    
                    # If successful, replace the original file
                    if os.path.exists(temp_path):
                        try:
                            # Verify the temp file is valid
                            with fitz.open(temp_path) as _:
                                pass
                            os.replace(temp_path, pdf_file)
                            processed += 1
                            self._safe_log(self.tx_stamp_log, f"Updated file {processed}/{len(pdf_files)}: {os.path.basename(pdf_file)}\n")
                        except Exception as e:
                            self._safe_log(self.tx_stamp_log, f"Failed to replace {os.path.basename(pdf_file)}: {str(e)}\n")
                            continue
                except Exception as e:
                    self._safe_log(self.tx_stamp_log, f"Error processing {os.path.basename(pdf_file)}: {str(e)}\n")
                    continue
            
            # Clean up temporary folder
            try:
                os.rmdir(temp_folder)
            except:
                pass
            
            if processed > 0:
                self._safe_log(self.tx_stamp_log, f"\nSuccessfully updated {processed} file(s).\n")
                self._safe_show_info("Complete", f"Successfully updated {processed} files.")
            else:
                self._safe_log(self.tx_stamp_log, "\nNo files were updated successfully.\n")
                self._safe_show_error("Warning", "No files were updated successfully.")
                
        except Exception as e:
            self._safe_log(self.tx_stamp_log, f"Error processing folder: {str(e)}\n")
            self._safe_show_error("Processing Failed", str(e))
            self._safe_log(self.tx_stamp_log, f"A critical error occurred: {e}\n")

    def _tab_task_stamping(self):
        card = ctk.CTkFrame(self.content, corner_radius=12)
        card.grid(row=0, column=0, sticky="new", padx=8, pady=8)
        card.grid_columnconfigure(1, weight=1)
        
        # Package Relations Frame
        pkg_frame = ctk.CTkFrame(self.content, corner_radius=12)
        pkg_frame.grid(row=1, column=0, sticky="new", padx=8, pady=8)
        pkg_frame.grid_columnconfigure(1, weight=1)
        
        ctk.CTkLabel(pkg_frame, text="Package Relations", font=("Segoe UI", 18, "bold"), text_color=COLOR_PRIMARY).grid(row=0, column=0, columnspan=2, padx=16, pady=(16, 8), sticky="w")
        
        # Add Create A10 Structure button
        ctk.CTkButton(pkg_frame, text="Create A10 Package Structure", fg_color=COLOR_OK, 
                      command=lambda: self._create_package_structure("A10")).grid(row=1, column=0, columnspan=2, padx=16, pady=8, sticky="w")

        ctk.CTkLabel(card, text="Task Splitting & Stamping", font=("Segoe UI", 18, "bold"), text_color=COLOR_PRIMARY).grid(row=0, column=0, columnspan=3, padx=16, pady=(16, 8), sticky="w")
        
        ctk.CTkButton(card, text="Select Input Folder...", fg_color=COLOR_OK, command=self._stamp_browse_input_folder).grid(row=1, column=0, padx=16, pady=8, sticky="w")
        self.entry_stamp_input_folder = ctk.CTkEntry(card, placeholder_text="Path to the folder with PDF files")
        self.entry_stamp_input_folder.grid(row=1, column=1, columnspan=2, padx=8, pady=8, sticky="ew")

        ctk.CTkButton(card, text="Select Output Folder...", fg_color=COLOR_ACCENT, command=self._stamp_browse_output_folder).grid(row=2, column=0, padx=16, pady=8, sticky="w")
        self.entry_stamp_output_folder = ctk.CTkEntry(card, placeholder_text="Path to save stamped tasks")
        self.entry_stamp_output_folder.grid(row=2, column=1, columnspan=2, padx=8, pady=8, sticky="ew")
        
        ctk.CTkLabel(card, text="Station (Optional):").grid(row=3, column=0, padx=16, pady=8, sticky="w")
        self.entry_stamp_station = ctk.CTkEntry(card, placeholder_text="Enter station name to be stamped, e.g., CAI")
        self.entry_stamp_station.grid(row=3, column=1, columnspan=2, padx=8, pady=8, sticky="ew")

        ctk.CTkLabel(card, text="Date (Optional):").grid(row=4, column=0, padx=16, pady=8, sticky="w")
        self.entry_stamp_date = ctk.CTkEntry(card, placeholder_text="Enter date to be stamped, e.g., 24-SEP-2025")
        self.entry_stamp_date.grid(row=4, column=1, columnspan=2, padx=8, pady=8, sticky="ew")

        ctk.CTkButton(card, text="Start Processing", height=44, fg_color=COLOR_WARN, command=self._stamp_start_process).grid(row=5, column=0, padx=16, pady=(16, 16), sticky="w")
        
        log_frame = ctk.CTkFrame(self.content, corner_radius=12)
        log_frame.grid(row=1, column=0, sticky="nsew", padx=8, pady=8)
        log_frame.grid_rowconfigure(0, weight=1); log_frame.grid_columnconfigure(0, weight=1)
        self.tx_stamp_log = tk.Text(log_frame, height=10, wrap="word", relief="sunken", borderwidth=1, bg="#fdfdfd", font=("Segoe UI", 9))
        self.tx_stamp_log.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        scrollbar = ttk.Scrollbar(log_frame, command=self.tx_stamp_log.yview)
        scrollbar.grid(row=0, column=1, sticky="ns"); self.tx_stamp_log.configure(yscrollcommand=scrollbar.set)

    def _stamp_browse_input_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.entry_stamp_input_folder.delete(0, "end"); self.entry_stamp_input_folder.insert(0, folder_path)

    def _stamp_browse_output_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.entry_stamp_output_folder.delete(0, "end"); self.entry_stamp_output_folder.insert(0, folder_path)

    def _stamp_start_process(self):
        input_folder = self.entry_stamp_input_folder.get()
        output_folder = self.entry_stamp_output_folder.get()
        station_value = self.entry_stamp_station.get().strip()
        date_value = self.entry_stamp_date.get().strip() # Get date value

        if not input_folder or not output_folder:
            self._safe_show_error("Error", "Please select an input folder and an output folder."); return
        if not os.path.isdir(input_folder):
            self._safe_show_error("Error", "The selected input path is not a valid folder."); return
        if not os.path.isdir(output_folder):
            self._safe_show_error("Error", "The selected output path is not a valid folder."); return

        self.tx_stamp_log.delete(1.0, tk.END)
        self._safe_log(self.tx_stamp_log, "Starting process...\n")
        
        # Start the folder processing in a new thread
        threading.Thread(target=self._stamp_process_folder, args=(input_folder, output_folder, station_value, date_value), daemon=True).start()

    # ***** END OF MODIFICATION *****


    # ---------- Effectivity / Util / Check ----------
    def _tab_effectivity(self):
        """
        EFFECTIVITY / TCM Module - Currently Commented Out (Development Mode)
        
        This function was originally designed to create the EFFECTIVITY and TCM interface
        for maintenance documentation management, but is currently disabled and
        only shows an "Under Development" message in English.
        """
        # Show Under Development message
        print("[DEV] Development Mode - EFFECTIVITY / TCM Module")
        messagebox.showinfo(
            "EFFECTIVITY / TCM", 
            "🛠️ UNDER DEVELOPMENT 🛠️\n\n"
            "✈️ Advanced Effectivity & TCM System\n\n"
            "This module is currently under development\n"
            "and will be available in a future release.\n\n"
            "Thank you for your patience!\n"
            "🚀 REDSEA Airlines Aviation Toolkit"
        )
        
        # Original code commented out during development
        # card = ctk.CTkFrame(self.content, corner_radius=12)
        # card.grid(row=0,column=0,sticky="ew",padx=8,pady=8)
        # card.grid_columnconfigure(0,weight=1)
        # ctk.CTkLabel(card,text="EFFECTIVITY / TCM",font=("Segoe UI",18,"bold"),text_color=COLOR_PRIMARY).grid(row=0,column=0,padx=16,pady=(16,8),sticky="w")

    def _load_excel_generic(self, tag):
        """
        Generic Excel File Loader - Currently Commented Out (Development Mode)
        
        This function was designed to load Excel files for effectivity analysis
        but is currently disabled and only shows an "Under Development" message.
        
        Args:
            tag (str): Identifier for the type of data ('fleet' or 'tcm')
        """
        # Show Under Development message
        print("[📊 Development Mode] - Excel Loading Function")
        messagebox.showinfo(
            "Excel File Loader", 
            "📂 UNDER DEVELOPMENT 📂\n\n"
            "📊 Excel Data Processing System\n\n"
            "This function is currently under development\n"
            "and will be available in a future release.\n\n"
            "Thank you for your patience!\n"
            "🚀 REDSEA Airlines Aviation Toolkit"
        )
        
        # Original code commented out during development
        # def job():
        #     f = filedialog.askopenfilename(filetypes=[("Excel","*.xlsx;*.xls")])
        #     if not f: return

    def _pick_chapters_folder(self):
        """
        Select Chapters Folder - Currently Commented Out (Development Mode)
        
        This function was designed to select maintenance manual chapters folders
        but is currently disabled and only shows an "Under Development" message.
        """
        # Show Under Development message
        print("[📁 Development Mode] - Chapter Selection Function")
        messagebox.showinfo(
            "Chapter Selection", 
            "📁 UNDER DEVELOPMENT 📁\n\n"
            "📚 Maintenance Manual Chapter System\n\n"
            "This function is currently under development\n"
            "and will be available in a future release.\n\n"
            "Thank you for your patience!\n"
            "🚀 REDSEA Airlines Aviation Toolkit"
        )
        
        # Original code commented out during development
        # d = filedialog.askdirectory()
        # if d: self.lbl_chapters.configure(text=d)

    def _log_effectivity(self, msg): 
        """
        EFFECTIVITY / TCM (Task Card Management) Module
        
        This function handles effectivity reporting and task card management for aviation maintenance.
        Features include compliance tracking, task scheduling, and maintenance documentation.
        Currently under development with enhanced capabilities coming soon.
        
        Args:
            msg (str): Message context for effectivity operation
        """
        # Display beautiful development status message
        print("[🚧 Development Mode] - EFFECTIVITY / TCM Module")
        messagebox.showinfo(
            "EFFECTIVITY / TCM", 
            " ️ UNDER DEVELOPMENT  ️\n\n"
            "✈️ Advanced Effectivity & TCM System\n\n"
            "🔧 Coming Soon:\n"
            "   • Automated Task Card Generation\n"
            "   • Compliance Tracking & Reporting\n"
            "   • Enhanced Documentation Management\n"
            "   • Real-time Effectivity Analysis\n\n"
            "🚀 Professional Aviation Maintenance Tools\n"
            "   Enhanced functionality in development!"
        )
    def _export_effectivity(self):
        """
        Export Effectivity Data - Currently Commented Out (Development Mode)
        
        This function was designed to export effectivity analysis results
        but is currently disabled and only shows an "Under Development" message.
        """
        # Show Under Development message
        print("[📤 Development Mode] - Export Effectivity Function")
        messagebox.showinfo(
            "Export Effectivity", 
            "📤 UNDER DEVELOPMENT 📤\n\n"
            "📊 Data Export System\n\n"
            "This function is currently under development\n"
            "and will be available in a future release.\n\n"
            "Thank you for your patience!\n"
            "🚀 REDSEA Airlines Aviation Toolkit"
        )
        
        # Original code commented out during development
        # if self.eff_df is None or self.eff_df.empty: messagebox.showwarning("Export","No data to export."); return
        # out = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel","*.xlsx")])

    # ---------- Check Control ----------
    def _tab_check_control(self):
        """
        Check Control Module - Currently Commented Out (Development Mode)
        
        This function was originally designed to create the Check Control interface
        for aircraft maintenance monitoring, but is currently disabled and
        only shows an "Under Development" message in English.
        """
        # Show Under Development message
        print("[DEV] Development Mode - Check Control Module")
        messagebox.showinfo(
            "Check Control", 
            "⚙️ UNDER DEVELOPMENT ⚙️\n\n"
            "✅ Advanced Check Control System\n\n"
            "This module is currently under development\n"
            "and will be available in a future release.\n\n"
            "Thank you for your patience!\n"
            "🚀 REDSEA Airlines Aviation Toolkit"
        )
        
        # Original code commented out during development
        # card = ctk.CTkFrame(self.content, corner_radius=12)
        # card.grid(row=0,column=0,sticky="ew",padx=8,pady=8)
        # card.grid_columnconfigure(0,weight=1)
        # ctk.CTkLabel(card,text="Check Control",font=("Segoe UI",18,"bold"),text_color=COLOR_PRIMARY).grid(row=0,column=0,padx=16,pady=(16,8),sticky="w")
    def _load_check_csv(self):
        """
        Check Control Data Management System
        
        This function manages maintenance check schedules, compliance monitoring,
        and check control workflows for comprehensive aircraft maintenance tracking.
        Provides automated scheduling and progress monitoring capabilities.
        Currently under development with enhanced features.
        """
        # Display beautiful development status message
        print("[🔍 Development Mode] - Check Control Module")
        messagebox.showinfo(
            "Check Control", 
            "⚙️ UNDER DEVELOPMENT ⚙️\n\n"
            "✅ Advanced Check Control System\n\n"
            "🔧 Coming Soon:\n"
            "   • Intelligent Check Scheduling\n"
            "   • Real-time Compliance Monitoring\n"
            "   • Automated Progress Tracking\n"
            "   • Advanced Reporting & Analytics\n\n"
            "🚀 Professional Maintenance Management\n"
            "   Enhanced check control capabilities in development!"
        )
        def job():
            f=filedialog.askopenfilename(filetypes=[("CSV","*.csv")])
            if not f: return
            try: df=pd.read_csv(f); self.check_df=df; self.after(0, lambda: df_to_tree(self.tree_check, df, max_rows=1000))
            except Exception as e: self._safe_show_error("CSV Load Error", str(e))
        threading.Thread(target=job, daemon=True).start()

    # ---------- Utilization ----------
    def _tab_utilization(self):
        """
        Utilization Module - Currently Commented Out (Development Mode)
        
        This function was originally designed to create the Utilization interface
        for aircraft flight hours and cycles tracking, but is currently disabled and
        only shows an "Under Development" message in English.
        """
        # Show Under Development message
        print("[DEV] Development Mode - Utilization Module")
        messagebox.showinfo(
            "Utilization", 
            "📈 UNDER DEVELOPMENT 📈\n\n"
            "✈️ Advanced Utilization Tracking System\n\n"
            "This module is currently under development\n"
            "and will be available in a future release.\n\n"
            "Thank you for your patience!\n"
            "🚀 REDSEA Airlines Aviation Toolkit"
        )
        
        # Original code commented out during development
        # top=ctk.CTkFrame(self.content, corner_radius=12)
        # top.grid(row=0,column=0,sticky="ew",padx=8,pady=8)
        # top.grid_columnconfigure((0,1,2,3,4),weight=1)
        # ctk.CTkLabel(top,text="Utilization",font=("Segoe UI",18,"bold"),text_color=COLOR_PRIMARY).grid(row=0,column=0,padx=16,pady=(16,8),sticky="w")
    def _load_util_csv(self):
        """
        Aircraft Utilization Tracking System
        
        This function manages aircraft utilization data including flight hours, cycles,
        maintenance intervals, and fleet utilization analysis for comprehensive
        aircraft lifecycle management and maintenance planning.
        Currently under development with advanced analytics capabilities.
        """
        # Display beautiful development status message
        print("[📊 Development Mode] - Utilization Tracking Module")
        messagebox.showinfo(
            "Utilization Tracking", 
            "📈 UNDER DEVELOPMENT 📈\n\n"
            "✈️ Advanced Utilization Analytics System\n\n"
            "🔧 Coming Soon:\n"
            "   • Smart Flight Hours Tracking\n"
            "   • Automated Cycle Monitoring\n"
            "   • Predictive Maintenance Scheduling\n"
            "   • Fleet Utilization Optimization\n\n"
            "🚀 Professional Aircraft Management\n"
            "   Enhanced utilization analytics in development!"
        )
        return  # Exit function after showing development message
        
        def job():
            f=filedialog.askopenfilename(filetypes=[("CSV","*.csv")]);
            if not f: return
            try: df=pd.read_csv(f); self.util_df=df; self.after(0, lambda: df_to_tree(self.tree_util, df, max_rows=None)); self.after(0, lambda: self.lbl_util_path.configure(text=f))
            except Exception as e: self._safe_show_error("CSV Load Error", str(e))
        threading.Thread(target=job, daemon=True).start()
    def _save_util_csv(self):
        """
        Save Aircraft Utilization Data
        
        This function exports aircraft utilization tracking data to CSV format
        for external analysis, reporting, and backup purposes.
        Currently under development with enhanced export capabilities.
        """
        # Display beautiful development status message
        print("[💾 Development Mode] - Save Utilization Data")
        messagebox.showinfo(
            "Save Utilization Data", 
            "📈 UNDER DEVELOPMENT 📈\n\n"
            "💾 Advanced Export System\n\n"
            "🔧 Coming Soon:\n"
            "   • Multiple Export Formats\n"
            "   • Custom Report Generation\n"
            "   • Automated Backup Features\n"
            "   • Enhanced Data Security\n\n"
            "🚀 Professional Export Tools!"
        )
    def _util_current_inputs_row(self): 
        """
        Get current input values as a dictionary
        
        Returns:
            dict: Current values from all utilization input fields
        """
        return {k:self.util_inputs[k].get().strip() for k in self.util_inputs}
    def _util_clear_inputs(self):
        """
        Clear Utilization Input Fields
        
        This function resets all utilization data input fields to provide
        a clean interface for new data entry operations.
        Currently under development with enhanced input management.
        """
        # Display beautiful development status message
        print("[🧹 Development Mode] - Clear Input Fields")
        messagebox.showinfo(
            "Clear Input Fields", 
            "📈 UNDER DEVELOPMENT 📈\n\n"
            "🧹 Advanced Input Management\n\n"
            "🔧 Coming Soon:\n"
            "   • Smart Field Validation\n"
            "   • Auto-save Draft Data\n"
            "   • Enhanced User Interface\n"
            "   • Advanced Input Controls\n\n"
            "🚀 Professional Input Management!"
        )
    def _util_add(self):
        """
        Add Aircraft Utilization Record
        
        This function processes and adds new aircraft utilization data to the tracking system.
        Handles flight hours, cycles, and maintenance intervals for comprehensive monitoring.
        Currently under development with enhanced data processing capabilities.
        """
        # Display beautiful development status message
        print("[📊 Development Mode] - Add Utilization Record")
        messagebox.showinfo(
            "Add Utilization Data", 
            "📈 UNDER DEVELOPMENT 📈\n\n"
            "✈️ Advanced Data Entry System\n\n"
            "🔧 Coming Soon:\n"
            "   • Smart Data Validation\n"
            "   • Automated Calculations\n"
            "   • Real-time Updates\n"
            "   • Enhanced Data Processing\n\n"
            "🚀 Professional Data Management Tools!"
        )
    def _util_delete_selected(self):
        """
        Delete Selected Utilization Records
        
        This function removes selected utilization records from the tracking database
        and refreshes the display to reflect changes.
        Currently under development with enhanced deletion capabilities.
        """
        # Display beautiful development status message
        print("[🗑️ Development Mode] - Delete Utilization Records")
        messagebox.showinfo(
            "Delete Records", 
            "📈 UNDER DEVELOPMENT 📈\n\n"
            "🗑️ Advanced Record Management\n\n"
            "🔧 Coming Soon:\n"
            "   • Smart Deletion Controls\n"
            "   • Data Recovery Features\n"
            "   • Enhanced Safety Checks\n"
            "   • Advanced Undo Capabilities\n\n"
            "🚀 Professional Data Management!"
        )

    # ---------- Hash Functions ----------
    def hash_function_md5(self):
        """
        MD5 Hash Generation Function
        
        This function provides MD5 hash generation capabilities for data integrity
        verification and security purposes in aviation maintenance documentation.
        Currently under development for enhanced cryptographic features.
        """
        # Feature currently under development - Advanced hash generation coming soon
        print("[Development Mode] - MD5 Hash Function")
        messagebox.showinfo("MD5 Hash Generator", 
                          "🔐 UNDER DEVELOPMENT 🔐\n\n"
                          "MD5 Hash Generation System is being developed.\n\n"
                          "This will include:\n"
                          "• File integrity verification\n"
                          "• Document authentication\n"
                          "• Data security validation\n\n"
                          "Enhanced cryptographic functionality coming soon!")

    def hash_function_sha256(self):
        """
        SHA-256 Hash Generation Function
        
        This function provides SHA-256 hash generation capabilities for enhanced
        security and data integrity verification in maintenance documentation systems.
        Currently under development for advanced cryptographic operations.
        """
        # Feature currently under development - Advanced SHA-256 capabilities coming soon
        print("[Development Mode] - SHA-256 Hash Function")
        messagebox.showinfo("SHA-256 Hash Generator", 
                          "🔐 UNDER DEVELOPMENT 🔐\n\n"
                          "SHA-256 Hash Generation System is being developed.\n\n"
                          "This will include:\n"
                          "• Advanced encryption algorithms\n"
                          "• Secure document verification\n"
                          "• Enhanced data protection\n\n"
                          "Professional cryptographic tools coming soon!")

    def hash_function_blake2(self):
        """
        BLAKE2 Hash Generation Function
        
        This function provides BLAKE2 hash generation capabilities for modern
        cryptographic operations and high-performance data integrity verification.
        Currently under development for next-generation security features.
        """
        # Feature currently under development - Modern BLAKE2 hash system coming soon
        print("[Development Mode] - BLAKE2 Hash Function")
        messagebox.showinfo("BLAKE2 Hash Generator", 
                          "🔐 UNDER DEVELOPMENT 🔐\n\n"
                          "BLAKE2 Hash Generation System is being developed.\n\n"
                          "This will include:\n"
                          "• Modern cryptographic algorithms\n"
                          "• High-performance hashing\n"
                          "• Advanced security protocols\n\n"
                          "Next-generation hash functionality coming soon!")

    # ---------- CMP / TCM Tasks (Indexed) + Cover Merge ----------
    def _tab_cmp_tcm_tasks(self):
        card=ctk.CTkFrame(self.content,corner_radius=16, fg_color=COLOR_CARD_BG); card.grid(row=0,column=0,sticky="ew",padx=12,pady=12); card.grid_columnconfigure((0,1,2,3),weight=1)
        ctk.CTkLabel(card,text="CMP / TCM Tasks (Indexed)",font=("Segoe UI",22,"bold"),text_color=COLOR_PRIMARY).grid(row=0,column=0,padx=20,pady=(20,12),sticky="w")
        ctk.CTkButton(card,text="📊 MPD RSD Excel…",fg_color=COLOR_OK,hover_color="#2f855a",height=40,font=("Segoe UI",12,"bold"),corner_radius=10,command=self._pick_mpd_rsd_excel).grid(row=1,column=0,padx=20,pady=8,sticky="w"); self.lbl_mpd_rsd_excel=ctk.CTkLabel(card,text="No Excel selected", font=("Segoe UI",11), text_color=COLOR_TEXT_DARK); self.lbl_mpd_rsd_excel.grid(row=1,column=1,columnspan=3,padx=12,pady=8,sticky="w")
      #  ctk.CTkLabel(card,text="Sheet Name:", font=("Segoe UI",12,"bold"), text_color=COLOR_PRIMARY).grid(row=2,column=0,padx=20,pady=8,sticky="w"); self.entry_sheet_name=ctk.CTkEntry(card,placeholder_text="e.g., MPD RSD, Sheet1, Sheet2", height=35, font=("Segoe UI",11), corner_radius=8); self.entry_sheet_name.grid(row=2,column=1,columnspan=3,padx=12,pady=8,sticky="ew")
        ctk.CTkButton(card,text="📁 TCM Folder…",fg_color=COLOR_OK,hover_color="#2f855a",height=40,font=("Segoe UI",12,"bold"),corner_radius=10,command=self._pick_tcm_dir).grid(row=3,column=0,padx=20,pady=8,sticky="w"); self.lbl_tcm_dir=ctk.CTkLabel(card,text="No TCM folder selected", font=("Segoe UI",11), text_color=COLOR_TEXT_DARK); self.lbl_tcm_dir.grid(row=3,column=1,columnspan=3,padx=12,pady=8,sticky="w")
        ctk.CTkButton(card,text="📂 Output Folder…",fg_color=COLOR_ACCENT,hover_color=COLOR_HOVER,height=40,font=("Segoe UI",12,"bold"),corner_radius=10,command=self._pick_output_folder_cmp_tcm).grid(row=4,column=0,padx=20,pady=8,sticky="w"); self.lbl_out_cmp_tcm=ctk.CTkLabel(card,text="Output: not set", font=("Segoe UI",11), text_color=COLOR_TEXT_DARK); self.lbl_out_cmp_tcm.grid(row=4,column=1,columnspan=3,padx=12,pady=8,sticky="w")

        # Rebuild Index button
        ctk.CTkButton(card,text="🔧 Rebuild TCM Index",fg_color=COLOR_WARN,hover_color="#dd6b20",height=40,font=("Segoe UI",12,"bold"),corner_radius=10,command=self._rebuild_tcm_index).grid(row=5,column=0,padx=20,pady=8,sticky="w")

        # Modern options with Refresh button
        opt = ctk.CTkFrame(card, fg_color=COLOR_BG_LIGHT, corner_radius=12); opt.grid(row=6,column=0,columnspan=4,sticky="ew",padx=20,pady=(8,16)); opt.grid_columnconfigure((0,1,2,3,4),weight=1)
        
        # Check selection row
        ctk.CTkLabel(opt,text="Check:", font=("Segoe UI",12,"bold"), text_color=COLOR_PRIMARY).grid(row=0,column=0,padx=12,pady=12,sticky="w")
        # Initialize with default values
        default_checks = [f"A{i}" for i in range(1,12)] + ["120DY","240DY","12MO","16MO"] + [f"C{i}" for i in range(1,7)]
        self.cmb_check2 = ctk.CTkComboBox(opt, values=default_checks, height=35, font=("Segoe UI",11), corner_radius=8); self.cmb_check2.grid(row=0,column=1,padx=12,pady=12,sticky="ew")
        
        # Refresh button
        ctk.CTkButton(opt, text="🔄 Refresh", width=90, height=35, fg_color=COLOR_ACCENT, hover_color=COLOR_HOVER, font=("Segoe UI",11,"bold"), corner_radius=8, command=self._refresh_available_checks).grid(row=0,column=2,padx=12,pady=12,sticky="w")
        
        ctk.CTkLabel(opt,text="Aircraft:", font=("Segoe UI",12,"bold"), text_color=COLOR_PRIMARY).grid(row=0,column=3,padx=12,pady=12,sticky="w")
        self.cmb_aircraft2=ctk.CTkComboBox(opt, values=["SU-RSA","SU-RSB","SU-RSC","SU-RSD"], height=35, font=("Segoe UI",11), corner_radius=8); self.cmb_aircraft2.grid(row=0,column=4,padx=12,pady=12,sticky="ew")

        ctk.CTkButton(card,text="🚀 Generate Task Cards",height=50,fg_color=COLOR_OK,hover_color="#2f855a",font=("Segoe UI",14,"bold"),corner_radius=12,command=self._generate_task_cards_indexed).grid(row=7,column=0,padx=20,pady=(0,20),sticky="w")
        
        # Modern log area with Copy button
        log_frame = ctk.CTkFrame(self.content, corner_radius=16, fg_color=COLOR_CARD_BG)
        log_frame.grid(row=1,column=0,sticky="nsew",padx=12,pady=12)
        log_frame.grid_columnconfigure(0, weight=1)
        log_frame.grid_rowconfigure(1, weight=1)
        
        # Copy button frame
        btn_frame = ctk.CTkFrame(log_frame, fg_color="transparent")
        btn_frame.grid(row=0, column=0, sticky="ew", padx=15, pady=8)
        btn_frame.grid_columnconfigure(0, weight=1)
        
        ctk.CTkButton(btn_frame, text="📋 Copy Log", command=self._copy_log, height=35, font=("Segoe UI",11,"bold"), corner_radius=8).grid(row=0, column=0, padx=(0, 12), sticky="w")
        # ctk.CTkButton(btn_frame, text="🔍 Extract All Tasks", command=self._extract_all_tasks).grid(row=0, column=1, padx=8, sticky="w")
        
        # Information about column usage
        info_label = ctk.CTkLabel(btn_frame, text="📍 Column 24 for CHECK codes", font=("Segoe UI", 10), text_color="gray")
        info_label.grid(row=0, column=1, padx=12, sticky="e")
        
        self.tx_cmp_tcm_log = tk.Text(log_frame, height=14, font=("Consolas",10), bg=COLOR_BG_LIGHT, fg=COLOR_TEXT_DARK, relief="flat", bd=0)
        self.tx_cmp_tcm_log.grid(row=1, column=0, sticky="nsew", padx=15, pady=15)

    def _copy_log(self):
        """Copy log content to clipboard"""
        try:
            log_content = self.tx_cmp_tcm_log.get("1.0", tk.END)
            self.clipboard_clear()
            self.clipboard_append(log_content)
            self._safe_log(self.tx_cmp_tcm_log, "📋 Log copied to clipboard!\n")
        except Exception as e:
            self._safe_log(self.tx_cmp_tcm_log, f"❌ Error copying log: {e}\n")

    def _extract_all_tasks(self):
        """Extract tasks for all check codes found in the Excel file"""
        try:
            if not hasattr(self, 'mpd_rsd_excel') or not self.mpd_rsd_excel or not os.path.exists(self.mpd_rsd_excel):
                self._safe_log(self.tx_cmp_tcm_log, "❌ Please select MPD RSD Excel file first\n")
                return
            
            sheet_name = self.entry_sheet_name.get().strip() if hasattr(self, 'entry_sheet_name') else None
            
            self._safe_log(self.tx_cmp_tcm_log, f"🔍 Extracting all tasks from Excel...\n")
            self._safe_log(self.tx_cmp_tcm_log, f"📊 File: {os.path.basename(self.mpd_rsd_excel)}\n")
            if sheet_name:
                self._safe_log(self.tx_cmp_tcm_log, f"📋 Sheet: {sheet_name}\n")
            
            # Note: This function would need to be re-implemented since the original was removed
            # For now, we'll show a message that this feature needs to be implemented
            self._safe_log(self.tx_cmp_tcm_log, "ℹ️ Extract All Tasks feature is currently not implemented\n")
            self._safe_log(self.tx_cmp_tcm_log, "Please use the individual check extraction instead\n")
            
        except Exception as e:
            self._safe_log(self.tx_cmp_tcm_log, f"❌ Error during task extraction: {e}\n")

    # CMP/TCM-specific stamping process
    def _stamp_document(self, doc, station_value, date_value, tail_number):
        """
        Updates the document with the provided information by replacing content
        
        Args:
            doc: PDF document to update
            station_value: Station value to set
            date_value: Date value to set
            tail_number: Tail number to set
        """
        for page_num, page in enumerate(doc):
            try:
                airline_card_no = "RC"
                found_number = False
                
                # Try to log - use appropriate log widget based on context
                log_widget = getattr(self, 'tx_cmp_tcm_log', None) or getattr(self, 'tx_stamp_log', None)
                if log_widget:
                    self._safe_log(log_widget, f"   - Page {page_num + 1}: Processing...\n")
                
                # --- Method 1: Location-based text search and replace ---
                search_terms = ["BOEING CARD NO.", "BOEING CARD NO"]
                search_instances = []
                for term in search_terms:
                    search_instances = page.search_for(term)
                    if search_instances:
                        if log_widget:
                            self._safe_log(log_widget, f"     > Method 1: Found label '{term}'.\n")
                        break
                
                if search_instances:
                    label_rect = search_instances[0]
                    search_area = fitz.Rect(label_rect.x1 - 5, label_rect.y0 - 5, label_rect.x1 + 300, label_rect.y1 + 5)
                    extracted_text = page.get_text("text", clip=search_area).strip()
                    
                    if extracted_text:
                        if log_widget:
                            self._safe_log(log_widget, f"     > Text found in area: '{extracted_text.replace(chr(10), ' ')}'\n")
                        boeing_match = re.search(r'([\d-]+)', extracted_text)
                        if boeing_match:
                            full_boeing_no = boeing_match.group(1).strip()
                            if log_widget:
                                self._safe_log(log_widget, f"     > Extracted number: {full_boeing_no}\n")
                            parts = full_boeing_no.split('-')
                            if len(parts) >= 3:
                                airline_card_no = f"RC{parts[0]}-{parts[1]}-{parts[2]}"
                                found_number = True
                            elif len(parts) >= 2:
                                airline_card_no = f"RC{parts[0]}-{parts[1]}"
                                found_number = True
                        elif log_widget:
                            self._safe_log(log_widget, f"     > WARNING: No number pattern found in the area.\n")
                    else:
                        self._safe_log(log_widget, f"     > WARNING: Found label, but the area next to it is empty.\n")

                # --- Method 2: Fallback to full page text search ---
                if not found_number:
                    if log_widget:
                        self._safe_log(log_widget, f"     > Method 1 failed. Trying Method 2 (Full page search).\n")
                    page_text = page.get_text()
                    boeing_match = re.search(r"BOEING\s+CARD\s+NO\.?\s*([\d-]+)", page_text, re.IGNORECASE)
                    if boeing_match:
                        full_boeing_no = boeing_match.group(1).strip()
                        if log_widget:
                            self._safe_log(log_widget, f"     > Method 2 found number: {full_boeing_no}\n")
                            parts = full_boeing_no.split('-')
                            if len(parts) >= 3:
                                airline_card_no = f"RC{parts[0]}-{parts[1]}-{parts[2]}"
                            elif len(parts) >= 2:
                                airline_card_no = f"RC{parts[0]}-{parts[1]}"
                    else:
                        if log_widget:
                            self._safe_log(log_widget, f"     > Method 2 failed. Number will be 'RC'.\n")
                
                # --- Content Replacement ---
                # First try to find each label and update its value
                items_to_update = [
                    ("TAIL NUMBER", tail_number), 
                    ("AIRLINE CARD NO", airline_card_no), 
                    ("STATION", station_value),
                    ("DATE", date_value)
                ]
                
                found_labels = False
                for label, value in items_to_update:
                    if not value: continue
                    instances = page.search_for(label)
                    if instances:
                        found_labels = True
                        for inst in instances:
                            # Create rectangle area after the label for replacing content
                            rect = fitz.Rect(inst.x0, inst.y1 - 5, inst.x1 + 200, inst.y1 + 20)
                            # Remove any existing content in the area
                            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                            # Add new text with better formatting
                            page.insert_text((inst.x0, inst.y1 + 12), value, fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")
                
                # If no labels found, try to find common form fields
                if not found_labels:
                    # Try to find common form field labels
                    alternative_labels = {
                        "TAIL NUMBER": ["TAIL", "A/C", "AIRCRAFT", "REG"],
                        "AIRLINE CARD NO": ["CARD NO", "CARD NUMBER", "AIRLINE CARD", "DOCUMENT NO"],
                        "STATION": ["STA", "BASE", "LOCATION"],
                        "DATE": ["DATE", "DUE DATE", "ISSUE DATE"]
                    }
                    
                    for original_label, value in items_to_update:
                        if not value: continue
                        alt_found = False
                        
                        for alt_label in alternative_labels.get(original_label, []):
                            instances = page.search_for(alt_label)
                            if instances:
                                alt_found = True
                                for inst in instances:
                                    # Clear existing content and insert new text
                                    rect = fitz.Rect(inst.x1, inst.y0 - 2, inst.x1 + 200, inst.y1 + 15)
                                    page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                                    page.insert_text((inst.x1 + 5, inst.y0 + ((inst.y1 - inst.y0) / 2)), 
                                                    value, fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")
                    
                    # Special handling for CMP/TCM task cards - common layout formats
                    # Try to determine if this is a Boeing task card format
                    if page.search_for("BOEING") or page.search_for("TASK CARD"):
                        # This appears to be a Boeing format card - look for structured layout
                        for i in range(page.rect.width // 3):  # Scan left third of page
                            for j in range(page.rect.height // 3):  # Scan top third of page
                                rect = fitz.Rect(i, j, i+200, j+20)  # Small window
                                text = page.get_text("text", clip=rect).lower()
                                
                                # Check for common field locations and replace content
                                content_rect = fitz.Rect(rect.x1, rect.y0, rect.x1 + 200, rect.y1)
                                if "tail" in text or "aircraft" in text:
                                    page.draw_rect(content_rect, color=(1, 1, 1), fill=(1, 1, 1))
                                    page.insert_text((rect.x1 + 5, rect.y0 + ((rect.y1 - rect.y0) / 2)), 
                                                    tail_number, fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")
                                    alt_found = True
                                
                                if "airline" in text or "card" in text:
                                    page.draw_rect(content_rect, color=(1, 1, 1), fill=(1, 1, 1))
                                    page.insert_text((rect.x1 + 5, rect.y0 + ((rect.y1 - rect.y0) / 2)), 
                                                    airline_card_no, fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")
                                    alt_found = True
                                    
                                if "station" in text or "sta" in text:
                                    page.draw_rect(content_rect, color=(1, 1, 1), fill=(1, 1, 1))
                                    page.insert_text((rect.x1 + 5, rect.y0 + ((rect.y1 - rect.y0) / 2)), 
                                                    station_value, fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")
                                    alt_found = True
                                    
                                if "date" in text:
                                    page.insert_text((rect.x0, rect.y1 + 15), date_value, fontsize=12, 
                                                    color=(0, 0, 0), fontname="Helvetica-Bold")
                                    alt_found = True
                        
                        # If we still didn't find locations, use standard Boeing card positions
                        if not alt_found:
                            # For Boeing task cards - standard positions for these fields
                            header_y = 100  # Y position in header section
                            if page.search_for("BOEING CARD NO"):
                                # Standard Boeing task card format - clear areas and insert new text
                                standard_positions = [
                                    (160, 220, 400, 260, tail_number),      # TAIL NUMBER position
                                    (160, 160, 400, 200, airline_card_no),  # CARD NO position
                                    (160, 280, 400, 320, station_value),    # STATION position
                                    (160, 100, 400, 140, date_value)        # DATE position
                                ]
                                
                                for x, y1, x2, y2, value in standard_positions:
                                    if value:
                                        rect = fitz.Rect(x, y1, x2, y2)
                                        page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                                        page.insert_text((x + 5, y1 + ((y2 - y1) / 2)), 
                                                        value, fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")
                                alt_found = True
                    
                    # If still not found, put the stamp in common locations based on the label type
                    if not alt_found:
                        fallback_positions = [
                            (50, 130, 400, 170, "TAIL", tail_number),
                            (50, 150, 400, 190, "CARD", airline_card_no),
                            (50, 170, 400, 210, "STA", station_value),
                            (50, 190, 400, 230, "DATE", date_value)
                        ]
                        
                        for x, y1, x2, y2, label, value in fallback_positions:
                            if value:
                                rect = fitz.Rect(x, y1, x2, y2)
                                page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                                page.insert_text((x, y1 + ((y2 - y1) / 2)), 
                                                f"{label}: {value}", fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")  
                        if station_value:
                            page.insert_text((50, 190), f"STA: {station_value}", fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")
                        if date_value:
                            page.insert_text((50, 210), f"DATE: {date_value}", fontsize=12, color=(0, 0, 0), fontname="Helvetica-Bold")

            except Exception as e:
                if log_widget:
                    self._safe_log(log_widget, f"     > Error stamping page {page_num + 1}: {str(e)}\n")

    def _stamp_cmp_tcm_files(self, input_folder, output_folder, station_value, date_value):
        """Process all PDF files in the input folder for CMP/TCM module."""
        try:
            # Make sure output folder exists
            os.makedirs(output_folder, exist_ok=True)
            
            # Get all PDF files in the input folder - search recursively through subdirectories
            pdf_files = []
            total_files = 0
            
            # Walk through all subdirectories recursively
            for root, dirs, files in os.walk(input_folder):
                for file in files:
                    if file.lower().endswith('.pdf'):
                        pdf_files.append(os.path.join(root, file))
                        total_files += 1
            
            self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Found {total_files} PDF files to process.\n")
            
            if total_files == 0:
                self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] No PDF files found in folder: {input_folder}\n")
                self._safe_show_info("No Files Found", f"No PDF files were found in the selected folder.")
                return
            
            # Process each PDF file
            processed_count = 0
            for input_path in pdf_files:
                self._stamp_cmp_tcm_single_pdf(input_path, output_folder, station_value, date_value)
                processed_count += 1
                # Update progress every 5 files
                if processed_count % 5 == 0:
                    self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Progress: {processed_count}/{total_files} files processed...\n")
                
            self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Completed stamping process for all {total_files} files.\n")
            self._safe_show_info("Stamping Complete", f"Successfully completed stamping {total_files} files.")
            
        except Exception as e:
            self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Error during folder processing: {str(e)}\n")
            self._safe_show_error("Stamping Error", f"An error occurred during stamping: {str(e)}")
    
    # Define the single PDF stamping function for CMP/TCM module
    def _stamp_cmp_tcm_single_pdf(self, input_file, output_folder, station_value, date_value):
        try:
            self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Processing file: {os.path.basename(input_file)}\n")
            
            # Open the PDF
            doc = fitz.open(input_file)
            
            # Try to detect tasks from the PDF
            tasks = []
            plane_code_from_cover = None
            
            # Get plane code from the cover page if possible
            if doc.page_count > 0:
                cover_page = doc[0]
                cover_text = cover_page.get_text()
                
                for code in self.TAIL_MAP.keys():
                    if code in cover_text:
                        plane_code_from_cover = code
                        self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Detected Plane Code from Cover: {plane_code_from_cover} -> {self.TAIL_MAP[plane_code_from_cover]}\n")
                        break
            
            # Scan all pages for task numbers
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                
                # Extract tasks using regex
                page_tasks = re.findall(r"(\d{2}-\d{3}-\d{2})", text)
                tasks.extend(page_tasks)
            
            # Remove duplicates
            tasks = list(set(tasks))
            
            if not tasks:
                self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] No tasks found in document: {os.path.basename(input_file)}\n")
                
                # If no tasks found, just stamp the document with the default plane code
                if plane_code_from_cover and plane_code_from_cover in self.TAIL_MAP:
                    # Process the document with stamping
                    self._stamp_document(doc, station_value, date_value, self.TAIL_MAP[plane_code_from_cover])
                    
                    # Save the stamped document (overwrite the original)
                    doc.save(input_file)
                    doc.close()
                    self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Updated: {os.path.basename(input_file)}\n")
                else:
                    # Try with first available plane code if not detected
                    first_code = list(self.TAIL_MAP.keys())[0]
                    self._stamp_document(doc, station_value, date_value, self.TAIL_MAP[first_code])
                    doc.save(input_file)
                    doc.close()
                    self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Updated with default plane code: {os.path.basename(input_file)}\n")
                return
            
            # Process tasks
            self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Found {len(tasks)} unique tasks in file.\n")
            
            # Determine which plane this task applies to
            plane_code = plane_code_from_cover
            
            # If we have a valid plane, stamp the document
            if plane_code and plane_code in self.TAIL_MAP:
                # Process the document with stamping
                self._stamp_document(doc, station_value, date_value, self.TAIL_MAP[plane_code])
                
                # For CMP/TCM module, we just update the existing file instead of creating copies
                doc.save(input_file)
                doc.close()
                self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Updated: {os.path.basename(input_file)}\n")
            else:
                # Try with first available plane code if not detected
                first_code = list(self.TAIL_MAP.keys())[0]
                self._stamp_document(doc, station_value, date_value, self.TAIL_MAP[first_code])
                doc.save(input_file)
                doc.close()
                self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Updated with default plane code: {os.path.basename(input_file)}\n")
        
        except Exception as e:
            self._safe_log(self.tx_cmp_tcm_log, f"[Stamping] Error processing file {os.path.basename(input_file)}: {str(e)}\n")
    
    def _pick_mpd_rsd_excel(self):
        f = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx;*.xls;*.xlsb")])
        if f:
            self.mpd_rsd_excel = f
            self.lbl_mpd_rsd_excel.configure(text=os.path.basename(f))
            self._safe_log(self.tx_cmp_tcm_log, f"Selected Excel: {f}\n")
            # Auto-refresh checks
            self._refresh_available_checks()

    def _pick_tcm_dir(self):
        d = filedialog.askdirectory()
        if d:
            self.tcm_dir = d
            self.lbl_tcm_dir.configure(text=d)
            self._safe_log(self.tx_cmp_tcm_log, f"Selected TCM Folder: {d}\n")
            # Auto-load index if exists
            self.indexer = TcmIndexer(d)
            if self.indexer.try_load_cache():
                self._safe_log(self.tx_cmp_tcm_log, f"Loaded existing index with {len(self.indexer.index)} PDFs.\n")
            else:
                self._safe_log(self.tx_cmp_tcm_log, "No index found. Please click 'Rebuild TCM Index'.\n")

    def _pick_output_folder_cmp_tcm(self):
        d = filedialog.askdirectory()
        if d:
            self.out_cmp_tcm = d
            self.lbl_out_cmp_tcm.configure(text=d)
            self._safe_log(self.tx_cmp_tcm_log, f"Output Folder: {d}\n")

    def _rebuild_tcm_index(self):
        if not self.tcm_dir or not os.path.isdir(self.tcm_dir):
            self._safe_show_error("Error", "Please select a valid TCM folder first.")
            return
        
        self.indexer = TcmIndexer(self.tcm_dir)
        
        def job():
            self._safe_log(self.tx_cmp_tcm_log, "Starting index build... (this may take a while)\n")
            def progress_cb(msg):
                self._safe_log(self.tx_cmp_tcm_log, msg)
                
            self.indexer.build_index(progress_callback=progress_cb)
            self._safe_log(self.tx_cmp_tcm_log, "Index build finished.\n")
            
        threading.Thread(target=job, daemon=True).start()

    def _extract_checks_from_excel(self, excel_path: str, sheet_name: str = None, check_column: int = 24):
        """
        Automatically extract CHECK list from specified column in Excel file
        
        Args:
            excel_path: Path to Excel file
            sheet_name: Sheet name (optional)
            check_column: Column number containing CHECK codes (default 24)
            
        Returns:
            List of CHECK codes found
        """
        try:
            if not excel_path or not os.path.exists(excel_path):
                return []
            
            # Read Excel file
            df = None
            try:
                if excel_path.lower().endswith('.xlsb'):
                    try:
                        df = pd.read_excel(excel_path, sheet_name=sheet_name, engine='pyxlsb')
                    except ImportError:
                        df = pd.read_excel(excel_path, sheet_name=sheet_name)
                else:
                    df = pd.read_excel(excel_path, sheet_name=sheet_name)
            except Exception:
                if not sheet_name:
                    df = pd.read_excel(excel_path, sheet_name=0)
                    
            if df is None or df.empty:
                return []
                
            # التأكد من وجود العمود المطلوب
            if check_column >= len(df.columns):
                return []
                
            # Extract values from specified column
            check_values = df.iloc[:, check_column].astype(str).unique()
            
            # تنظيف وفلترة القيم
            valid_checks = []
            check_patterns = [
                r'^A([1-9]|1[0-1])$',  # A1-A11
                r'^C[1-6]$',           # C1-C6
                r'^(120|240)\s*DY$',   # 120DY, 240DY
                r'^(12|16)\s*MO$',     # 12MO, 16MO
                r'^2000\s*FC$'         # 2000FC
            ]
            
            for value in check_values:
                value = str(value).strip().upper()
                if value and value != 'NAN':
                    # التحقق من أن القيمة تطابق أحد أنماط الـ CHECK
                    for pattern in check_patterns:
                        if re.match(pattern, value):
                            # تنظيف القيمة (إزالة المسافات من DY و MO)
                            clean_value = re.sub(r'\s+', '', value)
                            if clean_value not in valid_checks:
                                valid_checks.append(clean_value)
                            break
                            
            # ترتيب القائمة
            valid_checks.sort(key=lambda x: (
                0 if x.startswith('A') else 
                1 if x.endswith('DY') else 
                2 if x.endswith('MO') else 
                3 if x.endswith('FC') else 
                4 if x.startswith('C') else 5,
                x
            ))
            
            return valid_checks
            
        except Exception as e:
            if hasattr(self, 'tx_cmp_tcm_log'):
                self._safe_log(self.tx_cmp_tcm_log, f"❌ Error extracting CHECK list: {e}\n")
            return []
    
    def _normalize_check_code(self, check_code: str) -> str:
        """
        Standardize CHECK code format for easier comparison
        
        Args:
            check_code: Original CHECK code
            
        Returns:
            Standardized CHECK code format
        """
        if not check_code or not isinstance(check_code, str):
            return ""
            
        # تنظيف القيمة
        cleaned = str(check_code).strip().upper()
        
        if not cleaned or cleaned.lower() in ['nan', 'none']:
            return ""
        
        # تحويل التنسيقات المختلفة إلى تنسيق موحد
        transformations = [
            # Days: 120 DY, 120 DAYS -> 120DY
            (r'(\d+)\s*(DY|DAYS?)', r'\1DY'),
            # Months: 12 MO, 12 MONTHS -> 12MO  
            (r'(\d+)\s*(MO|MONTHS?)', r'\1MO'),
            # Flight Cycles: 2000 FC, 2000 CYCLES -> 2000FC
            (r'(\d+)\s*(FC|CYCLES?)', r'\1FC'),
            # Hours: 1000 HR, 1000 HOURS -> 1000HR
            (r'(\d+)\s*(HR|HOURS?)', r'\1HR'),
            # General: A 1, C 2 -> A1, C2
            (r'([A-Z])\s+(\d+)', r'\1\2'),
            # إزالة المسافات الزائدة
            (r'\s+', ''),
        ]
        
        result = cleaned
        for pattern, replacement in transformations:
            result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
        
        return result.upper()
    
    def _expand_tasks_with_subtasks(self, base_tasks: list, tcm_indexer=None):
        """
        Expand base tasks to include all related subtasks from TCM index
        
        Args:
            base_tasks: List of base task codes (e.g., ['52-020-00'])
            tcm_indexer: TCM indexer instance for finding subtasks
            
        Returns:
            Expanded list including base tasks and all found subtasks
        """
        if not base_tasks or not tcm_indexer:
            return base_tasks
            
        expanded_tasks = set(base_tasks)
        
        for base_task in base_tasks:
            try:
                # Enumerate exact subtasks as 4th segment: base_prefix-01..-10 and validate from index
                if '-' in base_task:
                    parts = base_task.split('-')
                    if len(parts) == 3:
                        # Base task is like "52-020-00", add subtasks 52-020-00-01..10
                        base_prefix = base_task
                        self._safe_log(self.tx_cmp_tcm_log, f"    🔍 Searching for subtasks of {base_task}...\n")
                        
                        found_count = 0
                        for i in range(1, 11):
                            subtask = f"{base_prefix}-{i:02d}"
                            pdf_path, run = tcm_indexer.find_best_occurrence_for_task(subtask)
                            if pdf_path and run:
                                expanded_tasks.add(subtask)
                                found_count += 1
                                self._safe_log(self.tx_cmp_tcm_log, f"      ✓ Found: {subtask}\n")
                        
                        if found_count > 0:
                            self._safe_log(self.tx_cmp_tcm_log, f"    ✅ Added {found_count} subtasks for {base_task}\n")
            except Exception as e:
                self._safe_log(self.tx_cmp_tcm_log, f"    ⚠️ Error expanding {base_task}: {e}\n")
                
        return sorted(list(expanded_tasks))
    
    def _extract_tasks_from_excel_mpd_rsd(self, excel_path: str, cmp_iss03: str, sheet_name: str = None, log_label: str = None, log_target: str = None):
        """
        Extract tasks from Excel MPD RSD based on CMP ISS03 (like A8, A1, etc.)
        
        Args:
            excel_path: Path to the Excel file
            cmp_iss03: Check code to search for (e.g., A1, A8, etc.)
            sheet_name: Optional sheet name to read from
            log_label: Optional label prefix for log messages
            log_target: Optional target for logging ('stamping' or default)
            
        Returns:
            List of task codes found for the specified check
        """
        # Configure logging target and label
        log_widget = getattr(self, 'tx_cmp_tcm_log', None)
        if (log_target or '').lower() == 'stamping' and hasattr(self, 'tx_stamp_log'):
            log_widget = self.tx_stamp_log
        if log_label is None and (log_target or '').lower() == 'stamping':
            log_label = 'Stamping'

        def _log(message: str):
            prefix = f"[{log_label}] " if log_label else ""
            try:
                self._safe_log(log_widget, prefix + message)
            except Exception:
                # Fallback to CMP/TCM log if specified widget is unavailable
                self._safe_log(self.tx_cmp_tcm_log, prefix + message)

        if not excel_path or not os.path.exists(excel_path):
            _log("❌ Excel file path is invalid or file does not exist\n")
            return []
            
        if not cmp_iss03 or not cmp_iss03.strip():
            _log("❌ CMP ISS03 check code is required\n")
            return []
        
        try:
            # Read Excel file (supports .xlsx, .xls, .xlsb)
            df = None
            
            if sheet_name and sheet_name.strip():
                try:
                    if excel_path.lower().endswith('.xlsb'):
                        try:
                            df = pd.read_excel(excel_path, sheet_name=sheet_name.strip(), engine='pyxlsb')
                        except ImportError:
                            df = pd.read_excel(excel_path, sheet_name=sheet_name.strip())
                    else:
                        df = pd.read_excel(excel_path, sheet_name=sheet_name.strip())
                    
                    _log(f"📊 Reading Excel file: {os.path.basename(excel_path)} (Sheet: {sheet_name.strip()})\n")
                    _log(f"📊 Excel shape: {df.shape[0]} rows x {df.shape[1]} columns\n")
                except Exception as e:
                    _log(f"⚠️ Failed to read specified sheet '{sheet_name}': {e}\n")
                    df = None
            
            if df is None:
                # Try to find MPD RSD sheet specifically
                try:
                    if excel_path.lower().endswith('.xlsb'):
                        xl_file = pd.ExcelFile(excel_path, engine='pyxlsb')
                        sheet_names = xl_file.sheet_names
                        _log(f"📋 Available sheets: {sheet_names}\n")
                        
                        mpd_sheet = None
                        for sheet in sheet_names:
                            if 'MPD RSD' in sheet.upper():
                                mpd_sheet = sheet
                                break
                        
                        if mpd_sheet:
                            df = pd.read_excel(excel_path, sheet_name=mpd_sheet, engine='pyxlsb')
                            _log(f"📊 Reading Excel file: {os.path.basename(excel_path)} (Sheet: {mpd_sheet})\n")
                        else:
                            df = pd.read_excel(excel_path, sheet_name=sheet_names[0], engine='pyxlsb')
                            _log(f"📊 Reading Excel file: {os.path.basename(excel_path)} (Sheet: {sheet_names[0]})\n")
                    else:
                        df = pd.read_excel(excel_path, sheet_name=0)
                        _log(f"📊 Reading Excel file: {os.path.basename(excel_path)} (Sheet: 0)\n")
                except Exception as e:
                    _log(f"❌ Failed to read Excel file: {e}\n")
            
            if df is None:
                _log(f"❌ Could not read any sheet from Excel file\n")
                return []
            
            # Convert all data to string for easier searching
            df_str = df.astype(str)
            
            # Find CMPISS03 section
            mpd_rsd_found = False
            mpd_rsd_start_row = -1
            
            for idx, row in df_str.iterrows():
                row_text = ' '.join(row.values)
                if re.search(r'\bCMPISS03\s+R1\b', row_text, re.IGNORECASE):
                    mpd_rsd_found = True
                    mpd_rsd_start_row = idx
                    _log(f"📍 Found CMPISS03 section at row {idx}\n")
                    break
            
            if not mpd_rsd_found:
                _log(f"⚠️ CMPISS03 R1 not found, searching for CMPISS03...\n")
                for idx, row in df_str.iterrows():
                    row_text = ' '.join(row.values)
                    if 'CMPISS03 R1' in row_text.upper():
                        mpd_rsd_found = True
                        mpd_rsd_start_row = idx
                        _log(f"📍 Found CMPISS03 R1 section at row {idx}\n")
                        break
            
            if not mpd_rsd_found:
                _log(f"❌ CMPISS03 section not found in Excel\n")
                return []
            
            # Find the requested check code in column 24
            target_check = cmp_iss03.upper()
            check_column = 24
            check_rows = []
            
            _log(f"🔍 Searching for '{target_check}' in column {check_column}...\n")
            
            for idx in range(len(df)):
                cell_value = str(df_str.iloc[idx, check_column]).strip()
                normalized_cell_value = re.sub(r'\s+', '', cell_value).upper()
                normalized_target = re.sub(r'\s+', '', target_check).upper()
                
                if normalized_cell_value == normalized_target:
                    check_rows.append(idx)
                    _log(f"📍 Found '{target_check}' at row {idx}, column {check_column}\n")
                    continue
                
                # Check with normalization
                if self._normalize_check_code(cell_value) == self._normalize_check_code(target_check):
                    check_rows.append(idx)
                    _log(f"📍 Found '{target_check}' (matched with '{cell_value}') at row {idx}\n")
            
            if not check_rows:
                _log(f"❌ No '{target_check}' found in column {check_column}\n")
                return []
            
            _log(f"✅ Found {len(check_rows)} occurrences of '{target_check}'\n")
            
            # Extract MPD ITEM tasks from column 0
            tasks = set()
            mpd_item_col = 0
            mpd_item_data = df_str.iloc[:, mpd_item_col].astype(str)
            
            for row_idx in check_rows:
                task_cell = mpd_item_data.iloc[row_idx]
                if str(task_cell).strip() and str(task_cell) != 'nan':
                    tasks.add(task_cell)
                    _log(f"     ✓ Found task: {task_cell} at row {row_idx}\n")
            
            sorted_tasks = sorted(tasks)
            _log(f"✅ Extracted {len(sorted_tasks)} tasks for '{cmp_iss03}': {sorted_tasks}\n")
            
            return sorted_tasks
        
        except Exception as e:
            _log(f"❌ Error extracting from Excel: {e}\n")
            return []
    
    def _refresh_available_checks(self):
        """
        Update available CHECK list from Excel file
        """
        if not hasattr(self, 'mpd_rsd_excel') or not self.mpd_rsd_excel:
            self._safe_show_error("Error", "Please select MPD RSD Excel file first")
            return
        
        if not os.path.exists(self.mpd_rsd_excel):
            self._safe_show_error("Error", "MPD RSD Excel file does not exist")
            return
        
        self._safe_log(self.tx_cmp_tcm_log, "🔄 Updating available CHECK list...\n")
        
        def job():
            try:
                sheet_name = None
                if hasattr(self, 'entry_sheet_name'):
                    sheet_name = self.entry_sheet_name.get().strip() or None
                
                available_checks = self._extract_available_checks_from_excel(self.mpd_rsd_excel, sheet_name)
                
                if available_checks:
                    # تحديث ComboBox في الخيط الرئيسي
                    def update_combo():
                        if hasattr(self, 'cmb_check2'):
                            current_value = self.cmb_check2.get()
                            self.cmb_check2.configure(values=available_checks)
                            # الاحتفاظ بالقيمة المحددة إذا كانت ما زالت متاحة
                            if current_value in available_checks:
                                self.cmb_check2.set(current_value)
                            elif available_checks:
                                self.cmb_check2.set(available_checks[0])
                        self._safe_log(self.tx_cmp_tcm_log, f"✅ CHECK list updated successfully ({len(available_checks)} items)\n")
                    
                    self.after(0, update_combo)
                else:
                    self.after(0, lambda: self._safe_log(self.tx_cmp_tcm_log, "⚠️ No valid CHECK values found in column 24\n"))
                    
            except Exception as e:
                self.after(0, lambda: self._safe_log(self.tx_cmp_tcm_log, f"❌ Error updating CHECK list: {e}\n"))
        
        threading.Thread(target=job, daemon=True).start()


    def _extract_available_checks_from_excel(self, excel_path: str, sheet_name: str = None):
        """
        Extract ALL values from column 24 in Excel file (no filtering)
        
        Args:
            excel_path: Path to Excel file
            sheet_name: Sheet name (optional)
            
        Returns:
            List of all unique values from column 24
        """
        if not excel_path or not os.path.exists(excel_path):
            self._safe_log(self.tx_cmp_tcm_log, "❌ Excel file path is invalid or file does not exist\n")
            return []
            
        try:
            # Read Excel file
            df = None
            
            if sheet_name and sheet_name.strip():
                try:
                    if excel_path.lower().endswith('.xlsb'):
                        try:
                            df = pd.read_excel(excel_path, sheet_name=sheet_name.strip(), engine='pyxlsb')
                        except ImportError:
                            df = pd.read_excel(excel_path, sheet_name=sheet_name.strip())
                    else:
                        df = pd.read_excel(excel_path, sheet_name=sheet_name.strip())
                    
                    self._safe_log(self.tx_cmp_tcm_log, f"📊 Reading Excel file: {os.path.basename(excel_path)} (Sheet: {sheet_name.strip()})\n")
                except Exception as e:
                    self._safe_log(self.tx_cmp_tcm_log, f"⚠️ Failed to read specified sheet '{sheet_name}': {e}\n")
                    df = None
            
            if df is None:
                # محاولة العثور على ورقة MPD RSD
                try:
                    if excel_path.lower().endswith('.xlsb'):
                        xl_file = pd.ExcelFile(excel_path, engine='pyxlsb')
                        sheet_names = xl_file.sheet_names
                        
                        mpd_sheet = None
                        for sheet in sheet_names:
                            if 'MPD RSD' in sheet.upper():
                                mpd_sheet = sheet
                                break
                        
                        if mpd_sheet:
                            df = pd.read_excel(excel_path, sheet_name=mpd_sheet, engine='pyxlsb')
                            self._safe_log(self.tx_cmp_tcm_log, f"📊 Reading Excel file: {os.path.basename(excel_path)} (Sheet: {mpd_sheet})\n")
                        else:
                            df = pd.read_excel(excel_path, sheet_name=sheet_names[0], engine='pyxlsb')
                            self._safe_log(self.tx_cmp_tcm_log, f"📊 Reading Excel file: {os.path.basename(excel_path)} (Sheet: {sheet_names[0]})\n")
                    else:
                        df = pd.read_excel(excel_path, sheet_name=0)
                        self._safe_log(self.tx_cmp_tcm_log, f"📊 Reading Excel file: {os.path.basename(excel_path)} (Sheet: 0)\n")
                except Exception as e:
                    self._safe_log(self.tx_cmp_tcm_log, f"❌ Failed to read Excel file: {e}\n")
                    return []
            
            if df is None:
                self._safe_log(self.tx_cmp_tcm_log, f"❌ Could not read any sheet from Excel file\n")
                return []
            
            # التحقق من وجود العمود 24
            if df.shape[1] <= 24:
                self._safe_log(self.tx_cmp_tcm_log, f"❌ ملف Excel لا يحتوي على العمود 24 (العدد الحالي للأعمدة: {df.shape[1]})\n")
                return []
            
            # استخراج كل القيم من العمود 24 بدون فلترة
            check_column = df.iloc[:, 24].astype(str)
            
            # تنظيف البيانات وإزالة القيم الفارغة أو nan فقط
            unique_checks = set()
            for value in check_column:
                cleaned_value = str(value).strip()
                # إزالة فقط القيم الفارغة و nan - باقي كل حاجة نضيفها
                if cleaned_value and cleaned_value.lower() not in ['nan', 'none', '']:
                    # الاحتفاظ بالقيمة كما هي بدون تغيير كبير
                    unique_checks.add(cleaned_value)
                    self._safe_log(self.tx_cmp_tcm_log, f"     ✓ Found: '{cleaned_value}'\n")
            
            sorted_checks = sorted(list(unique_checks))
            self._safe_log(self.tx_cmp_tcm_log, f"✅ Extracted {len(sorted_checks)} values from column 24\n")
            self._safe_log(self.tx_cmp_tcm_log, f"📋 All values: {', '.join(sorted_checks)}\n")
            
            return sorted_checks
            
        except Exception as e:
            self._safe_log(self.tx_cmp_tcm_log, f"❌ Error extracting values: {e}\n")
            return []

    def _generate_task_cards_indexed(self):
        """Generate task cards using indexed TCM extraction"""
        def job():
            self.operation_running = True
            
            selected = (self.cmb_check2.get() or "").strip()
            aircraft = (self.cmb_aircraft2.get() or "").strip()
            out_dir = self.out_cmp_tcm or (self.lbl_out_cmp_tcm.cget("text") if hasattr(self,"lbl_out_cmp_tcm") else None)
            
            # Check if we have MPD RSD Excel
            has_excel = hasattr(self, 'mpd_rsd_excel') and self.mpd_rsd_excel
            
            
            if not has_excel:
                self._safe_show_error("CMP/TCM","Select MPD RSD Excel")
                self.operation_running = False
                return
            if not self.tcm_dir or not os.path.isdir(self.tcm_dir):
                self._safe_show_error("CMP/TCM","Select valid TCM folder")
                self.operation_running = False
                return
            if not selected:
                self._safe_show_error("CMP/TCM","Select a Check")
                self.operation_running = False
                return
            if not aircraft:
                self._safe_show_error("CMP/TCM","Select an Aircraft")
                self.operation_running = False
                return
            if not out_dir:
                self._safe_show_error("CMP/TCM","Choose an Output folder")
                self.operation_running = False
                return

            if self.indexer is None:
                self.indexer = TcmIndexer(self.tcm_dir, threads=8, cache=True)
            if not self.indexer.index:
                loaded = self.indexer.try_load_cache()
                if not loaded:
                    self._safe_log(self.tx_cmp_tcm_log, "Index not found. Building index first...\n")
                    self.indexer.build_index(progress_callback=lambda s: self._safe_log(self.tx_cmp_tcm_log, s))
                    if self.indexer.stop_requested:
                        self._safe_log(self.tx_cmp_tcm_log, "❌ Operation cancelled by user.\n")
                        self.operation_running = False
                        return

            associated = expand_check(selected)
            self._safe_log(self.tx_cmp_tcm_log, f"[{time.strftime('%H:%M:%S')}] Generating for {selected} -> associated: {associated}\n")

            try:
                tasks_per_check = {}
                for chk in associated:
                    if self.indexer and self.indexer.stop_requested:
                        self._safe_log(self.tx_cmp_tcm_log, "❌ Operation cancelled by user.\n")
                        self.operation_running = False
                        return
                    
                    # Use Excel MPD RSD extraction
                    sheet_name = self.entry_sheet_name.get().strip() if hasattr(self, 'entry_sheet_name') else None
                    self._safe_log(self.tx_cmp_tcm_log, f"📊 Using Excel MPD RSD for {chk}\n")
                    base_tasks = self._extract_tasks_from_excel_mpd_rsd(self.mpd_rsd_excel, chk, sheet_name)
                    
                    # Expand base tasks with all related subtasks using the new expansion function
                    self._safe_log(self.tx_cmp_tcm_log, f"\n📋 Base tasks from Excel: {base_tasks}\n")
                    expanded_tasks = self._expand_tasks_with_subtasks(base_tasks, self.indexer)
                    tasks_per_check[chk] = expanded_tasks
                    self._safe_log(self.tx_cmp_tcm_log, f"✅ {chk}: {len(base_tasks)} base -> {len(expanded_tasks)} total\n")
                    self._safe_log(self.tx_cmp_tcm_log, f"📦 Final task list: {expanded_tasks}\n\n")

                total_saved = 0
                for chk, tasks in tasks_per_check.items():
                    if self.indexer and self.indexer.stop_requested:
                        self._safe_log(self.tx_cmp_tcm_log, "❌ Operation cancelled by user.\n")
                        self.operation_running = False
                        return
                    chk_folder = os.path.join(out_dir, chk.replace("/","_"))
                    safe_make_dir(chk_folder)
                    self._safe_log(self.tx_cmp_tcm_log, f"Processing check {chk}, tasks: {len(tasks)}\n")
                    for task_code in tasks:
                        if self.indexer and self.indexer.stop_requested:
                            self._safe_log(self.tx_cmp_tcm_log, "❌ Operation cancelled by user.\n")
                            self.operation_running = False
                            return
                        pdf_path, run = self.indexer.find_best_occurrence_for_task(task_code)
                        if not pdf_path or not run:
                            self._safe_log(self.tx_cmp_tcm_log, f"  - {task_code}: not found in index\n")
                            continue

                        start, end = run[0], run[1] if isinstance(run, (list,tuple)) and len(run)>=2 else (run[0], run[-1])
                        pages = list(range(start, end+1))

                        # Merge cover (if available)
                        cover_path = self._find_cover_for_task(task_code)
                        try:
                            src = fitz.open(pdf_path)
                        except Exception as e:
                            self._safe_log(self.tx_cmp_tcm_log, f"  - {task_code}: failed open {os.path.basename(pdf_path)} [{e}]\n")
                            continue
                        try:
                            out_doc = fitz.open()
                            if cover_path:
                                try:
                                    cover_doc = fitz.open(cover_path)
                                    out_doc.insert_pdf(cover_doc)
                                    cover_doc.close()
                                    self._safe_log(self.tx_cmp_tcm_log, f"    • merged cover: {os.path.basename(cover_path)}\n")
                                except Exception as e:
                                    self._safe_log(self.tx_cmp_tcm_log, f"    • cover merge failed [{e}] -> will save task without cover\n")
                            out_doc.insert_pdf(src, from_page=pages[0], to_page=pages[-1])
                            out_name = f"{task_code}.pdf"
                            out_path = unique_path(os.path.join(chk_folder, out_name))
                            out_doc.save(out_path)
                            out_doc.close()
                            total_saved += 1
                            self._safe_log(self.tx_cmp_tcm_log, f"  ✓ {task_code}: saved {len(pages)} page(s) -> {out_path}\n")
                        except Exception as e:
                            self._safe_log(self.tx_cmp_tcm_log, f"  - {task_code}: save failed [{e}]\n")
                        finally:
                            try: src.close()
                            except: pass

                self._safe_log(self.tx_cmp_tcm_log, f"Done. Total saved files: {total_saved}\n")
                self._safe_show_info("Task Completed", f"Successfully extracted {total_saved} files.\nLocation: {out_dir}")
            except Exception as e:
                self._safe_show_error("Generate failed", str(e))
            finally:
                self.operation_running = False

        threading.Thread(target=job, daemon=True).start()

    def _find_cover_for_task(self, task_code: str):
        """Search in covers_dir for a matching cover PDF for the task code."""
        if not self.covers_dir or not os.path.isdir(self.covers_dir):
            return None
        candidates = [
            f"{task_code}.pdf",
            f"{task_code.replace('-','_')}.pdf",
            f"{task_code.upper()}.pdf",
            f"{task_code.lower()}.pdf",
        ]
        for name in candidates:
            p = os.path.join(self.covers_dir, name)
            if is_pdf(p):
                return p
        return None

    def _tab_mail_merge(self):
        """Mail Merge (Covering) Module - Automated Word-Excel Integration for RC Cards"""
        
        # Check if required libraries are available
        if not DOCX_AVAILABLE:
            # Show error message
            error_frame = ctk.CTkFrame(self.content, corner_radius=16, fg_color=COLOR_CARD_BG)
            error_frame.grid(row=0, column=0, sticky="nsew", padx=12, pady=12)
            error_frame.grid_columnconfigure(0, weight=1)
            error_frame.grid_rowconfigure(0, weight=1)
            
            error_container = ctk.CTkFrame(error_frame, fg_color="transparent")
            error_container.grid(row=0, column=0, padx=40, pady=40)
            
            ctk.CTkLabel(
                error_container,
                text="📋 Mail Merge Module - Setup Required",
                font=("Segoe UI", 24, "bold"),
                text_color=COLOR_DANGER
            ).pack(pady=(0, 20))
            
            ctk.CTkLabel(
                error_container,
                text="The Mail Merge (Covering) module requires additional Python packages.",
                font=("Segoe UI", 14),
                text_color=COLOR_TEXT_DARK
            ).pack(pady=10)
            
            ctk.CTkLabel(
                error_container,
                text="Missing Package: python-docx",
                font=("Segoe UI", 12, "bold"),
                text_color=COLOR_DANGER
            ).pack(pady=10)
            
            install_frame = ctk.CTkFrame(error_container, fg_color=COLOR_BG_LIGHT, corner_radius=12)
            install_frame.pack(pady=20, padx=20, fill="x")
            
            ctk.CTkLabel(
                install_frame,
                text="Installation Instructions:",
                font=("Segoe UI", 14, "bold"),
                text_color=COLOR_PRIMARY
            ).pack(pady=(15, 10))
            
            instructions = [
                "1. Open Command Prompt or PowerShell",
                "2. Navigate to your project folder",
                "3. Run: pip install python-docx",
                "4. Restart the application"
            ]
            
            for instruction in instructions:
                ctk.CTkLabel(
                    install_frame,
                    text=instruction,
                    font=("Segoe UI", 11),
                    text_color=COLOR_TEXT_DARK,
                    anchor="w"
                ).pack(pady=5, padx=20, anchor="w")
            
            # Command to copy
            command_frame = ctk.CTkFrame(install_frame, fg_color="white", corner_radius=8)
            command_frame.pack(pady=15, padx=20, fill="x")
            
            command_label = ctk.CTkLabel(
                command_frame,
                text="pip install python-docx",
                font=("Consolas", 12, "bold"),
                text_color=COLOR_PRIMARY
            )
            command_label.pack(pady=10, padx=10)
            
            def copy_command():
                self.clipboard_clear()
                self.clipboard_append("pip install python-docx")
                messagebox.showinfo("Copied", "Installation command copied to clipboard!")
            
            ctk.CTkButton(
                install_frame,
                text="📋 Copy Installation Command",
                command=copy_command,
                fg_color=COLOR_OK,
                hover_color="#2f855a",
                height=40,
                font=("Segoe UI", 12, "bold")
            ).pack(pady=(0, 15))
            
            return
        
        # Original Mail Merge implementation
        # Main card for file selection and settings
        card = ctk.CTkFrame(self.content, corner_radius=16, fg_color=COLOR_CARD_BG)
        card.grid(row=0, column=0, sticky="ew", padx=12, pady=12)
        card.grid_columnconfigure((1, 2), weight=1)
        
        # Title
        ctk.CTkLabel(
            card, 
            text="Mail Merge (Covering) - RC Card Generator", 
            font=("Segoe UI", 22, "bold"), 
            text_color=COLOR_PRIMARY
        ).grid(row=0, column=0, columnspan=3, padx=20, pady=(20, 12), sticky="w")
        
        # Word Template selection
        ctk.CTkButton(
            card, 
            text="📄 Word Template…", 
            fg_color=COLOR_OK, 
            hover_color="#2f855a", 
            height=40, 
            font=("Segoe UI", 12, "bold"), 
            corner_radius=10,
            command=self._mm_select_word_template
        ).grid(row=1, column=0, padx=20, pady=10, sticky="w")
        
        self.mm_word_path_label = ctk.CTkLabel(
            card, 
            text="No template selected", 
            font=("Segoe UI", 11), 
            text_color=COLOR_TEXT_DARK
        )
        self.mm_word_path_label.grid(row=1, column=1, columnspan=2, padx=12, pady=10, sticky="w")
        
        # Excel Data selection
        ctk.CTkButton(
            card, 
            text="📊 Excel Data…", 
            fg_color=COLOR_OK, 
            hover_color="#2f855a", 
            height=40, 
            font=("Segoe UI", 12, "bold"), 
            corner_radius=10,
            command=self._mm_select_excel_file
        ).grid(row=2, column=0, padx=20, pady=10, sticky="w")
        
        self.mm_excel_path_label = ctk.CTkLabel(
            card, 
            text="No data file selected", 
            font=("Segoe UI", 11), 
            text_color=COLOR_TEXT_DARK
        )
        self.mm_excel_path_label.grid(row=2, column=1, columnspan=2, padx=12, pady=10, sticky="w")
        
        # Field Mapping Button
        ctk.CTkButton(
            card, 
            text="🔗 Configure Field Mapping", 
            fg_color=COLOR_PRIMARY, 
            hover_color=COLOR_ACCENT_DARK, 
            height=40, 
            font=("Segoe UI", 12, "bold"), 
            corner_radius=10,
            command=self._mm_show_field_mapping
        ).grid(row=3, column=0, padx=20, pady=10, sticky="w")
        
        self.mm_field_status_label = ctk.CTkLabel(
            card, 
            text="Auto-mapping enabled", 
            font=("Segoe UI", 11), 
            text_color=COLOR_OK
        )
        self.mm_field_status_label.grid(row=3, column=1, columnspan=2, padx=12, pady=10, sticky="w")
        
        # MPD/Key Input section
        mpd_frame = ctk.CTkFrame(card, fg_color=COLOR_BG_LIGHT, corner_radius=12)
        mpd_frame.grid(row=4, column=0, columnspan=3, padx=20, pady=(12, 16), sticky="ew")
        mpd_frame.grid_columnconfigure(1, weight=1)
        
        ctk.CTkLabel(
            mpd_frame, 
            text="Unique Key (MPD/SEQ):", 
            font=("Segoe UI", 12, "bold"), 
            text_color=COLOR_PRIMARY
        ).grid(row=0, column=0, padx=12, pady=12, sticky="w")
        
        self.mm_mpd_entry = ctk.CTkEntry(
            mpd_frame, 
            placeholder_text="Enter MPD, SEQ, or unique identifier", 
            height=40, 
            font=("Segoe UI", 12), 
            corner_radius=8
        )
        self.mm_mpd_entry.grid(row=0, column=1, padx=12, pady=12, sticky="ew")
        
        # Action buttons
        action_frame = ctk.CTkFrame(card, fg_color="transparent")
        action_frame.grid(row=5, column=0, columnspan=3, padx=20, pady=(12, 20), sticky="ew")
        action_frame.grid_columnconfigure((0, 1, 2), weight=1)
        
        ctk.CTkButton(
            action_frame, 
            text="🔍 Preview Data", 
            height=50, 
            fg_color=COLOR_PRIMARY, 
            hover_color=COLOR_HOVER, 
            font=("Segoe UI", 13, "bold"), 
            corner_radius=12,
            command=self._mm_preview_data
        ).grid(row=0, column=0, padx=10, pady=10, sticky="ew")
        
        ctk.CTkButton(
            action_frame, 
            text="🚀 Generate Document", 
            height=50, 
            fg_color=COLOR_WARN, 
            hover_color="#dd6b20", 
            font=("Segoe UI", 13, "bold"), 
            corner_radius=12,
            command=self._mm_generate_document
        ).grid(row=0, column=1, padx=10, pady=10, sticky="ew")
        
     #   ctk.CTkButton(
      #      action_frame, 
       #     text="📦 Generate All Rows", 
        #    height=50, 
         #   fg_color=COLOR_OK, 
          #  hover_color="#2f855a", 
           # font=("Segoe UI", 13, "bold"), 
            #corner_radius=12,
            #command=self._mm_batch_generate_all
       # ).grid(row=0, column=2, padx=10, pady=10, sticky="ew")
        
        # Output/Log section
        log_frame = ctk.CTkFrame(self.content, corner_radius=16, fg_color=COLOR_CARD_BG)
        log_frame.grid(row=1, column=0, sticky="nsew", padx=12, pady=12)
        log_frame.grid_rowconfigure(1, weight=1)
        log_frame.grid_columnconfigure(0, weight=1)
        
        ctk.CTkLabel(
            log_frame, 
            text="Process Log & Results", 
            font=("Segoe UI", 16, "bold"), 
            text_color=COLOR_PRIMARY
        ).grid(row=0, column=0, padx=15, pady=(15, 5), sticky="w")
        
        # Log text widget with scrollbar
        self.mm_log_text = tk.Text(
            log_frame, 
            height=15, 
            font=("Consolas", 10), 
            bg=COLOR_BG_LIGHT, 
            fg=COLOR_TEXT_DARK, 
            relief="flat", 
            bd=0,
            wrap=tk.WORD
        )
        self.mm_log_text.grid(row=1, column=0, sticky="nsew", padx=15, pady=15)
        
        scrollbar = ttk.Scrollbar(log_frame, command=self.mm_log_text.yview)
        scrollbar.grid(row=1, column=1, sticky="ns", pady=15)
        self.mm_log_text.configure(yscrollcommand=scrollbar.set)
        
        # Initialize variables
        self.mm_word_path = None
        self.mm_excel_path = None
        self.mm_excel_data = None
        self.mm_field_mapping = {}  # Manual field mapping: {word_field: excel_column}
        self.mm_detected_fields = []  # Fields detected in Word template
        self.mm_auto_mapping = True  # Use automatic field mapping by default
        self.mm_unique_key_column = 'MPD'  # Default unique key column
        
        # Welcome message
        self._mm_log("✨ Mail Merge (Covering) Module Initialized\n")
        self._mm_log("📌 Instructions:\n")
        self._mm_log("   1. Select Word Template (RC Card) file (.docx)\n")
        self._mm_log("   2. Select Excel Data file (.xlsx) with task data\n")
        self._mm_log("   3. Configure field mapping (auto or manual)\n")
        self._mm_log("   4. Enter unique key (MPD/SEQ) or generate all rows\n")
        self._mm_log("   5. Preview data and generate documents\n\n")

    def _mm_log(self, message):
        """Helper to log messages to the Mail Merge log widget"""
        try:
            self.mm_log_text.insert(tk.END, message)
            self.mm_log_text.see(tk.END)
            self.mm_log_text.update()
        except Exception:
            pass

    def _mm_load_excel_ignoring_names(self, file_path):
        with open(file_path, "rb") as source:
            original_bytes = source.read()
        cleaned_buffer = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(original_bytes)) as zin, zipfile.ZipFile(cleaned_buffer, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "xl/workbook.xml":
                    try:
                        root = ET.fromstring(data)
                        ns = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
                        defined_names = root.find("main:definedNames", ns)
                        if defined_names is not None:
                            root.remove(defined_names)
                            data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                    except Exception:
                        pass
                zout.writestr(item, data)
        cleaned_buffer.seek(0)
        workbook = openpyxl.load_workbook(cleaned_buffer, data_only=True, read_only=True, keep_links=False)
        sheet_name = "RC" if "RC" in workbook.sheetnames else workbook.sheetnames[0]
        worksheet = workbook[sheet_name]
        data_rows = list(worksheet.iter_rows(values_only=True))
        workbook.close()
        if not data_rows:
            return pd.DataFrame()
        headers = []
        used = set()
        for idx, value in enumerate(data_rows[0]):
            candidate = "" if value is None else str(value).strip()
            if not candidate:
                candidate = f"Column{idx+1}"
            while candidate in used:
                candidate = f"{candidate}_{idx+1}"
            headers.append(candidate)
            used.add(candidate)
        rows = [list(row) for row in data_rows[1:]]
        frame = pd.DataFrame(rows, columns=headers)
        frame = frame.dropna(how="all")
        return frame

    def _mm_select_word_template(self):
        """Select Word template file for Mail Merge"""
        try:
            file_path = filedialog.askopenfilename(
                title="Select Word Template (RC Card)",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
            )
            
            if file_path and os.path.exists(file_path):
                self.mm_word_path = file_path
                filename = os.path.basename(file_path)
                self.mm_word_path_label.configure(text=filename)
                self._mm_log(f"✅ Word Template Selected: {filename}\n")
            else:
                self._mm_log("⚠️ Word template selection cancelled\n")
                
        except Exception as e:
            self._mm_log(f"❌ Error selecting Word template: {e}\n")
            messagebox.showerror("Error", f"Failed to select Word template:\n{e}")

    def _mm_select_excel_file(self):
        """Select Excel data file for Mail Merge"""
        try:
            file_path = filedialog.askopenfilename(
                title="Select Excel or CSV Data File",
                filetypes=[
                    ("Excel Files", "*.xlsx;*.xls"),
                    ("CSV Files", "*.csv"),
                    ("All Files", "*.*")
                ]
            )
            
            if file_path and os.path.exists(file_path):
                self.mm_excel_path = file_path
                filename = os.path.basename(file_path)
                self.mm_excel_path_label.configure(text=filename)
                self._mm_log(f"✅ Excel File Selected: {filename}\n")
                
                # Load Excel data
                try:
                    # Try multiple engines to handle different Excel formats
                    excel_loaded = False
                    error_messages = []
                    
                    # Check if it's a CSV file
                    if file_path.lower().endswith('.csv'):
                        try:
                            self.mm_excel_data = pd.read_csv(file_path)
                            excel_loaded = True
                            self._mm_log("📄 CSV file detected and loaded\n")
                        except Exception as e_csv:
                            error_messages.append(f"CSV: {str(e_csv)[:100]}")
                    else:
                        # Try openpyxl first (for .xlsx files)
                        try:
                            self.mm_excel_data = pd.read_excel(file_path, engine='openpyxl')
                            excel_loaded = True
                        except Exception as e1:
                            error_messages.append(f"openpyxl: {str(e1)[:100]}")
                            
                            # Try xlrd for older .xls files
                            try:
                                self.mm_excel_data = pd.read_excel(file_path, engine='xlrd')
                                excel_loaded = True
                            except Exception as e2:
                                error_messages.append(f"xlrd: {str(e2)[:100]}")
                    
                    if not excel_loaded and file_path.lower().endswith(('.xlsx', '.xlsm', '.xltx', '.xltm')):
                        try:
                            self.mm_excel_data = self._mm_load_excel_ignoring_names(file_path)
                            excel_loaded = True
                            self._mm_log("🛠️ Excel fallback loader used\n")
                        except Exception as fallback_error:
                            error_messages.append(f"fallback: {str(fallback_error)[:100]}")
                    if excel_loaded:
                        rows, cols = self.mm_excel_data.shape
                        self._mm_log(f"✅ Excel data loaded successfully!\n")
                        self._mm_log(f"📊 Rows: {rows}, Columns: {cols}\n")
                        self._mm_log(f"📋 Columns: {', '.join(self.mm_excel_data.columns.tolist())}\n\n")
                        
                        # Check if MPD column exists
                        if 'MPD' not in self.mm_excel_data.columns:
                            self._mm_log("⚠️ Warning: 'MPD' column not found in Excel file\n")
                            self._mm_log(f"Available columns: {', '.join(self.mm_excel_data.columns.tolist())}\n\n")
                            messagebox.showwarning(
                                "Column Missing", 
                                "The Excel file doesn't contain an 'MPD' column.\nPlease ensure your data has an 'MPD' column."
                            )
                    else:
                        # Excel loading failed with all engines
                        self._mm_log("❌ Failed to load Excel file\n")
                        self._mm_log("The file may be corrupted or contain invalid XML/data.\n\n")
                        self._mm_log("💡 Suggested solutions:\n")
                        self._mm_log("   1. Open Excel → Formulas → Name Manager\n")
                        self._mm_log("      Delete any names with errors (#REF!, #NAME?)\n")
                        self._mm_log("   2. Save as new Excel file (.xlsx)\n")
                        self._mm_log("   3. Export to CSV and use that instead\n")
                        self._mm_log("   4. Remove special formatting/charts and save\n\n")
                        
                        messagebox.showerror(
                            "Excel File Error", 
                            "❌ Unable to read the Excel file.\n\n"
                            "The file contains corrupted data or invalid XML.\n\n"
                            "🔧 QUICK FIX:\n"
                            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                            "1. Open file in Excel\n"
                            "2. File → Save As → CSV (Comma delimited)\n"
                            "3. Use the CSV file instead\n\n"
                            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                            "OR:\n"
                            "1. Formulas → Name Manager\n"
                            "2. Delete names with errors (#REF!, #NAME?)\n"
                            "3. Save as new .xlsx file\n\n"
                            f"Technical error:\n{error_messages[0] if error_messages else 'Unknown error'}"
                        )
                    
                except Exception as e:
                    self._mm_log(f"❌ Error loading Excel data: {e}\n")
                    self._mm_log("\n💡 Try opening the file in Excel and save it as a new file.\n\n")
                    messagebox.showerror("Error", f"Failed to load Excel data:\n{e}")
            else:
                self._mm_log("⚠️ Excel file selection cancelled\n")
                
        except Exception as e:
            self._mm_log(f"❌ Error selecting Excel file: {e}\n")
            messagebox.showerror("Error", f"Failed to select Excel file:\n{e}")

    def _mm_preview_data(self):
        """Preview the data for the selected MPD"""
        try:
            # Validation
            if self.mm_excel_data is None:
                self._mm_log("❌ Please select an Excel file first\n")
                messagebox.showwarning("Missing File", "Please select an Excel data file first.")
                return
            
            mpd_value = self.mm_mpd_entry.get().strip()
            if not mpd_value:
                self._mm_log("❌ Please enter an MPD number\n")
                messagebox.showwarning("Missing MPD", "Please enter an MPD number to preview.")
                return
            
            # Search for MPD in data
            self._mm_log(f"\n🔍 Searching for MPD: {mpd_value}...\n")
            
            if 'MPD' not in self.mm_excel_data.columns:
                self._mm_log("❌ 'MPD' column not found in Excel data\n")
                messagebox.showerror("Column Missing", "The Excel file doesn't contain an 'MPD' column.")
                return
            
            # Find matching row
            matching_rows = self.mm_excel_data[self.mm_excel_data['MPD'].astype(str) == mpd_value]
            
            if matching_rows.empty:
                self._mm_log(f"❌ MPD '{mpd_value}' not found in Excel data\n")
                messagebox.showinfo("Not Found", f"MPD '{mpd_value}' was not found in the Excel data.")
                return
            
            # Display data
            row_data = matching_rows.iloc[0]
            self._mm_log(f"✅ MPD '{mpd_value}' found! Data preview:\n")
            self._mm_log("=" * 60 + "\n")
            
            for column, value in row_data.items():
                self._mm_log(f"  {column}: {value}\n")
            
            self._mm_log("=" * 60 + "\n\n")
            
        except Exception as e:
            self._mm_log(f"❌ Error previewing data: {e}\n")
            messagebox.showerror("Error", f"Failed to preview data:\n{e}")

    def _mm_manual_replace(self, doc, context):
        """
        Manually replace «KEY» patterns in the document using python-docx.
        This is a fallback for fields that are not true Merge Fields (plain text).
        """
        def replace_text(text):
            if '«' not in text:
                return text
            for key, value in context.items():
                if value is None: value = ""
                # Try different patterns
                patterns = [f"«{key}»", f"« {key} »", f"«{key} »", f"« {key}»"]
                for p in patterns:
                    if p in text:
                        text = text.replace(p, str(value))
            return text

        def process_paragraph(paragraph):
            if '«' in paragraph.text:
                # Simple replacement strategy
                new_text = replace_text(paragraph.text)
                if new_text != paragraph.text:
                    paragraph.text = new_text

        # Process paragraphs in body
        for p in doc.paragraphs:
            process_paragraph(p)
            
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_paragraph(p)

    def _mm_generate_document(self):
        """Generate Word document with Mail Merge data (Enhanced with docx-mailmerge)"""
        try:
            # Validation
            if not self.mm_word_path:
                self._mm_log("❌ Please select a Word template first\n")
                messagebox.showwarning("Missing Template", "Please select a Word template file first.")
                return
            
            if self.mm_excel_data is None:
                self._mm_log("❌ Please select an Excel file first\n")
                messagebox.showwarning("Missing Data", "Please select an Excel data file first.")
                return
            
            mpd_value = self.mm_mpd_entry.get().strip()
            if not mpd_value:
                self._mm_log("❌ Please enter an MPD number\n")
                messagebox.showwarning("Missing MPD", "Please enter an MPD number to generate document.")
                return
            
            # Search for MPD in data
            self._mm_log(f"\n🚀 Generating document for MPD: {mpd_value}...\n")
            
            if 'MPD' not in self.mm_excel_data.columns:
                self._mm_log("❌ 'MPD' column not found in Excel data\n")
                messagebox.showerror("Column Missing", "The Excel file doesn't contain an 'MPD' column.")
                return
            
            # Find matching row
            # Ensure MPD column is string and stripped for comparison
            matching_rows = self.mm_excel_data[self.mm_excel_data['MPD'].astype(str).str.strip() == mpd_value]
            
            if matching_rows.empty:
                self._mm_log(f"❌ MPD '{mpd_value}' not found in Excel data\n")
                messagebox.showinfo("Not Found", f"MPD '{mpd_value}' was not found in the Excel data.")
                return
            
            row_data = matching_rows.iloc[0].to_dict()
            
            # Prepare context with normalization (A/C -> AC)
            context = {}
            for col, value in row_data.items():
                # Original value
                val_str = str(value).strip() if pd.notnull(value) and value != '' else ''
                
                # Force LTR for codes (containing hyphens and digits)
                # This fixes the "reversed number" issue (e.g. 53-844-00 showing as 00-844-53)
                if re.search(r'\d+-\d+', val_str):
                    # Wrap in LTR embedding characters: \u202A (LRE) ... \u202C (PDF)
                    val_str = f"\u202A{val_str}\u202C"
                
                context[col] = val_str
                
                # Normalized key (remove special chars)
                normalized = re.sub(r'[^A-Z0-9_]', '', str(col).upper())
                if normalized != str(col).upper():
                    context[normalized] = val_str
            
            # Use docx-mailmerge if available
            if MAILMERGE_AVAILABLE:
                self._mm_log("📝 Loading Word template (using docx-mailmerge)...\n")
                doc = MailMerge(self.mm_word_path)
                
                # Merge data
                doc.merge(**context)
                
                # Ask for save location
                output_path = filedialog.asksaveasfilename(
                    title="Save Generated Document",
                    defaultextension=".docx",
                    initialfile=f"RC_Card_{mpd_value.replace('/', '_')}.docx",
                    filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
                )
                
                if output_path:
                    # Hybrid Approach:
                    # 1. Save MailMerge output to temp file
                    temp_path = os.path.join(os.path.dirname(output_path), f"temp_{int(time.time())}.docx")
                    doc.write(temp_path)
                    
                    # 2. Open with python-docx for manual replacement of remaining text placeholders
                    try:
                        doc_final = Document(temp_path)
                        self._mm_manual_replace(doc_final, context)
                        doc_final.save(output_path)
                        self._mm_log(f"✅ Document generated successfully (Hybrid Mode)!\n")
                    except Exception as e:
                        self._mm_log(f"⚠️ Manual replacement failed, using standard output: {e}\n")
                        # Fallback: just rename temp to output
                        if os.path.exists(output_path): os.remove(output_path)
                        os.rename(temp_path, output_path)
                    
                    # Cleanup temp
                    if os.path.exists(temp_path):
                        try: os.remove(temp_path)
                        except: pass

                    self._mm_log(f"📁 Saved to: {output_path}\n\n")
                    
                    messagebox.showinfo(
                        "Success", 
                        f"RC Card generated successfully!\n\nSaved to:\n{output_path}"
                    )
                else:
                    self._mm_log("⚠️ Document generation cancelled\n")
            else:
                self._mm_log("❌ docx-mailmerge library not found.\n")
                messagebox.showerror("Missing Library", "Please install docx-mailmerge: pip install docx-mailmerge")
            
        except Exception as e:
            self._mm_log(f"❌ Error generating document: {e}\n")
            messagebox.showerror("Error", f"Failed to generate document:\n{e}")

    def _mm_detect_merge_fields(self, doc_path):
        """Detect all merge field placeholders in the Word document"""
        try:
            doc = Document(doc_path)
            detected_fields = set()
            
            # Common merge field patterns
            patterns = [
                re.compile(r'«([^»]+)»'),  # «FIELD»
                re.compile(r'<<([^>]+)>>'),  # <<FIELD>>
                re.compile(r'\{\{([^}]+)\}\}'),  # {{FIELD}}
                re.compile(r'\{merge\s+([^}]+)\}'),  # {merge FIELD}
            ]
            
            # Search in paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text
                for pattern in patterns:
                    matches = pattern.findall(text)
                    detected_fields.update(matches)
            
            # Search in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            for pattern in patterns:
                                matches = pattern.findall(text)
                                detected_fields.update(matches)
            
            return sorted(list(detected_fields))
        except Exception as e:
            self._mm_log(f"❌ Error detecting merge fields: {e}\n")
            return []

    def _mm_show_field_mapping(self):
        """Show field mapping configuration dialog"""
        try:
            # Validation
            if not self.mm_word_path:
                messagebox.showwarning("Missing Template", "Please select a Word template first.")
                return
            
            if self.mm_excel_data is None:
                messagebox.showwarning("Missing Data", "Please select an Excel file first.")
                return
            
            # Detect fields in template
            self._mm_log("\n🔍 Detecting merge fields in template...\n")
            self.mm_detected_fields = self._mm_detect_merge_fields(self.mm_word_path)
            
            if not self.mm_detected_fields:
                messagebox.showinfo(
                    "No Fields Detected",
                    "No merge field placeholders detected in the template.\n\n"
                    "Expected formats:\n"
                    "• «FIELDNAME»\n"
                    "• <<FIELDNAME>>\n"
                    "• {{FIELDNAME}}\n"
                    "• {merge FIELDNAME}"
                )
                return
            
            self._mm_log(f"✅ Found {len(self.mm_detected_fields)} merge fields\n")
            
            # Create mapping dialog
            mapping_window = ctk.CTkToplevel(self)
            mapping_window.title("Field Mapping Configuration")
            mapping_window.geometry("900x700")
            mapping_window.transient(self)
            mapping_window.grab_set()
            
            # Title
            ctk.CTkLabel(
                mapping_window,
                text="🔗 Configure Field Mapping",
                font=("Segoe UI", 20, "bold"),
                text_color=COLOR_PRIMARY
            ).pack(pady=(20, 10))
            
            ctk.CTkLabel(
                mapping_window,
                text="Map Word merge fields to Excel columns",
                font=("Segoe UI", 12),
                text_color=COLOR_TEXT_DARK
            ).pack(pady=(0, 20))
            
            # Main frame with scrollbar
            main_frame = ctk.CTkFrame(mapping_window, fg_color="transparent")
            main_frame.pack(fill="both", expand=True, padx=20, pady=(0, 10))
            
            # Canvas for scrolling
            canvas = tk.Canvas(main_frame, bg=COLOR_BG_LIGHT, highlightthickness=0)
            scrollbar = ctk.CTkScrollbar(main_frame, command=canvas.yview)
            scrollable_frame = ctk.CTkFrame(canvas, fg_color=COLOR_BG_LIGHT)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )
            
            canvas_frame = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            def on_canvas_configure(event):
                canvas.itemconfig(canvas_frame, width=event.width)
            
            canvas.bind("<Configure>", on_canvas_configure)
            
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            # Excel columns
            excel_columns = ['<Auto-Match>'] + list(self.mm_excel_data.columns)
            
            # Field mapping widgets
            mapping_widgets = {}
            
            # Header
            header_frame = ctk.CTkFrame(scrollable_frame, fg_color=COLOR_PRIMARY, corner_radius=8)
            header_frame.pack(fill="x", padx=10, pady=(10, 5))
            header_frame.grid_columnconfigure((0, 1, 2), weight=1)
            
            ctk.CTkLabel(
                header_frame,
                text="Word Field",
                font=("Segoe UI", 12, "bold"),
                text_color="white"
            ).grid(row=0, column=0, padx=15, pady=10, sticky="w")
            
            ctk.CTkLabel(
                header_frame,
                text="Excel Column",
                font=("Segoe UI", 12, "bold"),
                text_color="white"
            ).grid(row=0, column=1, padx=15, pady=10, sticky="w")
            
            ctk.CTkLabel(
                header_frame,
                text="Preview",
                font=("Segoe UI", 12, "bold"),
                text_color="white"
            ).grid(row=0, column=2, padx=15, pady=10, sticky="w")
            
            # Create mapping row for each field
            for idx, field in enumerate(self.mm_detected_fields):
                row_frame = ctk.CTkFrame(scrollable_frame, fg_color=COLOR_CARD_BG, corner_radius=8)
                row_frame.pack(fill="x", padx=10, pady=3)
                row_frame.grid_columnconfigure(1, weight=1)
                row_frame.grid_columnconfigure(2, weight=1)
                
                # Word field label
                ctk.CTkLabel(
                    row_frame,
                    text=f"«{field}»",
                    font=("Segoe UI", 11, "bold"),
                    text_color=COLOR_PRIMARY,
                    anchor="w"
                ).grid(row=0, column=0, padx=15, pady=12, sticky="w")
                
                # Excel column dropdown
                # Try to find exact or similar match
                default_value = '<Auto-Match>'
                if field in excel_columns:
                    default_value = field
                elif field.upper() in [col.upper() for col in excel_columns]:
                    # Case-insensitive match
                    for col in excel_columns:
                        if col.upper() == field.upper():
                            default_value = col
                            break
                
                combo = ctk.CTkComboBox(
                    row_frame,
                    values=excel_columns,
                    font=("Segoe UI", 11),
                    width=200
                )
                combo.set(default_value)
                combo.grid(row=0, column=1, padx=15, pady=12, sticky="ew")
                
                # Preview label
                preview_label = ctk.CTkLabel(
                    row_frame,
                    text="",
                    font=("Segoe UI", 10),
                    text_color=COLOR_TEXT_DARK,
                    anchor="w"
                )
                preview_label.grid(row=0, column=2, padx=15, pady=12, sticky="w")
                
                # Update preview when selection changes
                def update_preview(choice, field=field, label=preview_label):
                    if choice and choice != '<Auto-Match>' and choice in self.mm_excel_data.columns:
                        sample = str(self.mm_excel_data[choice].iloc[0])[:30]
                        label.configure(text=f"e.g., {sample}...")
                    else:
                        label.configure(text="Auto-detected")
                
                combo.configure(command=update_preview)
                update_preview(default_value)
                
                mapping_widgets[field] = combo
            
            # Unique key selection
            key_frame = ctk.CTkFrame(scrollable_frame, fg_color=COLOR_CARD_BG, corner_radius=8)
            key_frame.pack(fill="x", padx=10, pady=(15, 5))
            key_frame.grid_columnconfigure(1, weight=1)
            
            ctk.CTkLabel(
                key_frame,
                text="🔑 Unique Key Column:",
                font=("Segoe UI", 12, "bold"),
                text_color=COLOR_PRIMARY
            ).grid(row=0, column=0, padx=15, pady=12, sticky="w")
            
            key_combo = ctk.CTkComboBox(
                key_frame,
                values=list(self.mm_excel_data.columns),
                font=("Segoe UI", 11),
                width=200
            )
            key_combo.set(self.mm_unique_key_column if self.mm_unique_key_column in self.mm_excel_data.columns else excel_columns[1] if len(excel_columns) > 1 else excel_columns[0])
            key_combo.grid(row=0, column=1, padx=15, pady=12, sticky="w")
            
            ctk.CTkLabel(
                key_frame,
                text="(Column used to identify each row uniquely)",
                font=("Segoe UI", 10),
                text_color=COLOR_TEXT_DARK
            ).grid(row=0, column=2, padx=15, pady=12, sticky="w")
            
            # Buttons
            button_frame = ctk.CTkFrame(mapping_window, fg_color="transparent")
            button_frame.pack(fill="x", padx=20, pady=(10, 20))
            button_frame.grid_columnconfigure((0, 1, 2), weight=1)
            
            def save_mapping():
                self.mm_field_mapping = {}
                self.mm_auto_mapping = False
                
                for field, widget in mapping_widgets.items():
                    selected = widget.get()
                    if selected and selected != '<Auto-Match>':
                        self.mm_field_mapping[field] = selected
                    else:
                        self.mm_auto_mapping = True  # At least one field uses auto-mapping
                
                self.mm_unique_key_column = key_combo.get()
                
                manual_count = len(self.mm_field_mapping)
                auto_count = len(self.mm_detected_fields) - manual_count
                
                self._mm_log(f"\n✅ Field mapping configured!\n")
                self._mm_log(f"   Manual mappings: {manual_count}\n")
                self._mm_log(f"   Auto-mapped: {auto_count}\n")
                self._mm_log(f"   Unique key: {self.mm_unique_key_column}\n\n")
                
                status_text = f"Manual: {manual_count}, Auto: {auto_count} fields"
                self.mm_field_status_label.configure(text=status_text, text_color=COLOR_OK)
                
                mapping_window.destroy()
            
            def reset_mapping():
                for widget in mapping_widgets.values():
                    widget.set('<Auto-Match>')
                self._mm_log("🔄 Mapping reset to auto-detect\n")
            
            ctk.CTkButton(
                button_frame,
                text="🔄 Reset to Auto",
                command=reset_mapping,
                fg_color=COLOR_ACCENT,
                hover_color=COLOR_ACCENT_DARK,
                height=40,
                font=("Segoe UI", 12, "bold")
            ).grid(row=0, column=0, padx=10, sticky="ew")
            
            ctk.CTkButton(
                button_frame,
                text="❌ Cancel",
                command=mapping_window.destroy,
                fg_color=COLOR_DANGER,
                hover_color="#c53030",
                height=40,
                font=("Segoe UI", 12, "bold")
            ).grid(row=0, column=1, padx=10, sticky="ew")
            
            ctk.CTkButton(
                button_frame,
                text="✅ Save Mapping",
                command=save_mapping,
                fg_color=COLOR_OK,
                hover_color="#2f855a",
                height=40,
                font=("Segoe UI", 12, "bold")
            ).grid(row=0, column=2, padx=10, sticky="ew")
            
        except Exception as e:
            self._mm_log(f"❌ Error configuring field mapping: {e}\n")
            messagebox.showerror("Error", f"Failed to configure field mapping:\n{e}")

    def _mm_batch_generate_all(self):
        """Generate documents for all rows in Excel (Enhanced with docx-mailmerge)"""
        try:
            # Validation
            if not self.mm_word_path:
                self._mm_log("❌ Please select a Word template first\n")
                messagebox.showwarning("Missing Template", "Please select a Word template file first.")
                return
            
            if self.mm_excel_data is None:
                self._mm_log("❌ Please select an Excel file first\n")
                messagebox.showwarning("Missing Data", "Please select an Excel data file first.")
                return
            
            if not MAILMERGE_AVAILABLE:
                messagebox.showerror("Missing Library", "Please install docx-mailmerge: pip install docx-mailmerge")
                return

            # Ask for output directory
            output_dir = filedialog.askdirectory(
                title="Select Output Directory for Generated Documents"
            )
            
            if not output_dir:
                self._mm_log("⚠️ Batch generation cancelled\n")
                return
            
            total_rows = len(self.mm_excel_data)
            
            # Confirm
            confirm = messagebox.askyesno(
                "Confirm Batch Generation",
                f"Generate {total_rows} documents?\n\n"
                f"Output directory:\n{output_dir}\n\n"
                f"This may take a few minutes."
            )
            
            if not confirm:
                self._mm_log("⚠️ Batch generation cancelled\n")
                return
            
            self._mm_log(f"\n📦 Starting Batch Generation (using docx-mailmerge)...\n")
            self._mm_log(f"📁 Output Directory: {output_dir}\n")
            self._mm_log(f"📊 Total Rows: {total_rows}\n")
            self._mm_log("=" * 60 + "\n")
            
            success_count = 0
            error_count = 0
            
            for idx, row in self.mm_excel_data.iterrows():
                try:
                    row_data = row.to_dict()
                    
                    # Determine unique filename (MPD preferred)
                    if 'MPD' in row_data and pd.notnull(row_data['MPD']):
                        unique_key = str(row_data['MPD']).strip()
                    elif self.mm_unique_key_column in row_data:
                        unique_key = str(row_data[self.mm_unique_key_column])
                    else:
                        unique_key = f"Row_{idx+1}"
                    
                    # Skip empty keys
                    if not unique_key or unique_key.lower() == 'nan':
                        continue

                    # Prepare context with normalization
                    context = {}
                    for col, value in row_data.items():
                        val_str = str(value).strip() if pd.notnull(value) and value != '' else ''
                        
                        # Force LTR for codes (containing hyphens and digits)
                        if re.search(r'\d+-\d+', val_str):
                            val_str = f"\u202A{val_str}\u202C"
                        
                        context[col] = val_str
                        
                        # Normalized key
                        normalized = re.sub(r'[^A-Z0-9_]', '', str(col).upper())
                        if normalized != str(col).upper():
                            context[normalized] = val_str
                    
                    # Generate
                    doc = MailMerge(self.mm_word_path)
                    doc.merge(**context)
                    
                    # Save document
                    safe_filename = re.sub(r'[^\w\-]', '_', unique_key)
                    output_path = os.path.join(output_dir, f"RC_Card_{safe_filename}.docx")
                    
                    # Handle duplicate filenames
                    counter = 1
                    base_path = output_path
                    while os.path.exists(output_path):
                        output_path = base_path.replace('.docx', f'_{counter}.docx')
                        counter += 1
                    
                    # Hybrid Approach for Batch
                    temp_path = output_path.replace('.docx', '_temp.docx')
                    doc.write(temp_path)
                    
                    try:
                        doc_final = Document(temp_path)
                        self._mm_manual_replace(doc_final, context)
                        doc_final.save(output_path)
                    except Exception as e:
                        # Fallback
                        if os.path.exists(output_path): os.remove(output_path)
                        os.rename(temp_path, output_path)
                        
                    # Cleanup temp
                    if os.path.exists(temp_path):
                        try: os.remove(temp_path)
                        except: pass
                    
                    self._mm_log(f"  [{idx+1}/{total_rows}] ✅ {unique_key}\n")
                    success_count += 1
                    
                    # Update UI occasionally
                    if idx % 5 == 0:
                        self.update_idletasks()
                    
                except Exception as e:
                    self._mm_log(f"  [{idx+1}/{total_rows}] ❌ Error: {e}\n")
                    error_count += 1
            
            # Summary
            self._mm_log("\n" + "=" * 60 + "\n")
            self._mm_log(f"📊 Batch Generation Complete!\n")
            self._mm_log(f"   ✅ Successful: {success_count}\n")
            self._mm_log(f"   ❌ Errors: {error_count}\n")
            self._mm_log(f"   📁 Output: {output_dir}\n\n")
            
            messagebox.showinfo(
                "Batch Generation Complete",
                f"Batch generation completed!\n\n"
                f"Successful: {success_count}\n"
                f"Errors: {error_count}\n\n"
                f"Documents saved to:\n{output_dir}"
            )
            
        except Exception as e:
            self._mm_log(f"❌ Error in batch generation: {e}\n")
            messagebox.showerror("Error", f"Batch generation failed:\n{e}")

    def _mm_replace_merge_fields(self, doc, row_data):
        """Replace merge fields in document with data from row"""
        replaced_count = 0
        
        # Determine field mapping
        field_map = {}
        for field in self.mm_detected_fields:
            if field in self.mm_field_mapping:
                # Use manual mapping
                excel_col = self.mm_field_mapping[field]
                if excel_col in row_data:
                    field_map[field] = row_data[excel_col]
            elif field in row_data:
                # Use auto-mapping (exact match)
                field_map[field] = row_data[field]
            elif self.mm_auto_mapping:
                # Try case-insensitive match
                for col, value in row_data.items():
                    if col.upper() == field.upper():
                        field_map[field] = value
                        break
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for field, value in field_map.items():
                merge_field_patterns = [
                    f"«{field}»",
                    f"<<{field}>>",
                    f"{{{{{field}}}}}",
                    f"{{merge {field}}}"
                ]
                
                for pattern in merge_field_patterns:
                    if pattern in paragraph.text:
                        for run in paragraph.runs:
                            if pattern in run.text:
                                run.text = run.text.replace(pattern, str(value))
                                replaced_count += 1
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field, value in field_map.items():
                        merge_field_patterns = [
                            f"«{field}»",
                            f"<<{field}>>",
                            f"{{{{{field}}}}}",
                            f"{{merge {field}}}"
                        ]
                        
                        for pattern in merge_field_patterns:
                            if pattern in cell.text:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if pattern in run.text:
                                            run.text = run.text.replace(pattern, str(value))
                                            replaced_count += 1
        
        return replaced_count

    def _mm_batch_process(self):
        """Batch process multiple MPD numbers"""
        try:
            # Validation
            if not self.mm_word_path:
                self._mm_log("❌ Please select a Word template first\n")
                messagebox.showwarning("Missing Template", "Please select a Word template file first.")
                return
            
            if self.mm_excel_data is None:
                self._mm_log("❌ Please select an Excel file first\n")
                messagebox.showwarning("Missing Data", "Please select an Excel data file first.")
                return
            
            if 'MPD' not in self.mm_excel_data.columns:
                self._mm_log("❌ 'MPD' column not found in Excel data\n")
                messagebox.showerror("Column Missing", "The Excel file doesn't contain an 'MPD' column.")
                return
            
            # Ask for output directory
            output_dir = filedialog.askdirectory(
                title="Select Output Directory for Batch Processing"
            )
            
            if not output_dir:
                self._mm_log("⚠️ Batch processing cancelled\n")
                return
            
            self._mm_log(f"\n📦 Starting Batch Processing...\n")
            self._mm_log(f"📁 Output Directory: {output_dir}\n")
            self._mm_log("=" * 60 + "\n")
            
            # Get all unique MPD values
            mpd_list = self.mm_excel_data['MPD'].dropna().unique().tolist()
            total_mpds = len(mpd_list)
            
            self._mm_log(f"📊 Found {total_mpds} unique MPD entries\n")
            self._mm_log("🚀 Generating documents...\n\n")
            
            success_count = 0
            error_count = 0
            
            for idx, mpd_value in enumerate(mpd_list, 1):
                try:
                    # Find matching row
                    matching_rows = self.mm_excel_data[self.mm_excel_data['MPD'].astype(str) == str(mpd_value)]
                    
                    if matching_rows.empty:
                        self._mm_log(f"  [{idx}/{total_mpds}] ⚠️ Skipping {mpd_value} - No data found\n")
                        error_count += 1
                        continue
                    
                    row_data = matching_rows.iloc[0].to_dict()
                    
                    # Load Word document
                    doc = Document(self.mm_word_path)
                    
                    # Replace merge fields
                    for paragraph in doc.paragraphs:
                        for column, value in row_data.items():
                            merge_field_patterns = [
                                f"«{column}»",
                                f"<<{column}>>",
                                f"{{{{{column}}}}}",
                                f"{{merge {column}}}"
                            ]
                            
                            for pattern in merge_field_patterns:
                                if pattern in paragraph.text:
                                    for run in paragraph.runs:
                                        if pattern in run.text:
                                            run.text = run.text.replace(pattern, str(value))
                    
                    # Replace in tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for column, value in row_data.items():
                                    merge_field_patterns = [
                                        f"«{column}»",
                                        f"<<{column}>>",
                                        f"{{{{{column}}}}}",
                                        f"{{merge {column}}}"
                                    ]
                                    
                                    for pattern in merge_field_patterns:
                                        if pattern in cell.text:
                                            for paragraph in cell.paragraphs:
                                                for run in paragraph.runs:
                                                    if pattern in run.text:
                                                        run.text = run.text.replace(pattern, str(value))
                    
                    # Save document
                    safe_mpd = str(mpd_value).replace('/', '_').replace('\\', '_')
                    output_path = os.path.join(output_dir, f"RC_Card_{safe_mpd}.docx")
                    doc.save(output_path)
                    
                    self._mm_log(f"  [{idx}/{total_mpds}] ✅ {mpd_value} - Generated successfully\n")
                    success_count += 1
                    
                except Exception as e:
                    self._mm_log(f"  [{idx}/{total_mpds}] ❌ {mpd_value} - Error: {e}\n")
                    error_count += 1
            
            # Summary
            self._mm_log("\n" + "=" * 60 + "\n")
            self._mm_log(f"📊 Batch Processing Complete!\n")
            self._mm_log(f"   ✅ Successful: {success_count}\n")
            self._mm_log(f"   ❌ Errors: {error_count}\n")
            self._mm_log(f"   📁 Output: {output_dir}\n\n")
            
            messagebox.showinfo(
                "Batch Processing Complete",
                f"Batch processing completed!\n\n"
                f"Successful: {success_count}\n"
                f"Errors: {error_count}\n\n"
                f"Documents saved to:\n{output_dir}"
            )
            
        except Exception as e:
            self._mm_log(f"❌ Error in batch processing: {e}\n")
            messagebox.showerror("Error", f"Batch processing failed:\n{e}")

# ================== HELPER FUNCTIONS ==================

def df_to_tree(tree: ttk.Treeview, df: pd.DataFrame, max_rows=None):
    """
    Populate a Tkinter Treeview widget with data from a pandas DataFrame.
    
    Args:
        tree: The Treeview widget to populate
        df: The DataFrame containing the data
        max_rows: Maximum number of rows to display (None for all rows)
    """
    # Clear existing items
    for item in tree.get_children():
        tree.delete(item)
        
    if df.empty:
        return
        
    # Set columns
    tree["columns"] = list(df.columns)
    tree["show"] = "tree headings"
    
    for col in df.columns:
        tree.heading(col, text=col)
        tree.column(col, width=100)
        
    # Add rows
    limit = max_rows if max_rows else len(df)
    for i, row in df.head(limit).iterrows():
        tree.insert("", "end", values=list(row))

# ================== MAIN EXECUTION ==================

def parse_arguments():
    """Professional command line argument parsing"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='REDSEA Airlines Aviation Maintenance Toolkit',
        epilog='Professional maintenance management system for aviation industry'
    )
    
    parser.add_argument(
        '--full', 
        action='store_true', 
        help='Start directly in full system mode (bypasses mini launcher)'
    )
    
    parser.add_argument(
        '--mini', 
        action='store_true', 
        help='Force start in mini launcher mode (default behavior)'
    )
    
    parser.add_argument(
        '--module', 
        type=str, 
        choices=['task_extract', 'task_stamping', 'effectivity', 'check_control', 'utilization', 'cmp_tcm', 'mail_merge'],
        help='Open specific module directly in full system mode'
    )
    
    return parser.parse_args()

if __name__ == "__main__":
    try:
        # Parse command line arguments
        args = parse_arguments()
        
        # Configure startup mode (default is mini launcher)
        if args.full:
            # Force full system mode
            MINI_LAUNCHER_MODE = False
        elif args.mini:
            # Force mini launcher mode (redundant but explicit)
            MINI_LAUNCHER_MODE = True
        # Default is mini launcher mode (MINI_LAUNCHER_MODE = True)
        
        if args.module:
            # Specific module requested - use full system
            TARGET_MODULE = args.module
            MINI_LAUNCHER_MODE = False
        
        # Start professional application
        print("Starting REDSEA Airlines Aviation Maintenance Toolkit...")
        print(f"Mode: {'Mini Launcher' if MINI_LAUNCHER_MODE else 'Full System'}")
        if TARGET_MODULE:
            print(f"Target Module: {TARGET_MODULE}")
        
        # Create and run application
        app = RedseaApp()
        app.mainloop()
        
    except KeyboardInterrupt:
        print("\nApplication interrupted by user")
        sys.exit(0)
    except Exception as e:
        error_msg = f"Critical application error: {e}"
        print(error_msg)
        try:
            messagebox.showerror("Application Error", error_msg)
        except:
            pass
        sys.exit(1)